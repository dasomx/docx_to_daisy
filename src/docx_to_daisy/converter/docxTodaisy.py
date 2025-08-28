import zipfile
import os
import uuid
import argparse
import re
import logging
import html
import time
from docx import Document  # python-docx 라이브러리
from docx.oxml.ns import qn  # XML 네임스페이스 처리
from lxml import etree  # lxml 라이브러리
from datetime import datetime
from docx_to_daisy.markers import MarkerProcessor  # 마커 처리기 임포트
import gc
from docx.table import Table as DocxTable
from docx_to_daisy.converter.utils import find_all_images, split_text_to_words, analyze_image_context, html_escape, BR_PATTERN, TABLE_TITLE_PATTERN
from docx_to_daisy.converter.validator import DaisyValidator, ValidationResult

# 로깅 설정
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def create_daisy_book(docx_file_path, output_dir, book_title=None, book_author=None, book_publisher=None, book_language="ko"):
    """DOCX 파일을 DAISY 형식으로 변환합니다.

    Args:
        docx_file_path (str): 변환할 DOCX 파일 경로
        output_dir (str): 생성된 DAISY 파일 저장 폴더
        book_title (str, optional): 책 제목. 기본값은 None (DOCX 파일명 사용)
        book_author (str, optional): 저자. 기본값은 None
        book_publisher (str, optional): 출판사. 기본값은 None
        book_language (str, optional): 언어 코드 (ISO 639-1). 기본값은 "ko"
    """
    # 이미지 설명 처리 함수 정의
    def get_clean_description(desc_list):
        """이미지 설명 목록에서 중복을 제거하고 깔끔한 설명을 반환합니다"""
        if not desc_list:
            return None
        
        # 중복 제거
        unique_desc = []
        seen = set()
        for desc in desc_list:
            if desc not in seen:
                seen.add(desc)
                unique_desc.append(desc)
        
        # 설명이 너무 길면 첫 번째 항목만 반환
        if len(unique_desc) > 0:
            return unique_desc[0]
        return None
    
    # --- 출력 디렉토리 생성 ---
    timings = {}
    t0 = time.time()
    os.makedirs(output_dir, exist_ok=True)
    timings["init_output_dir"] = time.time() - t0

    # --- DOCX 파일 읽기 및 구조 분석 ---
    try:
        t0 = time.time()
        document = Document(docx_file_path)
        timings["load_docx"] = time.time() - t0
    except FileNotFoundError:
        print(f"오류: DOCX 파일을 찾을 수 없습니다 - {docx_file_path}")
        return None
    except Exception as e:
        print(f"오류: DOCX 파일을 읽는 중 오류가 발생했습니다 - {str(e)}")
        return None

    # --- 기본 정보 설정 ---
    # book_title 확인
    t_validate = time.time()
    if book_title is None or not isinstance(book_title, str) or len(book_title.strip()) == 0:
        raise ValueError("책 제목이 제공되지 않았거나 유효하지 않습니다. 변환을 진행할 수 없습니다.")
    
    # book_author 확인
    if book_author is None or not isinstance(book_author, str) or len(book_author.strip()) == 0:
        raise ValueError("저자 정보가 제공되지 않았거나 유효하지 않습니다. 변환을 진행할 수 없습니다.")
    timings["validate_metadata"] = time.time() - t_validate

    book_title = str(book_title)
    book_author = str(book_author)
    book_publisher = str(book_publisher) if book_publisher else "Unknown Publisher"
    
    # book_uid에 책 제목을 포함시켜 DAISY 열 때 표시될 이름 설정
    safe_title = book_title.replace(" ", "_").replace("/", "_").replace("\\", "_").replace(":", "_")
    
    book_uid = f"BOOK-{safe_title}"  # 제목을 포함한 식별자
    print("book_uid", book_uid)

    content_structure = []
    element_counter = 0
    sent_counter = 0

    # doctitle과 docauthor를 위한 ID 미리 할당
    doctitle_id = "id_1"
    docauthor_id = "id_2"
    sent_counter = 2  # doctitle과 docauthor 이후부터 시작
    element_counter = 2  # doctitle과 docauthor 이후부터 시작

    # 이미지 관련 정보 저장 변수
    image_info = {}  # 이미지 번호 -> {제목, 설명, 위치} 매핑
    
    # --- 포함관계 인덱싱/캐싱 구조 추가 ---
    # 1. 단락, 표, 셀의 id를 key로 하여 포함관계, 인덱스, 부모 정보를 dict로 저장
    paragraph_id_map = {}
    table_id_map = {}
    cell_id_map = {}
    element_parent_map = {}
    element_type_map = {}
    
    # 단락 인덱싱
    for para_idx, para in enumerate(document.paragraphs):
        pid = id(para._element)
        paragraph_id_map[pid] = {
            'index': para_idx,
            'object': para,
            'parent': para._element.getparent()
        }
        element_type_map[pid] = 'paragraph'
        element_parent_map[pid] = para._element.getparent()
    # 표/셀 인덱싱 (중첩 표까지 포함)
    def index_table(table, table_idx_hint=0):
        tid = id(table._element)
        table_id_map[tid] = {
            'index': table_idx_hint,
            'object': table,
            'parent': table._element.getparent()
        }
        element_type_map[tid] = 'table'
        element_parent_map[tid] = table._element.getparent()

        # 셀 인덱싱 및 중첩 표 재귀 인덱싱
        for row in table.rows:
            for cell in row.cells:
                cid = id(cell._element)
                cell_id_map[cid] = {
                    'object': cell,
                    'parent': cell._element.getparent(),
                    'table_id': tid
                }
                element_type_map[cid] = 'cell'
                element_parent_map[cid] = cell._element.getparent()

                # 셀 내부 중첩 표 인덱싱
                try:
                    for nested_idx, nested_table in enumerate(getattr(cell, 'tables', []) or []):
                        index_table(nested_table, nested_idx)
                except Exception:
                    pass

    for table_idx, table in enumerate(document.tables):
        index_table(table, table_idx)

    # 2. 이미지가 어느 표/셀/단락에 속하는지 O(1)로 판별할 수 있도록 parent chain을 미리 저장
    def get_ancestor_type_and_id(element):
        """element의 조상 중 table/cell/paragraph를 찾아 반환"""
        current = element
        while current is not None:
            cid = id(current)
            if cid in cell_id_map:
                return ('cell', cid)
            if cid in table_id_map:
                return ('table', cid)
            if cid in paragraph_id_map:
                return ('paragraph', cid)
            current = current.getparent() if hasattr(current, 'getparent') else None
        return (None, None)

    # 이미지의 실제 문서 내 위치를 계산하는 함수
    def get_image_document_position(para_idx, run_idx):
        """이미지의 실제 문서 내 위치를 계산합니다.
        단락 인덱스와 런 인덱스를 기반으로 정확한 위치를 반환합니다."""
        # 기본 위치는 단락 순서
        base_position = para_idx * 1000  # 단락별로 1000 단위로 구분
        
        # 런 인덱스를 더해서 같은 단락 내에서의 순서 결정
        run_position = run_idx
        
        return base_position + run_position

    # 병합된 영역에 속하는지 확인하는 함수
    def is_in_merged_area(row_idx, col_idx, table_data, cell_merge_info, table_id):
        """현재 위치가 다른 셀의 병합 영역에 속하는지 확인합니다."""
        # 해당 위치의 셀 정보 찾기
        for cell in table_data["cells"]:
            if cell["row"] == row_idx and cell["col"] == col_idx:
                # 병합 정보에서 확인
                cid = cell.get("cell_id")
                if cid and table_id in cell_merge_info and cid in cell_merge_info[table_id]:
                    merge_info = cell_merge_info[table_id][cid]
                    # 병합 영역에 속하지만 시작점이 아닌 경우
                    return merge_info.get('is_merged_area', False)
        
        return False

    # 1. 문서에서 모든 이미지 찾기 (표 안/밖 분류)
    t_images = time.time()
    print("문서에서 이미지 찾는 중...")
    images = find_all_images(document)
    print(f"총 {len(images)}개의 이미지 발견")

    # 표 안/밖 이미지 분류
    standalone_images = []
    table_images = []
    cell_images = []
    for img in images:
        # 이미지가 속한 최상위 조상 타입/ID 판별
        ancestor_type, ancestor_id = get_ancestor_type_and_id(img['paragraph']._element)
        img['ancestor_type'] = ancestor_type
        img['ancestor_id'] = ancestor_id
        if ancestor_type == 'cell':
            cell_images.append(img)
        elif ancestor_type == 'table':
            table_images.append(img)
        else:
            standalone_images.append(img)
    print(f"표 밖의 이미지: {len(standalone_images)}개")
    print(f"표 안의 이미지: {len(table_images)}개, 셀 안의 이미지: {len(cell_images)}개")

    # 이미지 조회 성능 최적화를 위한 사전 매핑 (결과 동일성 유지)
    cell_images_map = {}
    for _img in cell_images:
        cell_images_map.setdefault(_img['ancestor_id'], []).append(_img)

    table_images_map = {}
    for _img in table_images:
        table_images_map.setdefault(_img['ancestor_id'], []).append(_img)
    
    # 2. 문서에서 모든 이미지 관계 미리 수집 (성능 최적화)
    print(f"문서에서 이미지 관계 수집 중...")
    image_counter = 0
    image_relations = {}
    
    # 모든 이미지 관계를 딕셔너리로 미리 수집 (O(1) 접근)
    for rel_id, rel in document.part.rels.items():
        if "image" in rel.reltype:
            try:
                image_relations[rel_id] = rel
                print(f"이미지 관계 발견: {rel_id}, {rel.reltype}")
            except Exception as e:
                print(f"이미지 관계 처리 오류: {str(e)}")
    
    print(f"문서에서 {len(image_relations)}개의 이미지 관계 발견")
    
    # 이미지 매핑 정보 초기화
    image_mapping = {}  # 이미지 번호 -> 이미지 관계 매핑

    # 이미지 처리 - 실제 문서 위치 기반으로 처리
    for i, img in enumerate(standalone_images, 1):
        img_num = str(i)
        try:
            image_counter += 1
            element_counter += 1
            sent_counter += 1
            
            # 이미지 ID 생성
            elem_id = f"id_{element_counter}"
            sent_id = f"id_{sent_counter}"
            
            # 이미지 파일 저장
            image_ext = ".jpeg"
            try:
                if 'image_rid' in img:
                    rel = image_relations.get(img['image_rid'])  # 미리 수집한 관계 사용
                    if rel and hasattr(rel, 'target_ref'):
                        ext = os.path.splitext(rel.target_ref)[1]
                        if ext:
                            image_ext = ext
            except:
                pass
            
            image_filename = f"image{img_num}{image_ext}"
            image_path = os.path.join(output_dir, image_filename)
            
            # 이미지 데이터 저장(큰 버퍼)
            with open(image_path, "wb", buffering=1<<20) as img_file:
                img_file.write(img['image_data'])
            # 상세 콘솔 출력 제거 (성능)
            
            # 이미지 정보를 content_structure에 추가 - 실제 문서 위치 사용
            document_position = get_image_document_position(img['paragraph_index'], img['run_index'])
            content_structure.append({
                "type": "image",
                "src": image_filename,
                "alt_text": f"이미지 {img_num}",
                "id": elem_id,
                "sent_id": sent_id,
                "level": 0,
                "markers": [],
                "position": document_position,
                "image_number": i,  # 이미지 번호 추가
                "insert_before": False,
                "para_index": img['paragraph_index'],  # 단락 인덱스 추가
                "run_index": img['run_index']  # 런 인덱스 추가
            })
            # 상세 콘솔 출력 제거 (성능)
        except Exception as e:
            print(f"이미지 {img_num} 처리 중 오류 발생: {str(e)}")
            continue

    print(f"{image_counter}개 이미지 추출 완료.")
    timings["extract_images"] = time.time() - t_images

    # 메모리 정리 (images는 표 처리에서도 사용하므로 유지)
    del image_relations
    del image_mapping
    gc.collect()

    # DOCX의 단락(paragraph)을 순회하며 구조 파악
    t_paragraphs = time.time()
    print("DOCX 파일 분석 중...")
    print(f"총 {len(document.paragraphs)}개의 단락을 처리합니다.")
    
    # 단락 처리
    for para_idx, para in enumerate(document.paragraphs):
        # 진행 상황 로그 (100개 단락마다)
        if para_idx % 100 == 0:
            print(f"단락 처리 진행 중: {para_idx}/{len(document.paragraphs)} ({para_idx/len(document.paragraphs)*100:.1f}%)")
        
        text_raw = para.text
        style_name = para.style.name.lower()

        # <br/> 태그 기준으로 세그먼트를 분리 (성능 최적화)
        br_segments = BR_PATTERN.split(text_raw)

        # 세그먼트별 처리
        for seg_idx, seg_text in enumerate(br_segments):
            # <br/> 태그가 있었던 자리에 빈 문단을 생성
            if seg_idx > 0:
                element_counter += 1
                sent_counter += 1
                blank_elem_id = f"p_{element_counter}"
                blank_sent_id = f"sent_{sent_counter}"
                content_structure.append({
                    "type": "p",
                    "text": "<br/>",
                    "words": ["<br/>"],
                    "id": blank_elem_id,
                    "sent_id": blank_sent_id,
                    "level": 0,
                    "markers": [],
                    "position": para_idx * 1000 + seg_idx * 100,  # 단락 내 세그먼트 순서
                    "insert_before": False
                })

            # 세그먼트 자체가 비어 있으면(공백만) 넘어감
            if not seg_text.strip():
                continue

            # 마커 처리
            processed_text, markers = MarkerProcessor.process_text(seg_text.strip())

            # 페이지 마커가 있는 경우 먼저 처리
            page_markers = [marker for marker in markers if marker.type == "page"]
            for marker in page_markers:
                element_counter += 1
                sent_counter += 1
                elem_id = f"page_{element_counter}"
                sent_id = f"sent_{sent_counter}"
                content_structure.append({
                    "type": "pagenum",
                    "text": marker.value,
                    "words": [marker.value],
                    "id": elem_id,
                    "sent_id": sent_id,
                    "level": 0,
                    "markers": [marker],
                    "position": para_idx * 1000 + seg_idx * 100 + 50,  # 페이지 마커는 세그먼트보다 먼저
                    "insert_before": True
                })

                # 마커만 있고 실제 내용이 없는 경우 건너뜀
                if not processed_text.strip():
                    continue

            # 페이지 마커는 별도 요소로 처리했으므로 본문 마커 목록에서는 제거
            markers = [m for m in markers if m.type != "page"]

            element_counter += 1
            sent_counter += 1
            elem_id = f"p_{element_counter}"
            sent_id = f"sent_{sent_counter}"

            # 단어 분리
            words = split_text_to_words(processed_text)

            # 스타일 이름에 따른 구조 매핑 (견고하게 처리)
            element_type = "p"
            element_level = 0

            # 스타일 이름 정규화 및 다양한 변형 대응
            normalized_style = (style_name or "").strip().lower()
            normalized_style = re.sub(r"\s+", " ", normalized_style)
            # 예: "Heading 1", "heading1", "제목 1", "제목1", "Heading 2 + Bold", 등
            heading_match = re.search(r"(?:heading|제목)\s*([1-6])\b", normalized_style)
            if heading_match:
                level_num = int(heading_match.group(1))
                element_type = f"h{level_num}"
                element_level = level_num

            # 스타일 이름에 따른 구조 매핑
            content_structure.append({
                "type": element_type,
                "text": processed_text,
                "words": words,
                "id": elem_id,
                "sent_id": sent_id,
                "level": element_level,
                "markers": markers,
                "position": para_idx * 1000 + seg_idx * 100 + 100,  # 텍스트는 세그먼트 내에서 마지막
                "insert_before": False
            })
    
    print(f"단락 처리 완료: 총 {len(content_structure)}개의 구조 요소 생성")
    timings["parse_paragraphs"] = time.time() - t_paragraphs

    # ======== 표 처리 ========
    t_tables = time.time()
    print("표 처리 중...")
    
    # body_children에서 모든 요소의 위치를 미리 계산 (성능 최적화)
    body_children = list(document._element.body.iterchildren())
    element_positions = {}
    
    # 모든 요소의 위치를 미리 계산
    for idx, child in enumerate(body_children):
        child_id = id(child)
        element_positions[child_id] = idx
    
    # 단락 인덱스 매핑을 미리 계산 (성능 최적화)
    paragraph_index_map = {}
    for para_idx, para in enumerate(document.paragraphs):
        para_id = id(para._element)
        paragraph_index_map[para_id] = para_idx
    
    if len(document.tables) > 0:
        print(f"문서에 {len(document.tables)}개의 표 발견")
        
        # 병합 정보 계산 및 표 데이터 추출 헬퍼
        def compute_cell_merge_info_for_table(table_obj):
            """주어진 표에 대한 병합 정보를 계산하여 반환합니다."""
            info_map = {}

            # 1) 실제 XML tc 기준으로 각 행의 그리드 세그먼트 수집
            row_segments = []
            for row_idx, row in enumerate(table_obj.rows):
                segments = []
                g = 0
                try:
                    tc_list = list(row._tr.tc_lst)
                except Exception:
                    tc_list = []
                    seen = set()
                    for c in row.cells:
                        if id(c._tc) not in seen:
                            tc_list.append(c._tc)
                            seen.add(id(c._tc))

                for tc in tc_list:
                    try:
                        tc_pr = tc.get_or_add_tcPr()
                    except Exception:
                        tc_pr = getattr(tc, 'tcPr', None)
                    colspan = 1
                    try:
                        grid_span = tc_pr.find(qn('w:gridSpan')) if tc_pr is not None else None
                        if grid_span is not None and grid_span.get(qn('w:val')):
                            colspan = int(grid_span.get(qn('w:val')))
                    except Exception:
                        pass
                    v_state = None
                    try:
                        v_merge = tc_pr.find(qn('w:vMerge')) if tc_pr is not None else None
                        if v_merge is not None:
                            v_val = v_merge.get(qn('w:val'))
                            if v_val == 'restart':
                                v_state = 'restart'
                            else:
                                v_state = 'continue'
                    except Exception:
                        pass

                    seg = {
                        'grid_start': g,
                        'grid_end': g + colspan - 1,
                        'colspan': colspan,
                        'v_state': v_state
                    }
                    segments.append(seg)
                    g += colspan

                row_segments.append(segments)

            # 2) top-left 매핑, rowspan/colspan 계산
            top_left_of = {}
            width_map = {}
            rowspan_map = {}
            for r, segs in enumerate(row_segments):
                for seg in segs:
                    gs, ge = seg['grid_start'], seg['grid_end']
                    for gg in range(gs, ge + 1):
                        top_left_of[(r, gg)] = (r, gs)
                    width_map[(r, gs)] = seg['colspan']

            for r0, segs in enumerate(row_segments):
                for seg in segs:
                    if seg['v_state'] == 'restart':
                        gs, ge = seg['grid_start'], seg['grid_end']
                        rs = 1
                        rr = r0 + 1
                        while rr < len(row_segments):
                            target = None
                            for s in row_segments[rr]:
                                if s['grid_start'] <= gs <= s['grid_end']:
                                    target = s
                                    break
                            if target is None:
                                break
                            if target['v_state'] == 'continue':
                                rs += 1
                                for gg in range(gs, ge + 1):
                                    top_left_of[(rr, gg)] = (r0, gs)
                                rr += 1
                            else:
                                break
                        rowspan_map[(r0, gs)] = rs

            # 3) (row_idx, col_idx)별 병합 정보 생성
            for r, row in enumerate(table_obj.rows):
                num_cols = len(row.cells)
                for c in range(num_cols):
                    tl = top_left_of.get((r, c), (r, c))
                    if tl == (r, c):
                        colspan = width_map.get(tl, 1)
                        rowspan = rowspan_map.get(tl, 1)
                        info_map[(r, c)] = {
                            'rowspan': rowspan,
                            'colspan': colspan,
                            'is_merged': (rowspan > 1 or colspan > 1),
                            'is_merged_area': False
                        }
                    else:
                        info_map[(r, c)] = {
                            'rowspan': 1,
                            'colspan': 1,
                            'is_merged': False,
                            'is_merged_area': True
                        }
            return info_map

        def extract_table_data_from_xml(tbl_el):
            """w:tbl Element(XML)에서 표 데이터를 직접 추출합니다. 문단/중첩표 순서 보존 및 병합 반영."""
            W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            data = {
                'rows': [],
                'cols': [],
                'cells': [],
                'images': []
            }

            # 1) 병합 정보 계산
            row_segments = []
            for tr in tbl_el.iterfind('./{%s}tr' % W_NS):
                segments = []
                g = 0
                tcs = list(tr.iterfind('./{%s}tc' % W_NS))
                for tc in tcs:
                    tc_pr = tc.find('{%s}tcPr' % W_NS)
                    colspan = 1
                    v_state = None
                    if tc_pr is not None:
                        grid_span = tc_pr.find('{%s}gridSpan' % W_NS)
                        if grid_span is not None:
                            val = grid_span.get('{%s}val' % W_NS)
                            if val:
                                try:
                                    colspan = int(val)
                                except Exception:
                                    colspan = 1
                        v_merge = tc_pr.find('{%s}vMerge' % W_NS)
                        if v_merge is not None:
                            v_val = v_merge.get('{%s}val' % W_NS)
                            if v_val == 'restart':
                                v_state = 'restart'
                            else:
                                v_state = 'continue'
                    seg = {'grid_start': g, 'grid_end': g + colspan - 1, 'colspan': colspan, 'v_state': v_state}
                    segments.append(seg)
                    g += colspan
                row_segments.append(segments)

            top_left_of = {}
            width_map = {}
            rowspan_map = {}
            for r, segs in enumerate(row_segments):
                for seg in segs:
                    gs, ge = seg['grid_start'], seg['grid_end']
                    for gg in range(gs, ge + 1):
                        top_left_of[(r, gg)] = (r, gs)
                    width_map[(r, gs)] = seg['colspan']
            for r0, segs in enumerate(row_segments):
                for seg in segs:
                    if seg['v_state'] == 'restart':
                        gs, ge = seg['grid_start'], seg['grid_end']
                        rs = 1
                        rr = r0 + 1
                        while rr < len(row_segments):
                            target = None
                            for s in row_segments[rr]:
                                if s['grid_start'] <= gs <= s['grid_end']:
                                    target = s
                                    break
                            if target is None:
                                break
                            if target['v_state'] == 'continue':
                                rs += 1
                                for gg in range(gs, ge + 1):
                                    top_left_of[(rr, gg)] = (r0, gs)
                                rr += 1
                            else:
                                break
                        rowspan_map[(r0, gs)] = rs

            # 2) 행/셀 수집 (문단/표 순서 보존)
            for r_idx, tr in enumerate(tbl_el.iterfind('./{%s}tr' % W_NS)):
                row_texts = []
                tcs = list(tr.iterfind('./{%s}tc' % W_NS))
                for c_idx, tc in enumerate(tcs):
                    # 문단 텍스트 나열 (row_texts용)
                    p_texts = []
                    for p in tc.iterfind('./{%s}p' % W_NS):
                        p_texts.append(''.join((t.text or '') for t in p.iterfind('.//{%s}t' % W_NS)))
                    row_texts.append(' '.join(p_texts))

                    # 셀 병합 정보
                    tl = top_left_of.get((r_idx, c_idx), (r_idx, c_idx))
                    if tl == (r_idx, c_idx):
                        colspan = width_map.get(tl, 1)
                        rowspan = rowspan_map.get(tl, 1)
                        is_merged = (rowspan > 1 or colspan > 1)
                        is_merged_area = False
                    else:
                        colspan = 1
                        rowspan = 1
                        is_merged = False
                        is_merged_area = True

                    # 셀 content_sequence 구성
                    content_sequence = []
                    for child in tc.iterchildren():
                        lname = etree.QName(child).localname if hasattr(etree, 'QName') else child.tag.split('}')[-1]
                        if lname == 'p':
                            text_val = ''.join((t.text or '') for t in child.iterfind('.//{%s}t' % W_NS))
                            content_sequence.append({'type': 'p', 'text': text_val})
                        elif lname == 'tbl':
                            nested_data = extract_table_data_from_xml(child)
                            content_sequence.append({'type': 'table', 'table_data': nested_data})

                    data['cells'].append({
                        'row': r_idx,
                        'col': c_idx,
                        'text': row_texts[-1] if row_texts else '',
                        'is_merged': is_merged,
                        'rowspan': rowspan,
                        'colspan': colspan,
                        'is_merged_area': is_merged_area,
                        'cell_id': None,
                        'images': [],
                        'paragraphs': p_texts,
                        'content_sequence': content_sequence
                    })
                data['rows'].append(row_texts)

            # 3) 열 데이터
            if data['rows']:
                num_cols = max(len(r) for r in data['rows'])
                for col_idx in range(num_cols):
                    col_vals = []
                    for rr in data['rows']:
                        if col_idx < len(rr):
                            col_vals.append(rr[col_idx])
                    data['cols'].append(col_vals)

            return data

        def extract_table_data(table_obj):
            """표 데이터를 추출하고, 셀 내부의 중첩 표도 재귀적으로 추출합니다."""
            data = {
                'rows': [],
                'cols': [],
                'cells': [],
                'images': []
            }
            merge_info_map = compute_cell_merge_info_for_table(table_obj)

            # 행/셀 데이터
            for row_idx, row in enumerate(table_obj.rows):
                row_data = []
                for col_idx, cell in enumerate(row.cells):
                    para_texts = [para.text for para in cell.paragraphs]
                    cell_text = " ".join(para_texts)
                    row_data.append(cell_text)

                    cid = id(cell._element)
                    cell_images_in_cell = cell_images_map.get(cid, [])

                    mi = merge_info_map.get((row_idx, col_idx), {
                        'rowspan': 1,
                        'colspan': 1,
                        'is_merged': False,
                        'is_merged_area': False
                    })

                    # 셀 내 콘텐츠 순서 보존: 단순 XML 스캔 (p/tbl, sdt 포함)
                    content_sequence = []
                    try:
                        W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                        logger.info(f"[InnerTable] 셀({row_idx},{col_idx}) 블록 스캔 시작")
                        p_count = 0
                        tbl_count = 0
 
                        def iter_block_children(tc_el):
                            for child in tc_el.iterchildren():
                                lname = etree.QName(child).localname if hasattr(etree, 'QName') else child.tag.split('}')[-1]
                                if lname in ('p', 'tbl'):
                                    yield child
                                elif lname == 'sdt':
                                    sdt_content = child.find('.//{%s}sdtContent' % W_NS)
                                    if sdt_content is not None:
                                        for sub in sdt_content.iterchildren():
                                            lname2 = etree.QName(sub).localname if hasattr(etree, 'QName') else sub.tag.split('}')[-1]
                                            if lname2 in ('p', 'tbl'):
                                                yield sub

                        prev_p_text = None
                        for child in iter_block_children(cell._element):
                            localname = etree.QName(child).localname if hasattr(etree, 'QName') else child.tag.split('}')[-1]
                            if localname == 'p':
                                text_val = ''.join((t.text or '') for t in child.iterfind('.//{%s}t' % W_NS))
                                # 연속 중복 문단 제거
                                if prev_p_text is None or prev_p_text != text_val:
                                    content_sequence.append({'type': 'p', 'text': text_val})
                                    p_count += 1
                                    prev_p_text = text_val
                            elif localname == 'tbl':
                                logger.info(f"[InnerTable] 셀({row_idx},{col_idx}) 내부표 감지 (XML)")
                                nested_data = extract_table_data_from_xml(child)
                                content_sequence.append({'type': 'table', 'table_data': nested_data})
                                tbl_count += 1

                        # 표 총계 기록 (직속 및 sdtContent 내부)
                        extra_tbls = list(cell._element.iterfind('./{%s}tbl' % W_NS))
                        for sdt in cell._element.iterfind('.//{%s}sdtContent' % W_NS):
                            extra_tbls.extend(list(sdt.iterfind('./{%s}tbl' % W_NS)))
                        logger.info(f"[InnerTable] 셀({row_idx},{col_idx}) 완료: p={p_count}, tbl={tbl_count}, tbl(xpath)={len(extra_tbls)}")

                    except Exception as e:
                        logger.info(f"[InnerTable] 셀({row_idx},{col_idx}) 스캔 예외: {e}")
                        for t in para_texts:
                            content_sequence.append({'type': 'p', 'text': t})

                    data['cells'].append({
                        'row': row_idx,
                        'col': col_idx,
                        'text': cell_text,
                        'is_merged': mi['is_merged'],
                        'rowspan': mi['rowspan'],
                        'colspan': mi['colspan'],
                        'is_merged_area': mi['is_merged_area'],
                        'cell_id': cid,
                        'images': cell_images_in_cell,
                        'paragraphs': para_texts,
                        'content_sequence': content_sequence
                    })

                data['rows'].append(row_data)

            # 열 데이터
            for col_idx in range(len(table_obj.columns)):
                col_data = []
                for row in table_obj.rows:
                    if col_idx < len(row.cells):
                        cell_text = " ".join(para.text for para in row.cells[col_idx].paragraphs)
                        col_data.append(cell_text)
                data['cols'].append(col_data)

            return data
        
        for table_idx, table in enumerate(document.tables, 1):
            print(f"표 {table_idx} 처리 중...")
            element_counter += 1
            sent_counter += 1
            elem_id = f"table_{element_counter}"
            sent_id = f"sent_{sent_counter}"
            
            # 표 데이터 추출 (중첩 표 포함)
            table_data = extract_table_data(table)
            
            # 표 안의 이미지 찾기 (인덱싱 구조 활용)
            tid = id(table._element)
            this_table_images = table_images_map.get(tid, [])
            
            print(f"표 {table_idx} 안에 {len(this_table_images)}개의 이미지 발견")
            
            # (행/열 정보는 extract_table_data에서 수집됨)
            
            # 표 위치 계산 (개선된 로직)
            table_document_position = 0
            tid = id(table._element)
            table_body_index = element_positions.get(tid, -1)
            
            if table_body_index != -1:
                # 표 이전의 단락들을 확인하여 가장 큰 단락 인덱스 찾기
                max_para_before_table = -1
                for idx in range(table_body_index):
                    child = body_children[idx]
                    if child.tag.endswith('p'):  # 단락인 경우
                        # 미리 계산된 단락 인덱스 매핑 사용
                        para_idx = paragraph_index_map.get(id(child), -1)
                        if para_idx != -1:
                            max_para_before_table = max(max_para_before_table, para_idx)
                
                # 표의 위치를 가장 큰 단락 인덱스 + 1로 설정
                if max_para_before_table >= 0:
                    # 단락과 동일한 스케일로 계산
                    # 단락은 para_idx * 1000 + seg_idx * 100 + 100을 사용
                    # 표는 단락 다음에 오므로 para_idx * 1000 + 500으로 설정
                    table_document_position = (max_para_before_table + 1) * 1000 + 500
                    print(f"표 {table_idx} 위치 계산: 단락 {max_para_before_table} 다음, 위치 {table_document_position}")
                else:
                    # 단락을 찾지 못한 경우 표의 body 내 위치 기반으로 설정
                    table_document_position = table_body_index * 1000 + 500
                    print(f"표 {table_idx} 위치 계산: 단락 없음, body 위치 {table_body_index}, 위치 {table_document_position}")
            else:
                # 표 요소를 찾지 못한 경우 표 인덱스 기반으로 설정
                table_document_position = table_idx * 1000 + 500
                print(f"표 {table_idx} 위치 계산: body 위치 없음, 표 인덱스 {table_idx}, 위치 {table_document_position}")
            
            print(f"표 {table_idx} 문서 위치: {table_document_position}")
            
            table_title = f"표 {table_idx}"
            
            # 표 정보를 content_structure에 추가 (caption 요소 포함)
            content_structure.append({
                "type": "table",
                "table_data": table_data,
                "id": elem_id,
                "sent_id": sent_id,
                "level": 0,
                "markers": [],
                "position": table_document_position,
                "insert_before": False,
                "title": table_title,
                "table_number": table_idx,
                "text": table_title,
                "original_table_id": tid  # 원본 테이블 ID 추가
            })
            
            print(f"표 {table_idx} 처리 완료: {len(table_data['rows'])}행 x {len(table_data['cols'])}열, 문서 위치: {table_document_position}")
    else:
        print("문서에 표가 없습니다.")
    timings["process_tables"] = time.time() - t_tables

    # 메모리 정리 (모든 처리 완료 후)
    # 변수가 정의된 경우에만 삭제
    if 'images' in locals():
        del images
    if 'table_images' in locals():
        del table_images
    if 'cell_images' in locals():
        del cell_images
    if 'cell_merge_info' in locals():
        del cell_merge_info
    if 'paragraph_id_map' in locals():
        del paragraph_id_map
    if 'table_id_map' in locals():
        del table_id_map
    if 'cell_id_map' in locals():
        del cell_id_map
    if 'element_parent_map' in locals():
        del element_parent_map
    if 'element_type_map' in locals():
        del element_type_map
    # image_relations는 이후 OPF, DTBook 단계에서도 사용되므로 유지
    gc.collect()

    # 콘텐츠를 위치에 따라 정렬 - 이미지와 텍스트의 정확한 순서 보장
    t_sort = time.time()
    content_structure.sort(key=lambda x: (
        x["position"],  # 기본 위치 (단락 순서)
        x.get("run_index", 0) if x["type"] == "image" else 0,  # 이미지의 경우 런 인덱스 고려
        not x["insert_before"]  # insert_before가 False인 요소가 나중에
    ))

    print(f"총 {len(content_structure)}개의 구조 요소 분석 완료.")
    timings["sort_content"] = time.time() - t_sort

    # --- 1. DTBook XML 생성 (dtbook.xml) ---
    t_dtbook = time.time()
    print("DTBook 생성 중...")
    dtbook_ns = "http://www.daisy.org/z3986/2005/dtbook/"
    dc_ns = "http://purl.org/dc/elements/1.1/"

    # 페이지 카운터 초기화
    total_pages = 0
    max_page_number = 0
    
    # pagenum 요소 기준으로 페이지 정보 계산 (중복 방지)
    page_values = []
    for item in content_structure:
        if item.get("type") == "pagenum":
            val = str(item.get("text", "")).strip()
            if val:
                page_values.append(val)

    total_pages = len(page_values)

    # 최대 숫자 페이지 계산 (로마 숫자 등 비숫자 표기는 제외)
    for value in page_values:
        # 페이지 마커 값이 "1", "0-9", "8.1" 등 다양한 포맷일 수 있으므로
        # 규칙: '-'는 마지막 숫자 세그먼트, '.'는 첫 숫자 세그먼트를 페이지 번호로 본다.
        page_num = None
        if '-' in value:
            nums = re.findall(r'\d+', value)
            if nums:
                page_num = int(nums[-1])
        elif '.' in value:
            nums = re.findall(r'\d+', value)
            if nums:
                page_num = int(nums[0])
        else:
            try:
                page_num = int(value)
            except ValueError:
                page_num = None
        if page_num is not None:
            max_page_number = max(max_page_number, page_num)

    dtbook_root = etree.Element(
        "{%s}dtbook" % dtbook_ns,
        attrib={
            "version": "2005-3"
        },
        nsmap={
            None: dtbook_ns,
            "dc": dc_ns
        }
    )

    # head 요소 추가
    head = etree.SubElement(dtbook_root, "head")

    # 필수 메타데이터 추가
    meta_uid = etree.SubElement(head, "meta",
                                name="dtb:uid",
                                content=book_uid)
    meta_title = etree.SubElement(head, "meta",
                                  name="dc:Title",
                                  content=book_title)
    meta_author = etree.SubElement(head, "meta",
                                   name="dc:Creator",
                                   content=book_author)
    meta_publisher = etree.SubElement(head, "meta",
                                   name="dc:Publisher",
                                   content=book_publisher)
    meta_language = etree.SubElement(head, "meta",
                                     name="dc:Language",
                                     content=book_language)
    meta_date = etree.SubElement(head, "meta",
                                 name="dc:Date",
                                 content=datetime.now().strftime("%Y-%m-%d"))

    # 페이지 관련 메타데이터 추가
    etree.SubElement(head, "meta",
                     name="dtb:totalPageCount",
                     content=str(max_page_number))
    etree.SubElement(head, "meta",
                     name="dtb:maxPageNumber",
                     content=str(max_page_number))

    # book 요소 추가
    dtbook_book = etree.SubElement(dtbook_root, "book", showin="blp")

    # frontmatter 추가
    dtbook_frontmatter = etree.SubElement(dtbook_book, "frontmatter")

    # doctitle과 docauthor 추가
    doctitle_seq = etree.SubElement(dtbook_frontmatter, "doctitle",
                                    id="forsmil-1",
                                    smilref="dtbook.smil#sforsmil-1")
    doctitle_seq.text = book_title

    docauthor = etree.SubElement(dtbook_frontmatter, "docauthor",
                                 id="forsmil-2",
                                 smilref="dtbook.smil#sforsmil-2")
    docauthor.text = book_author

    # frontmatter에는 출판사 정보를 포함하지 않음 (요청사항)

    # bodymatter 추가
    dtbook_bodymatter = etree.SubElement(dtbook_book, "bodymatter")

    # 현재 level1 요소
    current_level1 = None
    current_level = 0

    # 계층 구조 관리를 위한 변수들
    level_elements = {}  # 각 레벨별 현재 요소를 추적
    current_level = 0
    
    # 콘텐츠 추가
    for item in content_structure:
        if item["type"] == "pagenum":
            # 페이지 번호는 현재 활성 레벨 요소에 추가
            parent_elem = level_elements.get(current_level, dtbook_bodymatter)
            pagenum = etree.SubElement(
                parent_elem,
                "pagenum",
                id=f"page_{item['text']}_{item['text']}",
                smilref=f"dtbook.smil#smil_par_page_{item['text']}_{item['text']}",
                page="normal"
            )
            pagenum.text = str(item["text"])
            continue
        elif item["type"] == "image":
            # 이미지는 현재 활성 레벨 요소에 추가
            parent_elem = level_elements.get(current_level, dtbook_bodymatter)
            if parent_elem is dtbook_bodymatter:
                # level1이 없는 경우 생성 (표 전용: smilref/id 부여하지 않음)
                level1 = etree.SubElement(dtbook_bodymatter, "level1")
                h1 = etree.SubElement(level1, "h1")
                h1.text = "제목 없음"
                level_elements[1] = level1
                current_level = 1
                parent_elem = level1

            # 이미지 그룹 생성
            imggroup = etree.SubElement(parent_elem, "imggroup",
                                      id=item["id"], class_="figure")

            # 이미지 요소 생성
            img = etree.SubElement(imggroup, "img",
                                 id=f"{item['id']}_img",
                                 src=item["src"],
                                 alt=item["alt_text"])
            
            # 이미지 크기를 적절히 설정
            img.set("width", "60%")
            img.set("height", "auto")
            
            # 이미지 캡션 추가
            caption = etree.SubElement(imggroup, "caption",
                                     id=f"{item['id']}_caption")
            sent = etree.SubElement(caption, "sent",
                                  id=item["sent_id"],
                                  smilref=f"dtbook.smil#smil_par_{item['sent_id']}")
            
            continue
        elif item["type"] == "table":
            # 표는 현재 활성 레벨 요소에 추가
            parent_elem = level_elements.get(current_level, dtbook_bodymatter)
            if parent_elem is dtbook_bodymatter:
                # level1이 없는 경우 생성
                level1 = etree.SubElement(dtbook_bodymatter, "level1",
                                        id=item["id"],
                                        smilref=f"dtbook.smil#smil_par_{item['id']}")
                h1 = etree.SubElement(level1, "h1")
                h1.text = "제목 없음"
                level_elements[1] = level1
                current_level = 1
                parent_elem = level1

            # 표 요소 생성 (스타일 속성 추가)
            table_elem = etree.SubElement(parent_elem, "table",
                                   class_="data-table",
                                   border="1",
                                   style="width: 100%; border-collapse: collapse; border: 3px double #000;")

            table_data = item["table_data"]

            # tbody 요소 생성
            tbody = etree.SubElement(table_elem, "tbody")

            # 재귀 렌더링 함수 (문단/표 순서 보존)
            def render_table_to_dtbook(parent_tbody, table_data_obj, base_id):
                for row_idx, row_data in enumerate(table_data_obj["rows"]):
                    tr = etree.SubElement(parent_tbody, "tr",
                                          style="border: 3px double #000;")
                    for col_idx, _ in enumerate(row_data):
                        cell_info = next((cell for cell in table_data_obj["cells"]
                                          if cell["row"] == row_idx and cell["col"] == col_idx), None)
                        if cell_info and cell_info.get("is_merged_area", False):
                            continue

                        cell_elem = etree.SubElement(tr, "td",
                                                    style="text-align: left; vertical-align: middle; font-weight: normal; border: 3px double #000;")

                        if cell_info and cell_info["is_merged"]:
                            if cell_info["rowspan"] > 1:
                                cell_elem.set("rowspan", str(cell_info["rowspan"]))
                            if cell_info["colspan"] > 1:
                                cell_elem.set("colspan", str(cell_info["colspan"]))

                        seq = cell_info.get('content_sequence', []) if cell_info else []
                        seq_para_counter = 0
                        for s_idx, s in enumerate(seq):
                            if s.get('type') == 'p':
                                para_text = s.get('text', '')
                                p = etree.SubElement(
                                    cell_elem,
                                    "p",
                                    id=f"table_{base_id}_cell_{row_idx}_{col_idx}_p_{seq_para_counter}",
                                    smilref=f"dtbook.smil#smil_par_{base_id}_cell_{row_idx}_{col_idx}_p_{seq_para_counter}",
                                    style="margin: 0; padding: 8px; text-align: left; vertical-align: middle; font-weight: normal;"
                                )
                                if para_text.strip() == "<br/>":
                                    etree.SubElement(p, "br")
                                elif para_text.strip():
                                    p.text = para_text.strip()
                                seq_para_counter += 1
                            elif s.get('type') == 'table':
                                nested_table_elem = etree.SubElement(cell_elem, "table",
                                                                     border="1",
                                                                     style="width: 100%; border-collapse: collapse; border: 3px double #000;")
                                nested_tbody = etree.SubElement(nested_table_elem, "tbody")
                                nested_base_id = f"{base_id}_cell_{row_idx}_{col_idx}_nested_{s_idx}"
                                render_table_to_dtbook(nested_tbody, s.get('table_data', {}), nested_base_id)

                        # 셀 이미지 출력 (문단/표 외부에 인라인 이미지가 잡힌 경우)
                        if cell_info and cell_info.get("images"):
                            for img_idx, img in enumerate(cell_info["images"]):
                                nonlocal image_counter
                                image_counter += 1
                                image_ext = ".jpeg"
                                try:
                                    if 'image_rid' in img:
                                        rel = image_relations.get(img['image_rid'])
                                        if rel and hasattr(rel, 'target_ref'):
                                            ext = os.path.splitext(rel.target_ref)[1]
                                            if ext:
                                                image_ext = ext
                                except Exception:
                                    pass

                                image_filename = f"table_{base_id}_cell_{row_idx}_{col_idx}_img_{img_idx}{image_ext}"
                                image_path = os.path.join(output_dir, image_filename)
                                with open(image_path, "wb") as img_file:
                                    img_file.write(img['image_data'])

                                img_elem = etree.SubElement(cell_elem, "img",
                                                            src=image_filename,
                                                            alt=f"표 셀 이미지 {img_idx+1}")
                                img_elem.set("width", "100%")
                                img_elem.set("height", "auto")
                                etree.SubElement(cell_elem, "br")

                        # 중첩 표 재귀 렌더링
                        nested_tables = cell_info.get('nested_tables', []) if cell_info else []
                        for n_idx, nested_table_data in enumerate(nested_tables):
                            nested_table_elem = etree.SubElement(cell_elem, "table",
                                                                 border="1",
                                                                 style="width: 100%; border-collapse: collapse; border: 3px double #000;")
                            nested_tbody = etree.SubElement(nested_table_elem, "tbody")
                            nested_base_id = f"{base_id}_cell_{row_idx}_{col_idx}_nested_{n_idx}"
                            render_table_to_dtbook(nested_tbody, nested_table_data, nested_base_id)

            render_table_to_dtbook(tbody, table_data, item['id'])
        elif item["type"].startswith("h"):
            level = int(item["type"][1])  # h1 -> 1, h2 -> 2, h3 -> 3, h4 -> 4, h5 -> 5, h6 -> 6

            if level == 1:
                # 새로운 level1 시작
                level1 = etree.SubElement(dtbook_bodymatter, "level1",
                                        id=item["id"],
                                        smilref=f"dtbook.smil#smil_par_{item['id']}")
                h1 = etree.SubElement(level1, "h1")
                h1.text = " ".join(item["words"])
                
                # 레벨 요소 업데이트
                level_elements[1] = level1
                current_level = 1
                
                # 더 낮은 레벨 요소들 제거 (새로운 level1이 시작되므로)
                for l in range(2, 7):
                    if l in level_elements:
                        del level_elements[l]
                        
            else:
                # level2~6 처리
                if 1 not in level_elements:
                    # level1이 없는 경우 생성
                    level1 = etree.SubElement(dtbook_bodymatter, "level1",
                                            id=item["id"],
                                            smilref=f"dtbook.smil#smil_par_{item['id']}")
                    h1 = etree.SubElement(level1, "h1")
                    h1.text = "제목 없음"
                    level_elements[1] = level1
                    current_level = 1

                # 부모 레벨 찾기 (현재 레벨보다 1 작은 레벨)
                parent_level = level - 1
                parent_elem = level_elements.get(parent_level)
                
                if parent_elem is None:
                    # 부모 레벨이 없으면 level1을 부모로 사용
                    parent_elem = level_elements[1]
                
                # 새로운 level 요소 생성
                new_level = etree.SubElement(parent_elem, f"level{level}",
                                           id=item["id"],
                                           smilref=f"dtbook.smil#smil_par_{item['id']}")
                
                # 제목 요소 생성
                heading = etree.SubElement(new_level, f"h{level}")
                heading.text = " ".join(item["words"])
                
                # 레벨 요소 업데이트
                level_elements[level] = new_level
                current_level = level
                
                # 더 높은 레벨 요소들 제거 (새로운 레벨이 시작되므로)
                for l in range(level + 1, 7):
                    if l in level_elements:
                        del level_elements[l]

            # 기타 마커 처리
            for marker in item.get("markers", []):
                if marker.type != "page":  # 페이지 마커는 이미 처리됨
                    elem_info = MarkerProcessor.create_dtbook_element(marker)
                    if elem_info:
                        # 현재 레벨 요소에 마커 추가
                        current_elem = level_elements.get(current_level, level_elements.get(1))
                        marker_elem = etree.SubElement(current_elem,
                                                     elem_info["tag"],
                                                     attrib=elem_info["attrs"])
                        marker_elem.text = elem_info["text"]
        else:
            # 일반 단락은 현재 활성 레벨 요소에 추가
            parent_elem = level_elements.get(current_level, dtbook_bodymatter)
            if parent_elem is dtbook_bodymatter:
                # level1이 없는 경우 생성
                level1 = etree.SubElement(dtbook_bodymatter, "level1",
                                        id=item["id"],
                                        smilref=f"dtbook.smil#smil_par_{item['id']}")
                h1 = etree.SubElement(level1, "h1")
                h1.text = "제목 없음"
                level_elements[1] = level1
                current_level = 1
                parent_elem = level1

            p = etree.SubElement(parent_elem, "p",
                               id=item["id"],
                               smilref=f"dtbook.smil#smil_par_{item['id']}")
            
            # <br/> 태그가 포함된 경우 실제 br 요소로 생성
            if item.get("text", "") == "<br/>":
                br_elem = etree.SubElement(p, "br")
            else:
                p.text = " ".join(item["words"])

            # 기타 마커 처리
            for marker in item.get("markers", []):
                if marker.type != "page":  # 페이지 마커는 이미 처리됨
                    elem_info = MarkerProcessor.create_dtbook_element(marker)
                    if elem_info:
                        marker_elem = etree.SubElement(parent_elem,
                                                     elem_info["tag"],
                                                     attrib=elem_info["attrs"])
                        marker_elem.text = elem_info["text"]

    # XML 파일 저장
    dtbook_filepath = os.path.join(output_dir, "dtbook.xml")
    tree = etree.ElementTree(dtbook_root)

    # XML 선언에 인코딩 명시적 지정
    with open(dtbook_filepath, 'wb') as f:
        # XML 선언
        f.write('<?xml version="1.0" encoding="utf-8" standalone="no"?>\n'.encode('utf-8'))
        # DTD 선언
        f.write('<!DOCTYPE dtbook PUBLIC "-//NISO//DTD dtbook 2005-3//EN" "http://www.daisy.org/z3986/2005/dtbook-2005-3.dtd">\n'.encode('utf-8'))
        # XML 트리 저장 - 명시적으로 인코딩 지정
        tree.write(f,
                  encoding='utf-8',
                  pretty_print=True,
                  xml_declaration=False,
                  method='xml')

    print(f"DTBook 생성 완료: {dtbook_filepath}")
    timings["generate_dtbook"] = time.time() - t_dtbook

    # --- 2. OPF 파일 생성 (dtbook.opf) ---
    t_opf = time.time()
    print("OPF 생성 중...")
    opf_ns = "http://openebook.org/namespaces/oeb-package/1.0/"
    dc_ns = "http://purl.org/dc/elements/1.1/"

    opf_root = etree.Element(
        "package",
        attrib={
            "unique-identifier": "uid"
        },
        nsmap={
            None: opf_ns
        }
    )
 
    # 메타데이터
    metadata = etree.SubElement(opf_root, "metadata")

    # DC 메타데이터
    dc_metadata = etree.SubElement(metadata, "dc-metadata")

    format_elem = etree.SubElement(
        dc_metadata, "{%s}Format" % dc_ns, nsmap={'dc': dc_ns})
    format_elem.text = "ANSI/NISO Z39.86-2005"

    lang_elem = etree.SubElement(
        dc_metadata, "{%s}Language" % dc_ns, nsmap={'dc': dc_ns})
    lang_elem.text = book_language

    date_elem = etree.SubElement(
        dc_metadata, "{%s}Date" % dc_ns, nsmap={'dc': dc_ns})
    date_elem.text = datetime.now().strftime("%Y-%m-%d")

    publisher_elem = etree.SubElement(
        dc_metadata, "{%s}Publisher" % dc_ns, nsmap={'dc': dc_ns})
    publisher_elem.text = book_publisher

    title_elem = etree.SubElement(
        dc_metadata, "{%s}Title" % dc_ns, nsmap={'dc': dc_ns})
    title_elem.text = book_title

    identifier_elem = etree.SubElement(
        dc_metadata, "{%s}Identifier" % dc_ns, nsmap={'dc': dc_ns}, id="uid")
    identifier_elem.text = book_uid

    creator_elem = etree.SubElement(
        dc_metadata, "{%s}Creator" % dc_ns, nsmap={'dc': dc_ns})
    creator_elem.text = book_author

    # X-Metadata
    x_metadata = etree.SubElement(metadata, "x-metadata")
    etree.SubElement(x_metadata, "meta",
                     name="dtb:multimediaType",
                     content="textNCX")
    etree.SubElement(x_metadata, "meta",
                     name="dtb:totalTime",
                     content="0:00:00")
    etree.SubElement(x_metadata, "meta",
                     name="dtb:multimediaContent",
                     content="text")

    # Manifest
    manifest = etree.SubElement(opf_root, "manifest")

    # OPF
    etree.SubElement(manifest, "item",
                     href="dtbook.opf",
                     id="opf",
                     **{"media-type": "text/xml"})

    # DTBook
    etree.SubElement(manifest, "item",
                     href="dtbook.xml",
                     id="opf-1",
                     **{"media-type": "application/x-dtbook+xml"})

    # SMIL 파일들
    etree.SubElement(manifest, "item",
                     href="dtbook.smil",
                     id="mo",
                     **{"media-type": "application/smil"})

    # NCX
    etree.SubElement(manifest, "item",
                     href="dtbook.ncx",
                     id="ncx",
                     **{"media-type": "application/x-dtbncx+xml"})

    # Resources
    etree.SubElement(manifest, "item",
                     href="dtbook.res",
                     id="resource",
                     **{"media-type": "application/x-dtbresource+xml"})

    # Spine
    spine = etree.SubElement(opf_root, "spine")
    etree.SubElement(spine, "itemref",
                     idref="mo")

    # OPF Manifest에 이미지 파일 추가
    for item in content_structure:
        if item["type"] == "image":
            image_filename = os.path.basename(item["src"])
            image_id = f"img_{item['id']}"
            extension = os.path.splitext(image_filename)[1][1:].lower()

            # 이미지 확장자에 따른 MIME 타입 설정
            mime_type = {
                'jpg': 'image/jpeg',
                'jpeg': 'image/jpeg',
                'png': 'image/png',
                'gif': 'image/gif',
                'bmp': 'image/bmp',
                'tiff': 'image/tiff',
                'tif': 'image/tiff'
            }.get(extension, f'image/{extension}')

            print(f"이미지 매니페스트 추가: {image_filename} (MIME: {mime_type})")
            etree.SubElement(manifest, "item",
                             href=item["src"],
                             id=image_id,
                             **{"media-type": mime_type})
    
    # 표 안의 이미지 파일들도 매니페스트에 추가 (중첩 표 재귀 지원)
    def add_table_images_to_manifest(table_data_obj, base_id):
        for cell in table_data_obj["cells"]:
            if cell.get("images"):
                for img_idx, img in enumerate(cell["images"]):
                    image_ext = ".jpeg"
                    try:
                        if 'image_rid' in img:
                            rel = image_relations.get(img['image_rid'])
                            if rel and hasattr(rel, 'target_ref'):
                                ext = os.path.splitext(rel.target_ref)[1]
                                if ext:
                                    image_ext = ext
                    except Exception:
                        pass
                    image_filename = f"table_{base_id}_cell_{cell['row']}_{cell['col']}_img_{img_idx}{image_ext}"
                    image_id = f"img_table_{base_id}_cell_{cell['row']}_{cell['col']}_{img_idx}"
                    extension = os.path.splitext(image_filename)[1][1:].lower()
                    mime_type = {
                        'jpg': 'image/jpeg',
                        'jpeg': 'image/jpeg',
                        'png': 'image/png',
                        'gif': 'image/gif',
                        'bmp': 'image/bmp',
                        'tiff': 'image/tiff',
                        'tif': 'image/tiff'
                    }.get(extension, f'image/{extension}')
                    print(f"표 안 이미지 매니페스트 추가: {image_filename} (MIME: {mime_type})")
                    etree.SubElement(manifest, "item",
                                     href=image_filename,
                                     id=image_id,
                                     **{"media-type": mime_type})

            # 중첩 표 재귀 처리: content_sequence 기반
            for s_idx, s in enumerate(cell.get('content_sequence', []) or []):
                if s.get('type') == 'table':
                    nested_base_id = f"{base_id}_cell_{cell['row']}_{cell['col']}_nested_{s_idx}"
                    add_table_images_to_manifest(s.get('table_data', {}), nested_base_id)

    for item in content_structure:
        if item["type"] == "table":
            add_table_images_to_manifest(item["table_data"], item['id'])

    # OPF 파일 저장
    opf_filepath = os.path.join(output_dir, "dtbook.opf")
    tree = etree.ElementTree(opf_root)

    with open(opf_filepath, 'wb') as f:
        f.write('<?xml version="1.0" encoding="utf-8" standalone="no"?>\n'.encode('utf-8'))
        f.write('<!DOCTYPE package PUBLIC "+//ISBN 0-9673008-1-9//DTD OEB 1.2 Package//EN" "http://openebook.org/dtds/oeb-1.2/oebpkg12.dtd">\n'.encode('utf-8'))
        tree.write(f,
                  encoding='utf-8',
                  pretty_print=True,
                  xml_declaration=False,
                  method='xml')

    print(f"OPF 생성 완료: {opf_filepath}")
    timings["generate_opf"] = time.time() - t_opf

    # --- 3. SMIL 파일 생성 (dtbook.smil) ---
    t_smil = time.time()
    print("SMIL 파일 생성 중...")

    smil_ns = "http://www.w3.org/2001/SMIL20/"

    smil_root = etree.Element(
        "smil",
        nsmap={
            None: smil_ns
        }
    )

    # head
    head = etree.SubElement(smil_root, "head")
    etree.SubElement(head, "meta",
                     name="dtb:uid",
                     content=book_uid)
    etree.SubElement(head, "meta",
                     name="dtb:totalElapsedTime",
                     content="0:00:00")
    etree.SubElement(head, "meta",
                     name="dtb:generator",
                     content="DAISY Pipeline 2")

    # 페이지 관련 메타데이터 추가
    etree.SubElement(head, "meta",
                     name="dtb:totalPageCount",
                     content=str(max_page_number))
    etree.SubElement(head, "meta",
                     name="dtb:maxPageNumber",
                     content=str(max_page_number))

    # body
    body = etree.SubElement(smil_root, "body")
    root_seq = etree.SubElement(body, "seq", id="root-seq")

    # doctitle과 docauthor 추가
    doctitle_par = etree.SubElement(root_seq, "par",
                                   id="sforsmil-1",
                                   **{"class": "doctitle"})
    etree.SubElement(doctitle_par, "text",
                     src="dtbook.xml#forsmil-1")

    docauthor_par = etree.SubElement(root_seq, "par",
                                    id="sforsmil-2",
                                    **{"class": "docauthor"})
    etree.SubElement(docauthor_par, "text",
                     src="dtbook.xml#forsmil-2")

    # 나머지 콘텐츠 추가
    for item in content_structure:
        # pagenum 타입은 이미 DTBook에서 처리되었으므로 SMIL에서만 처리
        if item["type"] == "pagenum":
            page_par = etree.SubElement(root_seq, "par",
                                      id=f"smil_par_page_{item['text']}_{item['text']}",
                                      **{"class": "pagenum"},
                                      customTest="pagenum")
            etree.SubElement(page_par, "text",
                           src=f"dtbook.xml#page_{item['text']}_{item['text']}")
            continue

        # 기본 콘텐츠
        if item["type"].startswith("h"):
            # 제목 요소일 경우 level로 처리
            level = int(item["type"][1])  # h1 -> 1, h2 -> 2, h3 -> 3, h4 -> 4, h5 -> 5, h6 -> 6
            par = etree.SubElement(root_seq, "par",
                                 id=f"smil_par_{item['id']}",
                                 **{"class": f"level{level}"})
            etree.SubElement(par, "text",
                           src=f"dtbook.xml#{item['id']}")
        else:
            # 일반 콘텐츠 처리 (table은 별도 처리하므로 건너뜀)
            if item["type"] != "table":
                par = etree.SubElement(root_seq, "par",
                                     id=f"smil_par_{item['id']}",
                                     **{"class": item["type"]})
                etree.SubElement(par, "text",
                               src=f"dtbook.xml#{item['id']}")

        # 표 처리
        if item["type"] == "table":
            table_data = item["table_data"]

            def render_table_to_smil(parent_seq, table_data_obj, base_id):
                for row_idx, row_data in enumerate(table_data_obj["rows"]):
                    for col_idx, _ in enumerate(row_data):
                        cell_info = next((cell for cell in table_data_obj["cells"]
                                          if cell["row"] == row_idx and cell["col"] == col_idx), None)
                        if cell_info and cell_info.get("is_merged_area", False):
                            continue

                        seq = cell_info.get('content_sequence', []) if cell_info else []
                        para_counter = 0
                        for s_idx, s in enumerate(seq):
                            if s.get('type') == 'p':
                                # 빈 문단은 SMIL 생성에서 제외
                                if (s.get('text') or '').strip():
                                    cell_par = etree.SubElement(parent_seq, "par",
                                                                id=f"smil_par_{base_id}_cell_{row_idx}_{col_idx}_p_{para_counter}",
                                                                **{"class": "p"})
                                    etree.SubElement(cell_par, "text",
                                                     src=f"dtbook.xml#table_{base_id}_cell_{row_idx}_{col_idx}_p_{para_counter}")
                                    para_counter += 1
                            elif s.get('type') == 'table':
                                nested_base_id = f"{base_id}_cell_{row_idx}_{col_idx}_nested_{s_idx}"
                                render_table_to_smil(parent_seq, s.get('table_data', {}), nested_base_id)

            render_table_to_smil(root_seq, table_data, item['id'])

        # 마커 처리 (페이지 마커 제외)
        for marker in item.get("markers", []):
            if marker.type != "page":  # 페이지 마커는 이미 처리됨
                elem_info = MarkerProcessor.create_smil_element(marker, item["id"])
                if elem_info:
                    marker_par = etree.SubElement(root_seq, "par",
                                                id=f"smil_par_{item['id']}_{marker.type}",
                                                **{"class": elem_info["par_class"]})
                    etree.SubElement(marker_par, "text",
                                   src=elem_info["text_src"])

    # SMIL 파일 저장
    smil_filepath = os.path.join(output_dir, "dtbook.smil")
    tree = etree.ElementTree(smil_root)

    with open(smil_filepath, 'wb') as f:
        f.write('<?xml version="1.0" encoding="utf-8" standalone="no"?>\n'.encode('utf-8'))
        f.write('<!DOCTYPE smil PUBLIC "-//NISO//DTD dtbsmil 2005-2//EN" "http://www.daisy.org/z3986/2005/dtbsmil-2005-2.dtd">\n'.encode('utf-8'))
        tree.write(f,
                  encoding='utf-8',
                  pretty_print=True,
                  xml_declaration=False,
                  method='xml')

    print(f"SMIL 파일 생성 완료: {smil_filepath}")
    timings["generate_smil"] = time.time() - t_smil

    # --- 4. NCX 파일 생성 (dtbook.ncx) ---
    t_ncx = time.time()
    print("NCX 생성 중...")
    ncx_ns = "http://www.daisy.org/z3986/2005/ncx/"

    ncx_root = etree.Element(
        "{%s}ncx" % ncx_ns,
        attrib={
            "version": "2005-1"
        },
        nsmap={
            None: ncx_ns
        }
    )

    # head
    head = etree.SubElement(ncx_root, "head")
    # 최대 제목 레벨 계산 (dtb:depth 반영)
    max_heading_level = 0
    for _item in content_structure:
        if isinstance(_item.get("type"), str) and _item["type"].startswith("h"):
            try:
                max_heading_level = max(max_heading_level, int(_item["type"][1]))
            except Exception:
                pass
    if max_heading_level <= 0:
        max_heading_level = 1

    etree.SubElement(head, "meta",
                     name="dc:Identifier",
                     content=book_uid)
    etree.SubElement(head, "meta",
                     name="dc:Title",
                     content=book_title)
    etree.SubElement(head, "meta",
                     name="dc:Date",
                     content=datetime.now().strftime("%Y-%m-%d"))
    etree.SubElement(head, "meta",
                     name="dc:Format",
                     content="ANSI/NISO Z39.86-2005")
    etree.SubElement(head, "meta",
                     name="dc:Language",
                     content=book_language)
    etree.SubElement(head, "meta",
                     name="dtb:depth",
                     content=str(min(6, max_heading_level)))
    etree.SubElement(head, "meta",
                     name="dtb:totalPageCount",
                     content=str(max_page_number))
    etree.SubElement(head, "meta",
                     name="dtb:maxPageNumber",
                     content=str(max_page_number))
    etree.SubElement(head, "meta",
                     name="dtb:uid",
                     content=book_uid)
    etree.SubElement(head, "meta",
                     name="dtb:generator",
                     content="docx_to_daisy")

    # smilCustomTest 추가
    etree.SubElement(head, "smilCustomTest",
                    id="pagenum",
                    defaultState="false",
                    override="visible",
                    bookStruct="PAGE_NUMBER")
    etree.SubElement(head, "smilCustomTest",
                    id="note",
                    defaultState="true",
                    override="visible",
                    bookStruct="NOTE")
    etree.SubElement(head, "smilCustomTest",
                    id="noteref",
                    defaultState="true",
                    override="visible",
                    bookStruct="NOTE_REFERENCE")
    etree.SubElement(head, "smilCustomTest",
                    id="table",
                    defaultState="true",
                    override="visible")

    # docTitle
    doc_title = etree.SubElement(ncx_root, "docTitle")
    text = etree.SubElement(doc_title, "text")
    text.text = book_title

    # docAuthor
    doc_author = etree.SubElement(ncx_root, "docAuthor")
    text = etree.SubElement(doc_author, "text")
    text.text = book_author

    # navMap
    nav_map = etree.SubElement(ncx_root, "navMap")

    # 목차 항목 생성 (제목만 포함, 표 제외)
    play_order = 1
    level_stack = [None] * 7  # 1~6 사용

    for item in content_structure:
        if isinstance(item.get("type"), str) and item["type"].startswith("h"):
            try:
                level = int(item["type"][1])  # h1 -> 1, ..., h6 -> 6
            except Exception:
                continue

            nav_point = etree.Element(
                "navPoint",
                id=f"ncx_{item['id']}",
                **{"class": f"level{level}"},
                playOrder=str(play_order)
            )
            nav_label = etree.SubElement(nav_point, "navLabel")
            text = etree.SubElement(nav_label, "text")
            label_text = item.get("text") or " ".join(item.get("words", [])) or "제목 없음"
            text.text = label_text
            etree.SubElement(nav_point, "content",
                             src=f"dtbook.smil#smil_par_{item['id']}")

            # 부모 찾기: 현재 레벨보다 작은 가장 가까운 상위 레벨
            parent = None
            for pl in range(level - 1, 0, -1):
                if level_stack[pl] is not None:
                    parent = level_stack[pl]
                    break

            if parent is None:
                nav_map.append(nav_point)
            else:
                parent.append(nav_point)

            # 스택 갱신
            level_stack[level] = nav_point
            for clr in range(level + 1, 7):
                level_stack[clr] = None

            play_order += 1

    # pageList (문서 내 '$#' 페이지 마커만을 순서대로 추가 - 표지 제외)
    page_targets = []

    for item in content_structure:
        if item.get("type") == "pagenum":
            page_value = str(item.get("text", "")).strip()
            if not page_value:
                continue
            page_targets.append({
                "id": f"p{page_value}",
                "value": page_value,
                "type": "normal",
                "play_order": play_order
            })
            play_order += 1

    if page_targets:
        page_list = etree.SubElement(ncx_root, "pageList", id="pages")
        nav_label = etree.SubElement(page_list, "navLabel")
        text = etree.SubElement(nav_label, "text")
        text.text = "Page numbers list"

        for page in page_targets:
            nav_point = etree.SubElement(page_list, "pageTarget",
                                        id=page["id"],
                                        **{"class": "pagenum"},
                                        type=page["type"],
                                        value=page["value"],
                                        playOrder=str(page["play_order"]))
            nav_label = etree.SubElement(nav_point, "navLabel")
            text = etree.SubElement(nav_label, "text")
            text.text = page["value"]
            content = etree.SubElement(nav_point, "content",
                                      src=f"dtbook.smil#smil_par_page_{page['value']}_{page['value']}")

    # navList (각주, 미주 등이 있는 경우 추가)
    note_targets = []
    for item in content_structure:
        for marker in item.get("markers", []):
            if marker.type in ["note", "annotation"]:
                note_targets.append({
                    "id": f"note_{marker.value}",
                    "text": marker.text,
                    "smil_file": item["smil_file"],
                    "item_id": item["id"],
                    "play_order": play_order
                })
                play_order += 1

    if note_targets:
        nav_list = etree.SubElement(ncx_root, "navList")
        nav_label = etree.SubElement(nav_list, "navLabel")
        text = etree.SubElement(nav_label, "text")
        text.text = "각주"

        for note in note_targets:
            nav_target = etree.SubElement(nav_list, "navTarget",
                                         id=note["id"],
                                         playOrder=str(note["play_order"]))
            nav_label = etree.SubElement(nav_target, "navLabel")
            text = etree.SubElement(nav_label, "text")
            text.text = note["text"]
            content = etree.SubElement(nav_target, "content",
                                      src=f"{note['smil_file']}#s{note['item_id']}")

    # NCX 파일 저장
    ncx_filepath = os.path.join(output_dir, "dtbook.ncx")
    tree = etree.ElementTree(ncx_root)

    with open(ncx_filepath, 'wb') as f:
        f.write(b'<?xml version="1.0" encoding="UTF-8"?>\n')
        f.write(b'<!DOCTYPE ncx PUBLIC "-//NISO//DTD ncx 2005-1//EN" "http://www.daisy.org/z3986/2005/ncx-2005-1.dtd" >\n')
        tree.write(f,
                   pretty_print=True,
                   encoding='utf-8',
                   xml_declaration=False)

    print(f"NCX 생성 완료: {ncx_filepath}")
    timings["generate_ncx"] = time.time() - t_ncx

    # --- 5. Resources 파일 생성 (dtbook.res) ---
    t_res = time.time()
    print("Resources 생성 중...")
    res_ns = "http://www.daisy.org/z3986/2005/resource/"

    res_root = etree.Element(
        "{%s}resources" % res_ns,
        attrib={
            "version": "2005-1"
        },
        nsmap={
            None: res_ns
        }
    )

    # NCX scope
    ncx_scope = etree.SubElement(res_root, "scope",
                                 nsuri="http://www.daisy.org/z3986/2005/ncx/")

    # Custom tests
    custom_tests = [
        ("page-set", "PAGE_NUMBER", "page"),
        ("note-set", "NOTE", "note"),
        ("notref-set", "NOTE_REFERENCE", "note"),
        ("annot-set", "ANNOTATION", "annotation"),
        ("line-set", "LINE_NUMBER", "line"),
        ("sidebar-set", "OPTIONAL_SIDEBAR", "sidebar"),
        ("prodnote-set", "OPTIONAL_PRODUCER_NOTE", "note")
    ]

    for test_id, book_struct, text in custom_tests:
        node_set = etree.SubElement(ncx_scope, "nodeSet",
                                    id=test_id,
                                    select=f"//smilCustomTest[@bookStruct='{book_struct}']")
        resource = etree.SubElement(node_set, "resource",
                                    **{"{http://www.w3.org/XML/1998/namespace}lang": "en"})
        text_elem = etree.SubElement(resource, "text")
        text_elem.text = text

    # SMIL scope
    smil_scope = etree.SubElement(res_root, "scope",
                                  nsuri="http://www.w3.org/2001/SMIL20/")

    # Math sets
    math_sets = [
        ("math-seq-set", "seq", "mathematical formula"),
        ("math-par-set", "par", "mathematical formula")
    ]

    for set_id, elem_type, text in math_sets:
        node_set = etree.SubElement(smil_scope, "nodeSet",
                                    id=set_id,
                                    select=f"//{elem_type}[@class='math']")
        resource = etree.SubElement(node_set, "resource",
                                    **{"{http://www.w3.org/XML/1998/namespace}lang": "en"})
        text_elem = etree.SubElement(resource, "text")
        text_elem.text = text

    # Resources 파일 저장
    res_filepath = os.path.join(output_dir, "dtbook.res")
    tree = etree.ElementTree(res_root)

    with open(res_filepath, 'wb') as f:
        f.write(b'<?xml version="1.0" encoding="utf-8"?>\n')
        f.write(b'<!DOCTYPE resources\n  PUBLIC "-//NISO//DTD resource 2005-1//EN" "http://www.daisy.org/z3986/2005/resource-2005-1.dtd">\n')
        tree.write(f,
                   pretty_print=True,
                   encoding='utf-8',
                   xml_declaration=False)

    print(f"Resources 생성 완료: {res_filepath}")
    timings["generate_resources"] = time.time() - t_res

    print("\n--- DAISY 기본 파일 생성 완료 ---")
    print(f"생성된 파일은 '{output_dir}' 폴더에 있습니다.")
    print("주의: 이 코드는 DOCX의 기본적인 제목/문단 구조만 변환하며,")
    print("      오디오, SMIL 동기화, 목록, 표, 이미지, 페이지 번호 등은 포함하지 않습니다.")

    # 최종 메모리 정리 (모든 DAISY 파일 생성 완료 후)
    if 'image_relations' in locals():
        del image_relations
    gc.collect()

    # 타이밍 반환 (상위 호출자 기록용)
    return timings

    
def create_daisy_book_with_validation(docx_file_path, output_dir, book_title=None, book_author=None, book_publisher=None, book_language="ko", progress_callback=None):
    """DOCX 파일을 DAISY 형식으로 변환하고 검증을 수행합니다.

    Args:
        docx_file_path (str): 변환할 DOCX 파일 경로
        output_dir (str): 생성된 DAISY 파일 저장 폴더
        book_title (str, optional): 책 제목. 기본값은 None (DOCX 파일명 사용)
        book_author (str, optional): 저자. 기본값은 None
        book_publisher (str, optional): 출판사. 기본값은 None
        book_language (str, optional): 언어 코드 (ISO 639-1). 기본값은 "ko"
        progress_callback (callable, optional): 진행 상황을 보고하는 콜백 함수
    """
    try:
        # DAISY 파일 생성
        create_daisy_book(docx_file_path, output_dir, book_title, book_author, book_publisher, book_language)
        
        # 검증 단계 시작
        if progress_callback:
            progress_callback(95, "DAISY 파일 검증 중...")
        
        print("DAISY 파일 검증 시작...")
        
        # DAISY 검증 수행
        validator = DaisyValidator(output_dir)
        validation_result = validator.validate_all()
        
        # 검증 결과 처리
        if not validation_result.is_valid:
            error_messages = [f"{error.category}: {error.message}" for error in validation_result.errors]
            error_summary = "; ".join(error_messages[:3])  # 처음 3개 오류만 표시
            if len(error_messages) > 3:
                error_summary += f" 외 {len(error_messages) - 3}개 오류"
            
            raise ValueError(f"DAISY 파일 검증 실패: {error_summary}")
        
        # 경고가 있는 경우 로그로 출력
        if validation_result.warnings:
            warning_count = len(validation_result.warnings)
            print(f"DAISY 파일 검증 완료: {warning_count}개의 경고가 있습니다.")
            for warning in validation_result.warnings[:5]:  # 처음 5개 경고만 출력
                print(f"  경고: {warning.category} - {warning.message}")
            if len(validation_result.warnings) > 5:
                print(f"  ... 외 {len(validation_result.warnings) - 5}개 경고")
        else:
            print("DAISY 파일 검증 완료: 모든 검증을 통과했습니다.")
        
        # 검증 완료 메시지
        if progress_callback:
            progress_callback(100, "변환 및 검증이 완료되었습니다.")
        
        return validation_result
        
    except Exception as e:
        if progress_callback:
            progress_callback(100, f"변환 실패: {str(e)}")
        raise e


def zip_daisy_output(source_dir, output_zip_filename):

    """
    지정된 폴더의 내용을 ZIP 파일로 압축합니다.

    Args:
        source_dir (str): 압축할 DAISY 파일들이 있는 폴더 경로.
        output_zip_filename (str): 생성될 ZIP 파일의 이름 (경로 포함 가능).
    """
    if not os.path.isdir(source_dir):
        print(f"오류: 소스 디렉토리를 찾을 수 없습니다 - {source_dir}")
        return

    try:
        print(f"'{source_dir}' 폴더를 '{output_zip_filename}' 파일로 압축 중...")
        # ZIP 파일 쓰기 모드로 열기 (압축 사용)
        with zipfile.ZipFile(output_zip_filename, 'w', zipfile.ZIP_DEFLATED) as zipf:
            # source_dir 내부의 모든 파일과 폴더를 순회
            for root, dirs, files in os.walk(source_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    # ZIP 파일 내부에 저장될 상대 경로 계산
                    # (source_dir 자체를 포함하지 않도록 함)
                    archive_name = os.path.relpath(file_path, source_dir)
                    print(f"  추가 중: {archive_name}")
                    zipf.write(file_path, arcname=archive_name)
        print(f"ZIP 파일 생성 완료: {output_zip_filename}")
    except Exception as e:
        print(f"ZIP 파일 생성 중 오류 발생: {e}")

