import zipfile
import os
import uuid
import argparse
import re
import logging
import html
from docx import Document  # python-docx 라이브러리
from lxml import etree  # lxml 라이브러리
from datetime import datetime
from ..markers import MarkerProcessor  # 마커 처리기 임포트
import gc

# 로깅 설정
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# 컴파일된 정규식 패턴들 (성능 최적화)
BR_PATTERN = re.compile(r'<br\s*/?>', flags=re.IGNORECASE)
HTML_TAG_PATTERN = re.compile(r'<[^>]*>')
PUNCTUATION_PATTERN = re.compile(r'[.。,，!！?？:：;；]')
BRACKET_PATTERN = re.compile(r'\[(.*?)\]')
TABLE_TITLE_PATTERN = re.compile(r'\[?표\s*\d+\.?\d*\]?', re.IGNORECASE)


def html_escape(text):
    """HTML 특수 문자를 이스케이프하고 HTML 태그를 제거하는 함수
    
    Args:
        text (str): 이스케이프할 텍스트
        
    Returns:
        str: 이스케이프된 텍스트 (HTML 태그 제거됨)
    """
    if not isinstance(text, str):
        return str(text)

    # HTML 태그 제거
    cleaned_text = HTML_TAG_PATTERN.sub('', text)

    # HTML 특수 문자 이스케이프
    escaped = html.escape(cleaned_text, quote=True)

    # 추가적인 이스케이프 처리
    escaped = escaped.replace('&amp;', '&amp;')  # 이미 이스케이프된 &는 유지
    escaped = escaped.replace('&lt;', '&lt;')    # 이미 이스케이프된 <는 유지
    escaped = escaped.replace('&gt;', '&gt;')    # 이미 이스케이프된 >는 유지
    escaped = escaped.replace('&quot;', '&quot;')  # 이미 이스케이프된 "는 유지
    escaped = escaped.replace('&#x27;', '&#x27;')  # 이미 이스케이프된 '는 유지

    return escaped


def split_text_to_words(text):
    """텍스트를 단어로 분리하는 함수
    
    Args:
        text (str): 분리할 텍스트
        
    Returns:
        list: 분리된 단어들의 리스트
    """
    # <br/> 태그 제거
    text = BR_PATTERN.sub(' ', text)
    
    # 문장 부호 패턴 정의
    punctuation_pattern = r'[.。,，!！?？:：;；]'

    # 1. 먼저 공백으로 단어들을 분리
    words = text.strip().split()

    result = []
    for word in words:
        # 문장 부호가 없는 경우 그대로 반환
        if not PUNCTUATION_PATTERN.search(word):
            result.append(word)
            continue

        # 문장 부호로 시작하는 경우만 분리
        start_match = re.match(r'^([.。,，!！?？:：;；]+)(.+)$', word)
        if start_match:
            result.append(start_match.group(1))
            result.append(start_match.group(2))
            continue

        # 문장 부호가 단어 끝에 있는 경우는 변경하지 않고 그대로 둠
        if re.search(r'[.。,，!！?？:：;；]+$', word):
            result.append(word)
            continue

        # 그 외(문장 부호가 단어 중간에 있는 경우)는 원형 유지
        result.append(word)

    return result


def find_all_images(document):
    """문서에서 모든 이미지와 그 위치를 찾습니다."""
    images = []
    # Word 문서의 네임스페이스 정의
    nsmap = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    }
    
    for para_idx, para in enumerate(document.paragraphs):
        # 현재 단락의 전체 텍스트
        current_para_text = para.text
        
        # 이전 단락의 텍스트 (있는 경우)
        prev_para_text = document.paragraphs[para_idx-1].text if para_idx > 0 else ""
        
        # 다음 단락의 텍스트 (있는 경우)
        next_para_text = document.paragraphs[para_idx+1].text if para_idx < len(document.paragraphs)-1 else ""
        
        for run_idx, run in enumerate(para.runs):
            # 이미지 요소 찾기 (네임스페이스 명시)
            drawing = run._element.find('.//w:drawing', namespaces=nsmap)
            pict = run._element.find('.//w:pict', namespaces=nsmap)
            
            if drawing is not None or pict is not None:
                # 이미지 관계 ID 찾기
                blip = None
                if drawing is not None:
                    blip = drawing.find('.//a:blip', namespaces=nsmap)
                elif pict is not None:
                    blip = pict.find('.//a:blip', namespaces=nsmap)
                
                if blip is not None:
                    # 이미지 관계 ID 추출
                    embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if embed:
                        # 이미지 데이터 가져오기
                        image_part = document.part.related_parts[embed]
                        image_data = image_part.blob
                        
                        # 이미지가 발견되면 앞뒤 텍스트 출력
                        print(f"\n이미지 발견 (위치: 단락 {para_idx}, 런 {run_idx})")
                        print(f"이전 단락: {prev_para_text}")
                        print(f"현재 단락: {current_para_text}")
                        print(f"다음 단락: {next_para_text}")
                        print(f"이미지 크기: {len(image_data)} bytes")
                        
                        images.append({
                            'paragraph_index': para_idx,
                            'run_index': run_idx,
                            'paragraph': para,
                            'run': run,
                            'image_data': image_data,
                            'image_rid': embed
                        })
    return images


def analyze_image_context(document, image_info, window_size=2):
    """이미지 주변의 텍스트를 분석하여 이미지 설명을 찾습니다."""
    para_idx = image_info['paragraph_index']
    para = image_info['paragraph']
    
    # 이미지가 있는 단락의 전체 텍스트
    current_para_text = para.text
    
    # 이전 단락들의 텍스트 (window_size개)
    previous_paras = []
    for i in range(max(0, para_idx - window_size), para_idx):
        previous_paras.append(document.paragraphs[i].text)
    
    # 이후 단락들의 텍스트 (window_size개)
    next_paras = []
    for i in range(para_idx + 1, min(len(document.paragraphs), para_idx + window_size + 1)):
        next_paras.append(document.paragraphs[i].text)
    
    # 이미지 설명 패턴 찾기
    # 1. 대괄호로 둘러싸인 텍스트 찾기
    bracket_matches = BRACKET_PATTERN.finditer(current_para_text)
    
    # 2. 이미지 관련 키워드 찾기
    image_keywords = ['그림', '사진', '이미지', 'QR', '코드', '차트', '표', '다이어그램']
    
    # 이미지 설명 후보들
    candidates = []
    
    # 현재 단락에서 찾기
    for match in bracket_matches:
        text = match.group(1)
        if any(keyword in text for keyword in image_keywords):
            candidates.append({
                'text': text,
                'position': 'current',
                'distance': abs(match.start() - image_info['run_index'])
            })
    
    # 이전 단락들에서 찾기
    for idx, prev_text in enumerate(previous_paras):
        for match in BRACKET_PATTERN.finditer(prev_text):
            text = match.group(1)
            if any(keyword in text for keyword in image_keywords):
                candidates.append({
                    'text': text,
                    'position': 'previous',
                    'distance': len(previous_paras) - idx
                })
    
    # 이후 단락들에서 찾기
    for idx, next_text in enumerate(next_paras):
        for match in BRACKET_PATTERN.finditer(next_text):
            text = match.group(1)
            if any(keyword in text for keyword in image_keywords):
                candidates.append({
                    'text': text,
                    'position': 'next',
                    'distance': idx + 1
                })
    
    # 가장 적절한 설명 선택
    if candidates:
        # 거리와 위치를 고려하여 가장 적절한 설명 선택
        best_candidate = min(candidates, key=lambda x: (x['distance'], 
            0 if x['position'] == 'current' else 
            1 if x['position'] == 'previous' else 2))
        return best_candidate['text']
    
    return None

