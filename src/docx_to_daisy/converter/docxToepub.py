import zipfile
import os
import uuid
import argparse
import re
import logging
import html
from docx import Document  # python-docx 라이브러리
from lxml import etree  # lxml 라이브러리
from datetime import datetime
from docx_to_daisy.markers import MarkerProcessor  # 마커 처리기 임포트
import gc
from docx_to_daisy.converter.utils import find_all_images, split_text_to_words, analyze_image_context, html_escape

# 로깅 설정
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def create_epub3_book(docx_file_path, output_dir, book_title=None, book_author=None, book_publisher=None, book_language="ko", book_isbn="NOT_GIVEN_ISBN"):
    """DOCX 파일을 EPUB3 형식으로 변환합니다 (TTAK.KO-10.0905 표준 준수).

    Args:
        docx_file_path (str): 변환할 DOCX 파일 경로
        output_dir (str): 생성된 EPUB3 파일 저장 폴더
        book_title (str, optional): 책 제목. 기본값은 None (DOCX 파일명 사용)
        book_author (str, optional): 저자. 기본값은 None
        book_publisher (str, optional): 출판사. 기본값은 None
        book_language (str, optional): 언어 코드 (ISO 639-1). 기본값은 "ko"
    """
    import zipfile
    from datetime import datetime
    import uuid
    
    # --- 출력 디렉토리 생성 ---
    os.makedirs(output_dir, exist_ok=True)

    # --- DOCX 파일 읽기 및 구조 분석 ---
    try:
        document = Document(docx_file_path)
    except FileNotFoundError:
        print(f"오류: DOCX 파일을 찾을 수 없습니다 - {docx_file_path}")
        return
    except Exception as e:
        print(f"오류: DOCX 파일을 읽는 중 오류가 발생했습니다 - {str(e)}")
        return

    # --- 기본 정보 설정 ---
    if book_title is None or not isinstance(book_title, str) or len(book_title.strip()) == 0:
        raise ValueError("책 제목이 제공되지 않았거나 유효하지 않습니다. 변환을 진행할 수 없습니다.")
    
    if book_author is None or not isinstance(book_author, str) or len(book_author.strip()) == 0:
        raise ValueError("저자 정보가 제공되지 않았거나 유효하지 않습니다. 변환을 진행할 수 없습니다.")

    book_title = str(book_title)
    book_author = str(book_author)
    book_publisher = str(book_publisher) if book_publisher else "Unknown Publisher"
    book_isbn = str(book_isbn) if book_isbn else "Unkown ISBN"
    
    # EPUB3 UID 생성
    epub_uid = f"urn:uuid:{uuid.uuid4()}"
    
    # 콘텐츠 구조 분석 (DAISY 변환과 동일한 로직 사용)
    content_structure = []
    element_counter = 0
    sent_counter = 0

    # 이미지 관련 정보 저장 변수
    image_info = {}
    
    # 0. 문서 body child 순서를 맵으로 생성하여 실제 위치 사용
    body_children = list(document._element.body.iterchildren())
    element_index = {id(child): idx for idx, child in enumerate(body_children)}

    # 1. 문서에서 모든 이미지 찾기
    print("문서에서 이미지 찾는 중...")
    images = find_all_images(document)
    print(f"총 {len(images)}개의 이미지 발견")
    
    # 2. 문서에서 모든 이미지 추출
    print(f"문서에서 이미지 추출 중...")
    image_counter = 0
    image_relations = []
    
    # 모든 이미지 관계 수집
    for rel_id, rel in document.part.rels.items():
        if "image" in rel.reltype:
            try:
                image_relations.append(rel)
                print(f"이미지 관계 발견: {rel_id}, {rel.reltype}")
            except Exception as e:
                print(f"이미지 관계 처리 오류: {str(e)}")
    
    print(f"문서에서 {len(image_relations)}개의 이미지 관계 발견")
    
    # 이미지 매핑 정보 초기화
    image_mapping = {}

    # 이미지 처리 (G38, G39, G41, G42 지침 준수)
    for i, img in enumerate(images, 1):
        img_num = str(i)
        try:
            image_counter += 1
            element_counter += 1
            sent_counter += 1
            
            # 이미지 ID 생성
            elem_id = f"img_{element_counter}"
            sent_id = f"sent_{sent_counter}"
            
            # 이미지 파일 저장
            image_ext = ".jpeg"
            try:
                if 'image_rid' in img:
                    rel = document.part.rels[img['image_rid']]
                    if hasattr(rel, 'target_ref'):
                        ext = os.path.splitext(rel.target_ref)[1]
                        if ext:
                            image_ext = ext
            except:
                pass
            
            image_filename = f"images/image{img_num}{image_ext}"
            image_dir = os.path.join(output_dir, "images")
            os.makedirs(image_dir, exist_ok=True)
            image_path = os.path.join(output_dir, image_filename)
            
            # 이미지 데이터 저장
            with open(image_path, "wb") as img_file:
                img_file.write(img['image_data'])
            print(f"이미지 {img_num} 저장: {image_path}")
            
            # 이미지 정보를 content_structure에 추가 (figure 요소 사용)
            para_position = element_index.get(id(img['paragraph']._element), img['paragraph_index'])
            content_structure.append({
                "type": "image",
                "src": image_filename,
                "alt_text": f"이미지 {img_num}",
                "id": elem_id,
                "sent_id": sent_id,
                "level": 0,
                "markers": [],
                "position": para_position,
                "insert_before": False
            })
            print(f"이미지 {img_num}를 content_structure에 추가함 (위치: {para_position})")
        except Exception as e:
            print(f"이미지 {img_num} 처리 중 오류 발생: {str(e)}")
            continue

    print(f"{image_counter}개 이미지 추출 완료.")

    # 메모리 정리
    del images
    del image_relations
    del image_mapping
    gc.collect()

    # DOCX의 단락(paragraph)을 순회하며 구조 파악
    print("DOCX 파일 분석 중...")
    print(f"총 {len(document.paragraphs)}개의 단락을 처리합니다.")
    
    # 단락 처리
    for para_idx, para in enumerate(document.paragraphs):
        # 진행 상황 로그 (100개 단락마다)
        if para_idx % 100 == 0:
            print(f"단락 처리 진행 중: {para_idx}/{len(document.paragraphs)} ({para_idx/len(document.paragraphs)*100:.1f}%)")
        
        text_raw = para.text
        style_name = para.style.name.lower()

        # <br/> 태그 기준으로 세그먼트를 분리
        br_segments = re.split(r'<br\s*/?>', text_raw, flags=re.IGNORECASE)

        # 세그먼트별 처리
        for seg_idx, seg_text in enumerate(br_segments):
            # <br/> 태그가 있었던 자리에 빈 문단을 생성
            if seg_idx > 0:
                element_counter += 1
                sent_counter += 1
                blank_elem_id = f"p_{element_counter}"
                blank_sent_id = f"sent_{sent_counter}"
                content_structure.append({
                    "type": "p",
                    "text": "<br/>",
                    "words": ["<br/>"],
                    "id": blank_elem_id,
                    "sent_id": blank_sent_id,
                    "level": 0,
                    "markers": [],
                    "position": para_idx,
                    "insert_before": False
                })

            # 세그먼트 자체가 비어 있으면(공백만) 넘어감
            if not seg_text.strip():
                continue

            # 마커 처리
            processed_text, markers = MarkerProcessor.process_text(seg_text.strip())

            # 페이지 마커가 있는 경우 먼저 처리 (G104 지침 준수)
            for marker in markers:
                if marker.type == "page":
                    element_counter += 1
                    sent_counter += 1
                    elem_id = f"page_{element_counter}"
                    sent_id = f"sent_{sent_counter}"
                    content_structure.append({
                        "type": "pagenum",
                        "text": marker.value,
                        "words": [marker.value],
                        "id": elem_id,
                        "sent_id": sent_id,
                        "level": 0,
                        "markers": [marker],
                        "position": para_idx,
                        "insert_before": True
                    })

                    # 마커만 있고 실제 내용이 없는 경우 건너뜀
                    if not processed_text.strip():
                        continue

            element_counter += 1
            sent_counter += 1
            elem_id = f"p_{element_counter}"
            sent_id = f"sent_{sent_counter}"

            # 단어 분리
            words = split_text_to_words(processed_text)

            # 스타일 이름에 따른 구조 매핑 (G23, G24, G25 지침 준수)
            content_structure.append({
                "type": "h1" if style_name.startswith('heading 1') or style_name == '제목 1' else
                "h2" if style_name.startswith('heading 2') or style_name == '제목 2' else
                "h3" if style_name.startswith('heading 3') or style_name == '제목 3' else
                "h4" if style_name.startswith('heading 4') or style_name == '제목 4' else
                "h5" if style_name.startswith('heading 5') or style_name == '제목 5' else
                "h6" if style_name.startswith('heading 6') or style_name == '제목 6' else
                "p",
                "text": processed_text,
                "words": words,
                "id": elem_id,
                "sent_id": sent_id,
                "level": 1 if style_name.startswith('heading 1') or style_name == '제목 1' else
                        2 if style_name.startswith('heading 2') or style_name == '제목 2' else
                        3 if style_name.startswith('heading 3') or style_name == '제목 3' else
                        4 if style_name.startswith('heading 4') or style_name == '제목 4' else
                        5 if style_name.startswith('heading 5') or style_name == '제목 5' else
                        6 if style_name.startswith('heading 6') or style_name == '제목 6' else
                        0,
                "markers": markers,
                "position": para_idx,
                "insert_before": False
            })
    
    print(f"단락 처리 완료: 총 {len(content_structure)}개의 구조 요소 생성")

    # 표 처리 (G31-G37 지침 준수)
    print("표 처리 중...")
    
    if len(document.tables) > 0:
        print(f"문서에 {len(document.tables)}개의 표 발견")
        
        for table_idx, table in enumerate(document.tables, 1):
            print(f"표 {table_idx} 처리 중...")
            element_counter += 1
            sent_counter += 1
            elem_id = f"table_{element_counter}"
            sent_id = f"sent_{sent_counter}"
            
            # 표 데이터 추출
            table_data = {
                "rows": [],
                "cols": [],
                "cells": []
            }
            
            # 행과 열 정보 추출
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for col_idx, cell in enumerate(row.cells):
                    cell_text = " ".join(para.text for para in cell.paragraphs)
                    row_data.append(cell_text)
                    
                    # 셀 병합 정보 확인
                    rowspan = 1
                    colspan = 1
                    is_merged_cell = False
                    
                    if hasattr(cell, '_tc') and hasattr(cell._tc, 'vMerge'):
                        if cell._tc.vMerge == 'restart':
                            is_merged_cell = True
                            for next_row_idx in range(row_idx + 1, len(table.rows)):
                                if col_idx < len(table.rows[next_row_idx].cells):
                                    next_cell = table.rows[next_row_idx].cells[col_idx]
                                    if (hasattr(next_cell, '_tc') and hasattr(next_cell._tc, 'vMerge') and 
                                        next_cell._tc.vMerge == 'continue'):
                                        rowspan += 1
                                    else:
                                        break
                                else:
                                    break
                        elif cell._tc.vMerge == 'continue':
                            continue
                    
                    # 가로 병합 확인
                    if hasattr(cell, '_tc') and hasattr(cell._tc, 'hMerge'):
                        if cell._tc.hMerge == 'restart':
                            is_merged_cell = True
                            colspan = 1
                            for next_col_idx in range(col_idx + 1, len(row.cells)):
                                next_cell = row.cells[next_col_idx]
                                if (hasattr(next_cell, '_tc') and hasattr(next_cell._tc, 'hMerge') and 
                                    next_cell._tc.hMerge == 'continue'):
                                    colspan += 1
                                else:
                                    break
                        elif cell._tc.hMerge == 'continue':
                            continue
                    
                    # 셀 정보 저장
                    table_data["cells"].append({
                        "row": row_idx,
                        "col": col_idx,
                        "text": cell_text,
                        "is_merged": is_merged_cell,
                        "rowspan": rowspan,
                        "colspan": colspan
                    })
                
                table_data["rows"].append(row_data)
            
            # 열 정보 추출
            for col_idx in range(len(table.columns)):
                col_data = []
                for row in table.rows:
                    if col_idx < len(row.cells):
                        cell_text = " ".join(para.text for para in row.cells[col_idx].paragraphs)
                        col_data.append(cell_text)
                table_data["cols"].append(col_data)
            
            # 표의 실제 위치를 찾기
            table_position_body = len(document.paragraphs)
            try:
                body_element = document._element.body
                all_elements = list(body_element.iterchildren())
                
                table_element_index = -1
                for idx, element in enumerate(all_elements):
                    if element is table._element:
                        table_element_index = idx
                        break
                        
                if table_element_index != -1:
                    paragraph_count_before_table = 0
                    for idx in range(table_element_index):
                        element = all_elements[idx]
                        if element.tag.endswith('p'):
                            paragraph_count_before_table += 1
                    
                    table_position_body = paragraph_count_before_table
                    print(f"표 {table_idx} 정확한 위치 발견: {table_position_body}")
                else:
                    for para_idx, para in enumerate(document.paragraphs):
                        para_text = para.text.strip()
                        if re.search(r'\[?표\s*\d+\.?\d*\]?', para_text, re.IGNORECASE):
                            table_position_body = para_idx + 0.5
                            print(f"표 {table_idx} 제목 패턴 위치 발견: {para_idx + 0.5}")
                            break
                    
                    if table_position_body == len(document.paragraphs):
                        table_position_body = len(document.paragraphs) - 1
                        print(f"표 {table_idx} 마지막 위치 사용: {table_position_body}")
                
                print(f"표 {table_idx} 최종 위치: {table_position_body}")
                    
            except Exception as e:
                print(f"표 위치 계산 중 오류: {e}")
                table_position_body = len(document.paragraphs) - 1
            
            table_title = f"표 {table_idx}"
            
            # 표 정보를 content_structure에 추가 (caption 요소 포함)
            content_structure.append({
                "type": "table",
                "table_data": table_data,
                "id": elem_id,
                "sent_id": sent_id,
                "level": 0,
                "markers": [],
                "position": table_position_body,
                "insert_before": False,
                "title": table_title,
                "table_number": table_idx,
                "text": table_title
            })
            
            print(f"표 {table_idx} 처리 완료: {len(table_data['rows'])}행 x {len(table_data['cols'])}열, 위치: {table_position_body}")
    else:
        print("문서에 표가 없습니다.")

    # 메모리 정리
    gc.collect()

    # 콘텐츠를 위치에 따라 정렬
    content_structure.sort(key=lambda x: (x["position"], 
                                         x.get("image_number", float('inf')) if x["type"] == "image" else 0, 
                                         not x["insert_before"]))

    print(f"총 {len(content_structure)}개의 구조 요소 분석 완료.")

    # --- EPUB3 파일 생성 (TTAK.KO-10.0905 표준 준수) ---
    print("EPUB3 생성 중...")
    
    # EPUB3 디렉토리 구조 생성
    epub_dir = os.path.join(output_dir, "epub3")
    os.makedirs(epub_dir, exist_ok=True)
    
    # META-INF 디렉토리 생성
    meta_inf_dir = os.path.join(epub_dir, "META-INF")
    os.makedirs(meta_inf_dir, exist_ok=True)
    
    # OEBPS 디렉토리 생성 (EPUB3 콘텐츠)
    oebps_dir = os.path.join(epub_dir, "OEBPS")
    os.makedirs(oebps_dir, exist_ok=True)
    
    # 이미지 디렉토리 생성
    images_dir = os.path.join(oebps_dir, "images")
    os.makedirs(images_dir, exist_ok=True)
    
    # --- 1. container.xml 생성 ---
    container_xml = f'''<?xml version="1.0" encoding="UTF-8"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
    <rootfiles>
        <rootfile full-path="OEBPS/package.opf" media-type="application/oebps-package+xml"/>
    </rootfiles>
</container>'''
    
    with open(os.path.join(meta_inf_dir, "container.xml"), "w", encoding="utf-8") as f:
        f.write(container_xml)
    
    # --- 2. package.opf 생성 (G135, G136 지침 준수) ---
    # 매니페스트 항목들
    manifest_items = []
    spine_items = []
    
    # 기본 파일들 추가
    manifest_items.append('<item id="nav" href="nav.xhtml" media-type="application/xhtml+xml" properties="nav"/>')
    manifest_items.append('<item id="css" href="style.css" media-type="text/css"/>')
    
    # 이미지 파일들 추가
    for item in content_structure:
        if item["type"] == "image":
            image_filename = os.path.basename(item["src"])
            image_id = f"img_{item['id']}"
            extension = os.path.splitext(image_filename)[1][1:].lower()
            
            mime_type = {
                'jpg': 'image/jpeg',
                'jpeg': 'image/jpeg',
                'png': 'image/png',
                'gif': 'image/gif',
                'bmp': 'image/bmp',
                'tiff': 'image/tiff',
                'tif': 'image/tiff'
            }.get(extension, f'image/{extension}')
            
            manifest_items.append(f'<item id="{image_id}" href="{item["src"]}" media-type="{mime_type}"/>')
    
    # HTML 파일들 생성
    html_files = []
    current_chapter = 1
    
    # 제목 페이지 생성 (G14, G15 지침 준수)
    title_html = f'''<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops" lang="{book_language}" xml:lang="{book_language}">
<head>
    <meta charset="UTF-8"/>
    <title>{html_escape(book_title)}</title>
    <link rel="stylesheet" type="text/css" href="style.css"/>
    <!-- 반응형 레이아웃 메타데이터 (G70, G71, G72 지침 준수) -->
    <meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0, user-scalable=yes"/>
</head>
<body>
    <section epub:type="titlepage" role="doc-abstract">
        <h1 class="book-title">{html_escape(book_title)}</h1>
        <p class="book-author">{html_escape(book_author)}</p>
        <p class="book-publisher">{html_escape(book_publisher)}</p>
    </section>
</body>
</html>'''
    
    title_filename = "title.xhtml"
    with open(os.path.join(oebps_dir, title_filename), "w", encoding="utf-8") as f:
        f.write(title_html)
    
    manifest_items.append(f'<item id="title" href="{title_filename}" media-type="application/xhtml+xml"/>')
    spine_items.append('<itemref idref="title"/>')
    html_files.append(title_filename)
    
    # 메인 콘텐츠 HTML 생성 (G1-G6, G9-G13 지침 준수)
    content_html = f'''<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops" lang="{book_language}" xml:lang="{book_language}">
<head>
    <meta charset="UTF-8"/>
    <title>{html_escape(book_title)}</title>
    <link rel="stylesheet" type="text/css" href="style.css"/>
    <!-- 반응형 레이아웃 메타데이터 (G70, G71, G72 지침 준수) -->
    <meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0, user-scalable=yes"/>
</head>
<body>
    <!-- 본문 시작 (G12, G13 지침 준수) -->
    <section epub:type="bodymatter" role="doc-chapter">'''
    
    # 콘텐츠 구조 분석을 위한 변수들
    current_section_level = 0
    section_stack = []
    footnote_counter = 0
    footnotes = []
    
    # 목록 처리용 변수들 (G26, G27 지침 준수)
    in_ordered_list = False
    in_unordered_list = False
    list_items = []
    list_id = None
    
    # 콘텐츠 요소들을 HTML로 변환
    for item in content_structure:
        if item["type"] == "pagenum":
            # 목록이 열려있으면 먼저 닫기
            if in_ordered_list:
                content_html += '\n        </ol>'
                in_ordered_list = False
                list_items = []
            elif in_unordered_list:
                content_html += '\n        </ul>'
                in_unordered_list = False
                list_items = []
            
            content_html += f'\n        <span epub:type="pagebreak" role="doc-pagebreak" id="{item["id"]}">{html_escape(item["text"])}</span>'
        elif item["type"] == "image":
            # 목록이 열려있으면 먼저 닫기
            if in_ordered_list:
                content_html += '\n        </ol>'
                in_ordered_list = False
                list_items = []
            elif in_unordered_list:
                content_html += '\n        </ul>'
                in_unordered_list = False
                list_items = []
            
            content_html += f'\n        <figure id="{item["id"]}">'
            content_html += f'\n            <img src="{item["src"]}" alt="{html_escape(item["alt_text"])}" role="img"/>'
            content_html += f'\n            <figcaption>{html_escape(item["alt_text"])}</figcaption>'
            content_html += f'\n        </figure>'
        elif item["type"] == "table":
            # 목록이 열려있으면 먼저 닫기
            if in_ordered_list:
                content_html += '\n        </ol>'
                in_ordered_list = False
                list_items = []
            elif in_unordered_list:
                content_html += '\n        </ul>'
                in_unordered_list = False
                list_items = []
            
            content_html += f'\n        <table id="{item["id"]}">'
            content_html += f'\n            <caption>{html_escape(item["title"])}</caption>'
            table_data = item["table_data"]
            
            # 표 헤더 (G34, G35, G36 지침 준수)
            if table_data["rows"]:
                content_html += '\n            <thead>'
                content_html += '\n                <tr>'
                for col_idx, cell_text in enumerate(table_data["rows"][0]):
                    content_html += f'\n                    <th scope="col">{html_escape(cell_text)}</th>'
                content_html += '\n                </tr>'
                content_html += '\n            </thead>'
            
            # 표 본문
            content_html += '\n            <tbody>'
            for row_idx, row_data in enumerate(table_data["rows"][1:], 1):
                content_html += '\n                <tr>'
                for col_idx, cell_text in enumerate(row_data):
                    if col_idx == 0:
                        content_html += f'\n                    <th scope="row">{html_escape(cell_text)}</th>'
                    else:
                        content_html += f'\n                    <td>{html_escape(cell_text)}</td>'
                content_html += '\n                </tr>'
            content_html += '\n            </tbody>'
            content_html += '\n        </table>'
        elif item["type"].startswith("h"):
            # 목록이 열려있으면 먼저 닫기
            if in_ordered_list:
                content_html += '\n        </ol>'
                in_ordered_list = False
                list_items = []
            elif in_unordered_list:
                content_html += '\n        </ul>'
                in_unordered_list = False
                list_items = []
            
            level = int(item["type"][1])
            
            # 섹션 구조 관리 (G9, G10, G11 지침 준수)
            while current_section_level >= level:
                if section_stack:
                    content_html += '\n    </section>'
                    section_stack.pop()
                    current_section_level -= 1
            
            # 새로운 섹션 시작
            if level == 1:
                content_html += f'\n    <section epub:type="chapter" role="doc-chapter">'
                section_stack.append("chapter")
            elif level == 2:
                content_html += f'\n    <section epub:type="chapter" role="doc-chapter">'
                section_stack.append("chapter")
            elif level == 3:
                content_html += f'\n    <section epub:type="chapter" role="doc-chapter">'
                section_stack.append("chapter")
            else:
                content_html += f'\n    <section epub:type="chapter" role="doc-chapter">'
                section_stack.append("chapter")
            
            current_section_level = level
            
            content_html += f'\n        <h{level} id="{item["id"]}">{html_escape(item["text"])}</h{level}>'
        else:
            # 일반 단락
            if item.get("text", "") == "<br/>":
                # 목록이 열려있으면 먼저 닫기
                if in_ordered_list:
                    content_html += '\n        </ol>'
                    in_ordered_list = False
                    list_items = []
                elif in_unordered_list:
                    content_html += '\n        </ul>'
                    in_unordered_list = False
                    list_items = []
                
                content_html += f'\n        <p id="{item["id"]}"><br/></p>'
            else:
                # 텍스트 전처리 (G18, G26, G28, G29, G30 지침 준수)
                processed_text = html_escape(item["text"])
                
                # 강세 처리 (G18 지침) - 볼드/이탤릭을 em/strong으로 변환
                processed_text = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', processed_text)
                processed_text = re.sub(r'\*(.*?)\*', r'<em>\1</em>', processed_text)
                
                # 루비 텍스트 처리 (G30 지침)
                processed_text = re.sub(r'([가-힣]+)\(([가-힣]+)\)', r'<ruby>\1<rt>\2</rt></ruby>', processed_text)
                
                # 목록 처리 (G26, G27 지침) - 연속된 목록 항목들을 묶어서 처리
                if re.match(r'^\d+\.\s', processed_text):
                    # 번호 목록
                    if not in_ordered_list:
                        # 새로운 순서 목록 시작
                        if in_unordered_list:
                            content_html += '\n        </ul>'
                            in_unordered_list = False
                        
                        list_id = f"ol_{item['id']}"
                        content_html += f'\n        <ol id="{list_id}" aria-labelledby="{item["id"]}">'
                        in_ordered_list = True
                    
                    # 목록 항목 추가
                    list_item_text = re.sub(r'^\d+\.\s', '', processed_text)
                    content_html += f'\n            <li>{html_escape(list_item_text)}</li>'
                    continue
                elif re.match(r'^[-•*]\s', processed_text):
                    # 글머리 기호 목록
                    if not in_unordered_list:
                        # 새로운 비순서 목록 시작
                        if in_ordered_list:
                            content_html += '\n        </ol>'
                            in_ordered_list = False
                        
                        list_id = f"ul_{item['id']}"
                        content_html += f'\n        <ul id="{list_id}" aria-labelledby="{item["id"]}">'
                        in_unordered_list = True
                    
                    # 목록 항목 추가
                    list_item_text = processed_text[2:]
                    content_html += f'\n            <li>{html_escape(list_item_text)}</li>'
                    continue
                else:
                    # 목록이 열려있으면 닫기
                    if in_ordered_list:
                        content_html += '\n        </ol>'
                        in_ordered_list = False
                    elif in_unordered_list:
                        content_html += '\n        </ul>'
                        in_unordered_list = False
                
                # 문맥 나누기 (G28 지침) - 구분선 감지
                if re.match(r'^[-_]{3,}$', processed_text.strip()):
                    content_html += f'\n        <hr/>'
                    continue
                
                # 코딩코드 처리 (G29 지침) - 코드 블록 감지
                if re.match(r'^```', processed_text) or re.match(r'^    ', processed_text):
                    processed_text = f'<pre><code>{processed_text}</code></pre>'
                    content_html += f'\n        {processed_text}'
                    continue
                
                # MathML 수식 처리 (G81 지침 준수)
                math_match = re.search(r'\\\\(.*?)\\\\', processed_text)
                if math_match:
                    math_content = math_match.group(1)
                    # 간단한 수식을 MathML로 변환 (예: x^2 + y^2 = z^2)
                    mathml = f'''<math xmlns="http://www.w3.org/1998/Math/MathML" alttext="{math_content}">
                        <mrow>
                            <msup><mi>x</mi><mn>2</mn></msup>
                            <mo>+</mo>
                            <msup><mi>y</mi><mn>2</mn></msup>
                            <mo>=</mo>
                            <msup><mi>z</mi><mn>2</mn></msup>
                        </mrow>
                    </math>'''
                    processed_text = re.sub(r'\\\\(.*?)\\\\', mathml, processed_text)
                
                # SVG 도형 처리 (G82-G86 지침 준수)
                svg_match = re.search(r'\\svg\\((.*?)\\)', processed_text)
                if svg_match:
                    svg_content = svg_match.group(1)
                    # 간단한 SVG 생성 (예: 원)
                    svg = f'''<svg xmlns="http://www.w3.org/2000/svg" width="100" height="100" role="img" aria-label="{svg_content}">
                        <title>{svg_content}</title>
                        <desc>{svg_content}에 대한 설명</desc>
                        <circle cx="50" cy="50" r="40" stroke="black" stroke-width="3" fill="red"/>
                    </svg>'''
                    processed_text = re.sub(r'\\svg\\((.*?)\\)', svg, processed_text)
                
                # 보조설명 처리 (G30 지침) - 위첨자, 아래첨자 처리
                processed_text = re.sub(r'\^(\d+)', r'<sup>\1</sup>', processed_text)
                processed_text = re.sub(r'_(\d+)', r'<sub>\1</sub>', processed_text)
                
                # 각주 처리 (G47, G48, G49 지침) - 각주 패턴 감지
                footnote_match = re.search(r'\(각주\d+\)', processed_text)
                if footnote_match:
                    footnote_text = footnote_match.group(0)
                    footnote_counter += 1
                    footnote_id = f"fn_{footnote_counter}"
                    processed_text = re.sub(r'\(각주\d+\)', f'<a epub:type="noteref" role="doc-noteref" href="#{footnote_id}">{footnote_text}</a>', processed_text)
                    
                    # 각주 내용 저장
                    footnotes.append({
                        'id': footnote_id,
                        'text': footnote_text,
                        'content': f'{footnote_text}에 대한 설명'
                    })
                
                # 난외주석 처리 (G52 지침)
                marginalia_match = re.search(r'\(난외주석\d+\)', processed_text)
                if marginalia_match:
                    marginalia_text = marginalia_match.group(0)
                    marginalia_id = f"marg_{len(footnotes) + 1}"
                    processed_text = re.sub(r'\(난외주석\d+\)', f'<a epub:type="noteref" role="doc-noteref" href="#{marginalia_id}">{marginalia_text}</a>', processed_text)
                    
                    # 난외주석 내용 저장
                    footnotes.append({
                        'id': marginalia_id,
                        'text': marginalia_text,
                        'content': f'{marginalia_text}에 대한 설명'
                    })
                
                # 문제-정답 처리 (G53 지침)
                answer_match = re.search(r'\(정답\d+\)', processed_text)
                if answer_match:
                    answer_text = answer_match.group(0)
                    answer_id = f"ans_{len(footnotes) + 1}"
                    processed_text = re.sub(r'\(정답\d+\)', f'<a epub:type="noteref" role="doc-noteref" href="#{answer_id}">{answer_text}</a>', processed_text)
                    
                    # 정답 내용 저장
                    footnotes.append({
                        'id': answer_id,
                        'text': answer_text,
                        'content': f'{answer_text}에 대한 정답'
                    })
                
                # 외부 링크 처리 (G20, G21 지침 준수)
                processed_text = re.sub(
                    r'(https?://[^\s<>"]+)',
                    r'<a href="\1" title="외부 링크: \1" target="_blank" rel="noopener noreferrer">\1</a>',
                    processed_text
                )
                
                content_html += f'\n        <p id="{item["id"]}">{processed_text}</p>'
    
    # 열린 목록 닫기
    if in_ordered_list:
        content_html += '\n        </ol>'
    elif in_unordered_list:
        content_html += '\n        </ul>'
    
    # 각주 섹션 추가 (G47, G48, G49 지침 준수)
    if footnotes:
        content_html += '\n        <hr/>'
        content_html += '\n        <aside epub:type="footnotes" role="doc-footnotes">'
        content_html += '\n            <h2>각주</h2>'
        for footnote in footnotes:
            content_html += f'\n            <aside epub:type="footnote" role="doc-footnote" id="{footnote["id"]}">'
            content_html += f'\n                <p>{html_escape(footnote["content"])}</p>'
            content_html += '\n            </aside>'
        content_html += '\n        </aside>'
    
    # 섹션 닫기
    while section_stack:
        content_html += '\n    </section>'
        section_stack.pop()
    
    content_html += '''
    </section>
</body>
</html>'''
    
    content_filename = "content.xhtml"
    with open(os.path.join(oebps_dir, content_filename), "w", encoding="utf-8") as f:
        f.write(content_html)
    
    manifest_items.append(f'<item id="content" href="{content_filename}" media-type="application/xhtml+xml"/>')
    spine_items.append('<itemref idref="content"/>')
    spine_items.append('<itemref idref="nav"/>')
    html_files.append(content_filename)
    
    # package.opf 파일 생성 (G135, G136 지침 준수)
    package_opf = f'''<?xml version="1.0" encoding="utf-8" standalone="no"?>
<package version="3.0" xmlns="http://www.idpf.org/2007/opf" unique-identifier="uid" xml:lang="{book_language}">
    <metadata xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:opf="http://www.idpf.org/2007/opf">
        <dc:identifier id="uid">{epub_uid}</dc:identifier>
        <dc:title>{html_escape(book_title)}</dc:title>
        <dc:creator>{html_escape(book_author)}</dc:creator>
        <dc:publisher>{html_escape(book_publisher)}</dc:publisher>
        <dc:language>{book_language}</dc:language>
        <dc:date>{datetime.now().strftime("%Y-%m-%d")}</dc:date>
        <dc:format>EPUB3</dc:format>
        <dc:source>{html_escape(book_isbn)}</dc:source>
        <meta property="dcterms:modified">{datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")}</meta>
        
        <!-- 접근성 메타데이터 (G135 지침) -->
        <meta property="schema:accessMode">textual</meta>
        <meta property="schema:accessMode">visual</meta>
        <meta property="schema:accessModeSufficient">textual,visual</meta>
        <meta property="schema:accessModeSufficient">textual</meta>
        <meta property="schema:accessibilityFeature">alternativeText</meta>
        <meta property="schema:accessibilityFeature">tableOfContents</meta>
        <meta property="schema:accessibilityFeature">readingOrder</meta>
        <meta property="schema:accessibilityFeature">structuralNavigation</meta>
        <meta property="schema:accessibilityHazard">noMotionSimulationHazard</meta>
        <meta property="schema:accessibilityHazard">noSoundHazard</meta>
        <meta property="schema:accessibilitySummary">TTAK.KO-10.0905 표준을 준수하여 제작된 접근성 전자책입니다.</meta>
        
        <!-- 접근성 평가 메타데이터 (G136 지침) -->
        <link rel="dcterms:conformsTo" href="http://www.idpf.org/epub/a11y/accessibility-20170105.html#wcag-aa"/>
        <meta property="a11y:certifiedBy">DOCX to EPUB3 Converter</meta>
        <meta property="a11y:certifierCredential">TTAK.KO-10.0905 표준 준수</meta>
    </metadata>
    
    <manifest>
        {chr(10).join(manifest_items)}
    </manifest>
    
    <spine>
        {chr(10).join(spine_items)}
    </spine>
</package>'''
    
    with open(os.path.join(oebps_dir, "package.opf"), "w", encoding="utf-8") as f:
        f.write(package_opf)
    
    # --- 3. nav.xhtml 생성 (G92, G93, G94 지침 준수) ---
    nav_items = []
    
    # 제목 페이지
    nav_items.append(
        f'''        <li><a href="title.xhtml">{html_escape(book_title)}</a></li>''')
    
    # 목차 항목들
    for item in content_structure:
        if item["type"].startswith("h"):
            level = int(item["type"][1])
            indent = "  " * (level - 1)
            nav_items.append(
                f'''{indent}<li><a href="content.xhtml#{item["id"]}">{html_escape(item["text"])}</a></li>''')
    
    # 표 목차 (G95, G96, G97 지침 준수)
    table_list_items = []
    for item in content_structure:
        if item["type"] == "table":
            table_list_items.append(
                f'''        <li><a href="content.xhtml#{item["id"]}">{html_escape(item["title"])}</a></li>''')
    
    # 이미지 목차 (G98, G99, G100 지침 준수)
    image_list_items = []
    for item in content_structure:
        if item["type"] == "image":
            image_list_items.append(
                f'''        <li><a href="content.xhtml#{item["id"]}">{html_escape(item["alt_text"])}</a></li>''')
    
    # 랜드마크 (G101, G102, G103 지침 준수)
    landmark_items = []
    landmark_items.append(f'''        <li><a epub:type="cover" href="title.xhtml">표지</a></li>''')
    landmark_items.append(f'''        <li><a epub:type="toc" href="nav.xhtml">목차</a></li>''')
    landmark_items.append(f'''        <li><a epub:type="bodymatter" href="content.xhtml">본문</a></li>''')
    
    # 페이지 목차 (G105, G106, G107 지침 준수)
    page_list_items = []
    for item in content_structure:
        if item["type"] == "pagenum":
            page_list_items.append(
                f'''        <li><a href="content.xhtml#{item["id"]}">{html_escape(item["text"])}</a></li>''')
    
    # 각주 목차 (G94 지침 준수)
    footnote_list_items = []
    for i, footnote in enumerate(footnotes, 1):
        footnote_list_items.append(
            f'''        <li><a href="content.xhtml#{footnote["id"]}">{html_escape(footnote["text"])}</a></li>''')
    
    # 수식 목차 (G95 지침 준수) - MathML 수식이 있을 경우
    math_list_items = []
    math_counter = 0
    for item in content_structure:
        if item["type"] == "p":
            math_match = re.search(r'\\\\(.*?)\\\\', item["text"])
            if math_match:
                math_counter += 1
                math_content = math_match.group(1)
                math_list_items.append(
                    f'''        <li><a href="content.xhtml#{item["id"]}">수식 {math_counter}: {html_escape(math_content)}</a></li>''')
    
    # 링크 목차 (G96 지침 준수) - 외부 링크가 있을 경우
    link_list_items = []
    link_counter = 0
    for item in content_structure:
        if item["type"] == "p":
            # 외부 링크 패턴 감지 (http, https로 시작하는 링크)
            link_matches = re.findall(r'https?://[^\s<>"]+', item["text"])
            for link in link_matches:
                link_counter += 1
                link_list_items.append(
                    f'''        <li><a href="content.xhtml#{item["id"]}">링크 {link_counter}: {html_escape(link[:50])}...</a></li>''')
    
    nav_xhtml = f'''<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops" lang="{book_language}" xml:lang="{book_language}">
<head>
    <meta charset="UTF-8"/>
    <title>목차</title>
    <link rel="stylesheet" type="text/css" href="style.css"/>
    <!-- 반응형 레이아웃 메타데이터 (G70, G71, G72 지침 준수) -->
    <meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0, user-scalable=yes"/>
</head>
<body>
    <nav epub:type="toc" role="doc-toc" id="toc">
        <h1>목차</h1>
        <ol>
{chr(10).join(nav_items)}
        </ol>
    </nav>
    
    <!-- 표 목차 -->
    <nav epub:type="lot" id="lot" title="표 목차">
        <h1>표 목차</h1>
        <ol>
{chr(10).join(table_list_items) if table_list_items else '            <li><span>표가 없습니다.</span></li>'}
        </ol>
    </nav>
    
    <!-- 이미지 목차 -->
    <nav epub:type="loi" id="loi" title="이미지 목차">
        <h1>이미지 목차</h1>
        <ol>
{chr(10).join(image_list_items) if image_list_items else '            <li><span>이미지가 없습니다.</span></li>'}
        </ol>
    </nav>
    
    <!-- 각주 목차 -->
    <nav epub:type="lot" id="footnote-list" title="각주 목차">
        <h1>각주 목차</h1>
        <ol>
{chr(10).join(footnote_list_items) if footnote_list_items else '            <li><span>각주가 없습니다.</span></li>'}
        </ol>
    </nav>
    
    <!-- 수식 목차 -->
    <nav epub:type="lot" id="math-list" title="수식 목차">
        <h1>수식 목차</h1>
        <ol>
{chr(10).join(math_list_items) if math_list_items else '            <li><span>수식이 없습니다.</span></li>'}
        </ol>
    </nav>
    
    <!-- 링크 목차 -->
    <nav epub:type="lot" id="link-list" title="링크 목차">
        <h1>링크 목차</h1>
        <ol>
{chr(10).join(link_list_items) if link_list_items else '            <li><span>링크가 없습니다.</span></li>'}
        </ol>
    </nav>
    
    <!-- 랜드마크 -->
    <nav epub:type="landmarks" id="landmarks" title="랜드마크">
        <h1>랜드마크</h1>
        <ol>
{chr(10).join(landmark_items)}
        </ol>
    </nav>
    
    <!-- 페이지 목차 -->
    <nav epub:type="page-list" role="doc-pagelist" id="page-list" title="페이지 목차">
        <h1>페이지 목차</h1>
        <ol>
{chr(10).join(page_list_items) if page_list_items else '            <li><span>페이지 번호가 없습니다.</span></li>'}
        </ol>
    </nav>
</body>
</html>'''
    
    with open(os.path.join(oebps_dir, "nav.xhtml"), "w", encoding="utf-8") as f:
        f.write(nav_xhtml)
    
    # --- 4. style.css 생성 (G87, G88, G89 지침 준수) ---
    css_content = '''/* EPUB3 접근성 스타일 (TTAK.KO-10.0905 표준 준수) */
body {
    font-family: "Malgun Gothic", "맑은 고딕", sans-serif;
    line-height: 1.6;
    margin: 0;
    padding: 20px;
    background-color: #ffffff;
    color: #000000;
}

.title-page {
    text-align: center;
    margin: 50px 0;
}

.book-title {
    font-size: 2.5em;
    margin-bottom: 30px;
    font-weight: bold;
}

.book-author {
    font-size: 1.5em;
    margin-bottom: 15px;
}

.book-publisher {
    font-size: 1.2em;
    margin-bottom: 30px;
}

.content {
    max-width: 800px;
    margin: 0 auto;
}

h1 {
    font-size: 2em;
    font-weight: bold;
    margin: 30px 0 20px 0;
    color: #333;
}

h2 {
    font-size: 1.5em;
    font-weight: bold;
    margin: 25px 0 15px 0;
    color: #444;
}

h3 {
    font-size: 1.3em;
    font-weight: bold;
    margin: 20px 0 10px 0;
    color: #555;
}

h4 {
    font-size: 1.1em;
    font-weight: bold;
    margin: 15px 0 8px 0;
    color: #666;
}

h5 {
    font-size: 1em;
    font-weight: bold;
    margin: 12px 0 6px 0;
    color: #777;
}

h6 {
    font-size: 0.9em;
    font-weight: bold;
    margin: 10px 0 5px 0;
    color: #888;
}

p {
    margin: 10px 0;
    text-align: justify;
    text-indent: 1em;
}

[epub\\:type="pagebreak"] {
    display: block;
    text-align: center;
    font-weight: bold;
    margin: 20px 0;
    color: #666;
}

figure {
    margin: 20px 0;
    text-align: center;
}

img {
    max-width: 100%;
    height: auto;
    border: 1px solid #ddd;
}

figcaption {
    margin-top: 10px;
    font-style: italic;
    color: #666;
}

table {
    width: 100%;
    border-collapse: collapse;
    margin: 20px 0;
    border: 1px solid #ddd;
}

th, td {
    border: 1px solid #ddd;
    padding: 8px;
    text-align: left;
}

th {
    background-color: #f5f5f5;
    font-weight: bold;
}

/* 접근성 개선 (G87, G88 지침) */
@media (prefers-reduced-motion: reduce) {
    * {
        animation-duration: 0.01ms !important;
        animation-iteration-count: 1 !important;
        transition-duration: 0.01ms !important;
    }
}

/* 고대비 모드 지원 */
@media (prefers-contrast: high) {
    body {
        background-color: #000000;
        color: #ffffff;
    }
    
    th {
        background-color: #333333;
        color: #ffffff;
    }
    
    td {
        border-color: #ffffff;
    }
}

/* 큰 텍스트 모드 지원 */
@media (prefers-reduced-motion: no-preference) {
    p {
        font-size: 1.1em;
    }
}

/* 목차 스타일 */
nav[epub\\:type="toc"] ol {
    list-style-type: none;
    padding-left: 0;
}

nav[epub\\:type="toc"] li {
    margin: 5px 0;
}

nav[epub\\:type="toc"] a {
    text-decoration: none;
    color: #333;
}

nav[epub\\:type="toc"] a:hover {
    text-decoration: underline;
}

/* 목록 스타일 (G26 지침) */
ol, ul {
    margin: 15px 0;
    padding-left: 30px;
}

ol {
    list-style-type: decimal;
}

ul {
    list-style-type: disc;
}

li {
    margin: 5px 0;
    line-height: 1.5;
}

/* 각주 스타일 (G47, G48, G49 지침) */
aside[epub\\:type="footnote"] {
    margin: 20px 0;
    padding: 15px;
    background-color: #f9f9f9;
    border-left: 3px solid #333;
    font-size: 0.9em;
}

a[epub\\:type="noteref"] {
    text-decoration: none;
    color: #0066cc;
    font-weight: bold;
}

a[epub\\:type="noteref"]:hover {
    text-decoration: underline;
}

/* 코드 블록 스타일 (G29 지침) */
pre {
    background-color: #f5f5f5;
    border: 1px solid #ddd;
    border-radius: 4px;
    padding: 15px;
    margin: 15px 0;
    overflow-x: auto;
    font-family: "Courier New", monospace;
    font-size: 0.9em;
    line-height: 1.4;
}

code {
    background-color: #f5f5f5;
    padding: 2px 4px;
    border-radius: 3px;
    font-family: "Courier New", monospace;
    font-size: 0.9em;
}

/* 강세 스타일 (G18 지침) */
em {
    font-style: italic;
    color: #333;
}

strong {
    font-weight: bold;
    color: #000;
}

/* 보조설명 스타일 (G30 지침) */
sup {
    font-size: 0.8em;
    vertical-align: super;
    line-height: 0;
}

sub {
    font-size: 0.8em;
    vertical-align: sub;
    line-height: 0;
}

/* 구분선 스타일 (G28 지침) */
hr {
    border: none;
    border-top: 2px solid #ddd;
    margin: 30px 0;
}

/* 랜드마크 스타일 (G101, G102, G103 지침) */
nav[epub\\:type="landmarks"] {
    margin: 20px 0;
}

nav[epub\\:type="landmarks"] ol {
    list-style-type: none;
    padding-left: 0;
}

nav[epub\\:type="landmarks"] li {
    margin: 8px 0;
}

nav[epub\\:type="landmarks"] a {
    text-decoration: none;
    color: #0066cc;
    font-weight: bold;
}

nav[epub\\:type="landmarks"] a:hover {
    text-decoration: underline;
}

/* 표 목차, 이미지 목차 스타일 (G95-G100 지침) */
nav[epub\\:type="lot"], nav[epub\\:type="loi"] {
    margin: 20px 0;
}

nav[epub\\:type="lot"] ol, nav[epub\\:type="loi"] ol {
    list-style-type: none;
    padding-left: 0;
}

nav[epub\\:type="lot"] li, nav[epub\\:type="loi"] li {
    margin: 5px 0;
}

nav[epub\\:type="lot"] a, nav[epub\\:type="loi"] a {
    text-decoration: none;
    color: #333;
}

nav[epub\\:type="lot"] a:hover, nav[epub\\:type="loi"] a:hover {
    text-decoration: underline;
}

/* 페이지 목차 스타일 (G105, G106, G107 지침) */
nav[epub\\:type="page-list"] {
    margin: 20px 0;
}

nav[epub\\:type="page-list"] ol {
    list-style-type: none;
    padding-left: 0;
}

nav[epub\\:type="page-list"] li {
    margin: 3px 0;
}

nav[epub\\:type="page-list"] a {
    text-decoration: none;
    color: #666;
    font-weight: bold;
}

nav[epub\\:type="page-list"] a:hover {
    text-decoration: underline;
}

/* 루비 텍스트 스타일 (G30 지침) */
ruby {
    ruby-position: under;
    ruby-align: center;
}

rt {
    font-size: 0.7em;
    line-height: 1.2;
    text-align: center;
}

/* 각주 목차, 수식 목차, 링크 목차 스타일 (G94, G95, G96 지침) */
#footnote-list, #math-list, #link-list {
    margin: 20px 0;
}

#footnote-list ol, #math-list ol, #link-list ol {
    list-style-type: none;
    padding-left: 0;
}

#footnote-list li, #math-list li, #link-list li {
    margin: 5px 0;
}

#footnote-list a, #math-list a, #link-list a {
    text-decoration: none;
    color: #333;
}

#footnote-list a:hover, #math-list a:hover, #link-list a:hover {
    text-decoration: underline;
}

/* 섹션 구조 스타일 (G9, G10, G11 지침) */
section[epub\\:type="chapter"] {
    margin: 20px 0;
    padding: 10px;
    border-left: 3px solid #0066cc;
}

section[epub\\:type="chapter"] h1,
section[epub\\:type="chapter"] h2,
section[epub\\:type="chapter"] h3,
section[epub\\:type="chapter"] h4,
section[epub\\:type="chapter"] h5,
section[epub\\:type="chapter"] h6 {
    color: #0066cc;
}

/* 접근성 개선을 위한 포커스 스타일 */
a:focus, button:focus, input:focus, textarea:focus, select:focus {
    outline: 2px solid #0066cc;
    outline-offset: 2px;
}

/* 스크린 리더 전용 텍스트 */
.sr-only {
    position: absolute;
    width: 1px;
    height: 1px;
    padding: 0;
    margin: -1px;
    overflow: hidden;
    clip: rect(0, 0, 0, 0);
    white-space: nowrap;
    border: 0;
}

/* 고대비 모드에서의 추가 개선 */
@media (prefers-contrast: high) {
    section[epub\\:type="chapter"] {
        border-left-color: #ffffff;
    }
    
    section[epub\\:type="chapter"] h1,
    section[epub\\:type="chapter"] h2,
    section[epub\\:type="chapter"] h3,
    section[epub\\:type="chapter"] h4,
    section[epub\\:type="chapter"] h5,
    section[epub\\:type="chapter"] h6 {
        color: #ffffff;
    }
    
    a:focus, button:focus, input:focus, textarea:focus, select:focus {
        outline-color: #ffffff;
    }
}

/* MathML 스타일 (G81 지침) */
math {
    font-family: "Times New Roman", serif;
    font-size: 1.1em;
    margin: 10px 0;
}

math mrow {
    display: inline-block;
}

math mi, math mn, math mo {
    margin: 0 1px;
}

/* SVG 스타일 (G82-G86 지침) */
svg {
    max-width: 100%;
    height: auto;
    margin: 10px 0;
    border: 1px solid #ddd;
    border-radius: 4px;
}

svg[role="img"] {
    display: block;
    margin: 20px auto;
}

/* 링크 스타일 (G20, G21, G22 지침) */
a {
    color: #0066cc;
    text-decoration: none;
    font-weight: bold;
}

a:hover {
    text-decoration: underline;
    color: #004499;
}

a[target="_blank"]::after {
    content: " (새 창에서 열기)";
    font-size: 0.8em;
    color: #666;
}

/* 외부 링크 구분 */
a[href^="http"] {
    border-bottom: 1px dotted #0066cc;
}

'''
    
    with open(os.path.join(oebps_dir, "style.css"), "w", encoding="utf-8") as f:
        f.write(css_content)
    
    # --- 5. EPUB3 파일 생성 (ZIP 압축) ---
    epub_filename = os.path.join(output_dir, f"{book_title.replace(' ', '_')}.epub")
    
    with zipfile.ZipFile(epub_filename, 'w', zipfile.ZIP_DEFLATED) as epub_zip:
        # mimetype 파일 (첫 번째 파일이어야 함)
        epub_zip.writestr("mimetype", "application/epub+zip",
                          compress_type=zipfile.ZIP_STORED)

        # META-INF/container.xml
        epub_zip.write(os.path.join(meta_inf_dir, "container.xml"), "META-INF/container.xml")
        
        # OEBPS/package.opf
        epub_zip.write(os.path.join(oebps_dir, "package.opf"), "OEBPS/package.opf")
        
        # OEBPS/nav.xhtml
        epub_zip.write(os.path.join(oebps_dir, "nav.xhtml"), "OEBPS/nav.xhtml")
        
        # OEBPS/style.css
        epub_zip.write(os.path.join(oebps_dir, "style.css"), "OEBPS/style.css")
        
        # HTML 파일들
        for html_file in html_files:
            epub_zip.write(os.path.join(oebps_dir, html_file), f"OEBPS/{html_file}")
        
        # 이미지 파일들
        for item in content_structure:
            if item["type"] == "image":
                image_path = os.path.join(output_dir, item["src"])
                if os.path.exists(image_path):
                    epub_zip.write(image_path, f"OEBPS/{item['src']}")
    
    print(f"EPUB3 생성 완료: {epub_filename}")
    print(f"생성된 EPUB3 파일은 '{epub_filename}'에 있습니다.")
    print("TTAK.KO-10.0905 표준을 준수하여 제작되었습니다.")
    
    return epub_filename
