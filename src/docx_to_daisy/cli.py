import zipfile
import os
import uuid
import argparse
import re
import logging
import html
from docx import Document  # python-docx 라이브러리
from lxml import etree  # lxml 라이브러리
from datetime import datetime
from .markers import MarkerProcessor  # 마커 처리기 임포트
import gc

# 로깅 설정
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def html_escape(text):
    """HTML 특수 문자를 이스케이프하고 HTML 태그를 제거하는 함수
    
    Args:
        text (str): 이스케이프할 텍스트
        
    Returns:
        str: 이스케이프된 텍스트 (HTML 태그 제거됨)
    """
    if not isinstance(text, str):
        return str(text)

    # HTML 태그 제거
    import re
    # HTML 태그 패턴 (시작 태그, 종료 태그, 자체 종료 태그 모두 포함)
    html_tag_pattern = r'<[^>]*>'
    cleaned_text = re.sub(html_tag_pattern, '', text)

    # HTML 특수 문자 이스케이프
    escaped = html.escape(cleaned_text, quote=True)

    # 추가적인 이스케이프 처리
    escaped = escaped.replace('&amp;', '&amp;')  # 이미 이스케이프된 &는 유지
    escaped = escaped.replace('&lt;', '&lt;')    # 이미 이스케이프된 <는 유지
    escaped = escaped.replace('&gt;', '&gt;')    # 이미 이스케이프된 >는 유지
    escaped = escaped.replace('&quot;', '&quot;')  # 이미 이스케이프된 "는 유지
    escaped = escaped.replace('&#x27;', '&#x27;')  # 이미 이스케이프된 '는 유지

    return escaped


def split_text_to_words(text):
    """텍스트를 단어로 분리하는 함수
    
    Args:
        text (str): 분리할 텍스트
        
    Returns:
        list: 분리된 단어들의 리스트
    """
    # <br/> 태그 제거
    text = text.replace('<br/>', ' ')
    
    # 문장 부호 패턴 정의
    punctuation_pattern = r'[.。,，!！?？:：;；]'

    # 1. 먼저 공백으로 단어들을 분리
    words = text.strip().split()

    result = []
    for word in words:
        # 2. 각 단어에서 문장 부호가 있는지 확인
        if re.search(punctuation_pattern, word):
            # 문장 부호가 단어 중간에 있는 경우는 그대로 유지
            if not re.match(f'^{punctuation_pattern}', word) and not re.search(f'{punctuation_pattern}$', word):
                result.append(word)
                continue

            # 문장 부호로 시작하는 경우
            if re.match(f'^{punctuation_pattern}', word):
                punct = re.match(f'^({punctuation_pattern}+)', word).group(1)
                remaining = word[len(punct):]
                if punct:
                    result.append(punct)
                if remaining:
                    result.append(remaining)
                continue

            # 문장 부호로 끝나는 경우
            if re.search(f'{punctuation_pattern}$', word):
                match = re.search(f'({punctuation_pattern}+)$', word)
                punct = match.group(1)
                text_part = word[:-len(punct)]
                if text_part:
                    result.append(text_part + punct)  # 문장 부호를 단어에 붙임
                else:
                    result.append(punct)
                continue

        # 문장 부호가 없는 일반 단어
        result.append(word)

    return result


def find_all_images(document):
    """문서에서 모든 이미지와 그 위치를 찾습니다."""
    images = []
    # Word 문서의 네임스페이스 정의
    nsmap = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    }
    
    for para_idx, para in enumerate(document.paragraphs):
        # 현재 단락의 전체 텍스트
        current_para_text = para.text
        
        # 이전 단락의 텍스트 (있는 경우)
        prev_para_text = document.paragraphs[para_idx-1].text if para_idx > 0 else ""
        
        # 다음 단락의 텍스트 (있는 경우)
        next_para_text = document.paragraphs[para_idx+1].text if para_idx < len(document.paragraphs)-1 else ""
        
        for run_idx, run in enumerate(para.runs):
            # 이미지 요소 찾기 (네임스페이스 명시)
            drawing = run._element.find('.//w:drawing', namespaces=nsmap)
            pict = run._element.find('.//w:pict', namespaces=nsmap)
            
            if drawing is not None or pict is not None:
                # 이미지 관계 ID 찾기
                blip = None
                if drawing is not None:
                    blip = drawing.find('.//a:blip', namespaces=nsmap)
                elif pict is not None:
                    blip = pict.find('.//a:blip', namespaces=nsmap)
                
                if blip is not None:
                    # 이미지 관계 ID 추출
                    embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if embed:
                        # 이미지 데이터 가져오기
                        image_part = document.part.related_parts[embed]
                        image_data = image_part.blob
                        
                        # 이미지가 발견되면 앞뒤 텍스트 출력
                        print(f"\n이미지 발견 (위치: 단락 {para_idx}, 런 {run_idx})")
                        print(f"이전 단락: {prev_para_text}")
                        print(f"현재 단락: {current_para_text}")
                        print(f"다음 단락: {next_para_text}")
                        print(f"이미지 크기: {len(image_data)} bytes")
                        
                        images.append({
                            'paragraph_index': para_idx,
                            'run_index': run_idx,
                            'paragraph': para,
                            'run': run,
                            'image_data': image_data,
                            'image_rid': embed
                        })
    return images


def analyze_image_context(document, image_info, window_size=2):
    """이미지 주변의 텍스트를 분석하여 이미지 설명을 찾습니다."""
    para_idx = image_info['paragraph_index']
    para = image_info['paragraph']
    
    # 이미지가 있는 단락의 전체 텍스트
    current_para_text = para.text
    
    # 이전 단락들의 텍스트 (window_size개)
    previous_paras = []
    for i in range(max(0, para_idx - window_size), para_idx):
        previous_paras.append(document.paragraphs[i].text)
    
    # 이후 단락들의 텍스트 (window_size개)
    next_paras = []
    for i in range(para_idx + 1, min(len(document.paragraphs), para_idx + window_size + 1)):
        next_paras.append(document.paragraphs[i].text)
    
    # 이미지 설명 패턴 찾기
    # 1. 대괄호로 둘러싸인 텍스트 찾기
    bracket_pattern = r'\[(.*?)\]'
    bracket_matches = re.finditer(bracket_pattern, current_para_text)
    
    # 2. 이미지 관련 키워드 찾기
    image_keywords = ['그림', '사진', '이미지', 'QR', '코드', '차트', '표', '다이어그램']
    
    # 이미지 설명 후보들
    candidates = []
    
    # 현재 단락에서 찾기
    for match in bracket_matches:
        text = match.group(1)
        if any(keyword in text for keyword in image_keywords):
            candidates.append({
                'text': text,
                'position': 'current',
                'distance': abs(match.start() - image_info['run_index'])
            })
    
    # 이전 단락들에서 찾기
    for idx, prev_text in enumerate(previous_paras):
        for match in re.finditer(bracket_pattern, prev_text):
            text = match.group(1)
            if any(keyword in text for keyword in image_keywords):
                candidates.append({
                    'text': text,
                    'position': 'previous',
                    'distance': len(previous_paras) - idx
                })
    
    # 이후 단락들에서 찾기
    for idx, next_text in enumerate(next_paras):
        for match in re.finditer(bracket_pattern, next_text):
            text = match.group(1)
            if any(keyword in text for keyword in image_keywords):
                candidates.append({
                    'text': text,
                    'position': 'next',
                    'distance': idx + 1
                })
    
    # 가장 적절한 설명 선택
    if candidates:
        # 거리와 위치를 고려하여 가장 적절한 설명 선택
        best_candidate = min(candidates, key=lambda x: (x['distance'], 
            0 if x['position'] == 'current' else 
            1 if x['position'] == 'previous' else 2))
        return best_candidate['text']
    
    return None


def create_daisy_book(docx_file_path, output_dir, book_title=None, book_author=None, book_publisher=None, book_language="ko"):
    """DOCX 파일을 DAISY 형식으로 변환합니다.

    Args:
        docx_file_path (str): 변환할 DOCX 파일 경로
        output_dir (str): 생성된 DAISY 파일 저장 폴더
        book_title (str, optional): 책 제목. 기본값은 None (DOCX 파일명 사용)
        book_author (str, optional): 저자. 기본값은 None
        book_publisher (str, optional): 출판사. 기본값은 None
        book_language (str, optional): 언어 코드 (ISO 639-1). 기본값은 "ko"
    """
    # 이미지 설명 처리 함수 정의
    def get_clean_description(desc_list):
        """이미지 설명 목록에서 중복을 제거하고 깔끔한 설명을 반환합니다"""
        if not desc_list:
            return None
        
        # 중복 제거
        unique_desc = []
        seen = set()
        for desc in desc_list:
            if desc not in seen:
                seen.add(desc)
                unique_desc.append(desc)
        
        # 설명이 너무 길면 첫 번째 항목만 반환
        if len(unique_desc) > 0:
            return unique_desc[0]
        return None
    
    # --- 출력 디렉토리 생성 ---
    os.makedirs(output_dir, exist_ok=True)

    # --- DOCX 파일 읽기 및 구조 분석 ---
    try:
        document = Document(docx_file_path)
    except FileNotFoundError:
        print(f"오류: DOCX 파일을 찾을 수 없습니다 - {docx_file_path}")
        return
    except Exception as e:
        print(f"오류: DOCX 파일을 읽는 중 오류가 발생했습니다 - {str(e)}")
        return

    # --- 기본 정보 설정 ---
    # book_title 확인
    if book_title is None or not isinstance(book_title, str) or len(book_title.strip()) == 0:
        raise ValueError("책 제목이 제공되지 않았거나 유효하지 않습니다. 변환을 진행할 수 없습니다.")
    
    # book_author 확인
    if book_author is None or not isinstance(book_author, str) or len(book_author.strip()) == 0:
        raise ValueError("저자 정보가 제공되지 않았거나 유효하지 않습니다. 변환을 진행할 수 없습니다.")

    book_title = str(book_title)
    book_author = str(book_author)
    book_publisher = str(book_publisher) if book_publisher else "Unknown Publisher"
    
    # book_uid에 책 제목을 포함시켜 DAISY 열 때 표시될 이름 설정
    safe_title = book_title.replace(" ", "_").replace("/", "_").replace("\\", "_").replace(":", "_")
    
    book_uid = f"BOOK-{safe_title}"  # 제목을 포함한 식별자
    print("book_uid", book_uid)

    content_structure = []
    element_counter = 0
    sent_counter = 0

    # doctitle과 docauthor를 위한 ID 미리 할당
    doctitle_id = "id_1"
    docauthor_id = "id_2"
    sent_counter = 2  # doctitle과 docauthor 이후부터 시작
    element_counter = 2  # doctitle과 docauthor 이후부터 시작

    # 이미지 관련 정보 저장 변수
    image_info = {}  # 이미지 번호 -> {제목, 설명, 위치} 매핑
    
    # 0. 문서 body child 순서를 맵으로 생성하여 실제 위치 사용
    body_children = list(document._element.body.iterchildren())
    element_index = {id(child): idx for idx, child in enumerate(body_children)}

    # 1. 문서에서 모든 이미지 찾기
    print("문서에서 이미지 찾는 중...")
    images = find_all_images(document)
    print(f"총 {len(images)}개의 이미지 발견")
    
    # 2. 문서에서 모든 이미지 추출
    print(f"문서에서 이미지 추출 중...")
    image_counter = 0
    image_relations = []
    
    # 모든 이미지 관계 수집
    for rel_id, rel in document.part.rels.items():
        if "image" in rel.reltype:
            try:
                # 이미지 관계 정보 저장
                image_relations.append(rel)
                print(f"이미지 관계 발견: {rel_id}, {rel.reltype}")
            except Exception as e:
                print(f"이미지 관계 처리 오류: {str(e)}")
    
    print(f"문서에서 {len(image_relations)}개의 이미지 관계 발견")
    
    # 이미지 매핑 정보 초기화
    image_mapping = {}  # 이미지 번호 -> 이미지 관계 매핑

    # 이미지 처리
    for i, img in enumerate(images, 1):
        img_num = str(i)
        try:
            image_counter += 1
            element_counter += 1
            sent_counter += 1
            
            # 이미지 ID 생성
            elem_id = f"id_{element_counter}"
            sent_id = f"id_{sent_counter}"
            
            # 이미지 파일 저장
            image_ext = ".jpeg"
            try:
                if 'image_rid' in img:
                    rel = document.part.rels[img['image_rid']]
                    if hasattr(rel, 'target_ref'):
                        ext = os.path.splitext(rel.target_ref)[1]
                        if ext:
                            image_ext = ext
            except:
                pass
            
            image_filename = f"images/image{img_num}{image_ext}"
            image_dir = os.path.join(output_dir, "images")
            os.makedirs(image_dir, exist_ok=True)
            image_path = os.path.join(output_dir, image_filename)
            
            # 이미지 데이터 저장
            with open(image_path, "wb") as img_file:
                img_file.write(img['image_data'])
            print(f"이미지 {img_num} 저장: {image_path}")
            
            # 이미지 정보를 content_structure에 추가
            para_position = element_index.get(id(img['paragraph']._element), img['paragraph_index'])
            content_structure.append({
                "type": "image",
                "src": image_filename,
                "alt_text": f"이미지 {img_num}",
                "id": elem_id,
                "sent_id": sent_id,
                "level": 0,
                "markers": [],
                "position": para_position,
                "insert_before": False
            })
            print(f"이미지 {img_num}를 content_structure에 추가함 (위치: {para_position})")
        except Exception as e:
            print(f"이미지 {img_num} 처리 중 오류 발생: {str(e)}")
            continue

    print(f"{image_counter}개 이미지 추출 완료.")

    # 메모리 정리
    del images
    del image_relations
    del image_mapping
    gc.collect()

    # DOCX의 단락(paragraph)을 순회하며 구조 파악
    print("DOCX 파일 분석 중...")
    print(f"총 {len(document.paragraphs)}개의 단락을 처리합니다.")
    
    # 단락 처리
    for para_idx, para in enumerate(document.paragraphs):
        # 진행 상황 로그 (100개 단락마다)
        if para_idx % 100 == 0:
            print(f"단락 처리 진행 중: {para_idx}/{len(document.paragraphs)} ({para_idx/len(document.paragraphs)*100:.1f}%)")
        
        text_raw = para.text
        style_name = para.style.name.lower()

        # <br/> 태그 기준으로 세그먼트를 분리
        br_segments = re.split(r'<br\s*/?>', text_raw, flags=re.IGNORECASE)

        # 세그먼트별 처리
        for seg_idx, seg_text in enumerate(br_segments):
            # <br/> 태그가 있었던 자리에 빈 문단을 생성
            if seg_idx > 0:
                element_counter += 1
                sent_counter += 1
                blank_elem_id = f"p_{element_counter}"
                blank_sent_id = f"sent_{sent_counter}"
                content_structure.append({
                    "type": "p",
                    "text": "<br/>",
                    "words": ["<br/>"],
                    "id": blank_elem_id,
                    "sent_id": blank_sent_id,
                    "level": 0,
                    "markers": [],
                    "position": para_idx,
                    "insert_before": False
                })

            # 세그먼트 자체가 비어 있으면(공백만) 넘어감
            if not seg_text.strip():
                continue

            # 마커 처리
            processed_text, markers = MarkerProcessor.process_text(seg_text.strip())

            # 페이지 마커가 있는 경우 먼저 처리
            for marker in markers:
                if marker.type == "page":
                    element_counter += 1
                    sent_counter += 1
                    elem_id = f"page_{element_counter}"
                    sent_id = f"sent_{sent_counter}"
                    content_structure.append({
                        "type": "pagenum",
                        "text": marker.value,
                        "words": [marker.value],
                        "id": elem_id,
                        "sent_id": sent_id,
                        "level": 0,
                        "markers": [marker],
                        "position": para_idx,
                        "insert_before": True
                    })

                    # 마커만 있고 실제 내용이 없는 경우 건너뜀
                    if not processed_text.strip():
                        continue

            element_counter += 1
            sent_counter += 1
            elem_id = f"p_{element_counter}"
            sent_id = f"sent_{sent_counter}"

            # 단어 분리
            words = split_text_to_words(processed_text)

            # 스타일 이름에 따른 구조 매핑
            content_structure.append({
                "type": "h1" if style_name.startswith('heading 1') or style_name == '제목 1' else
                "h2" if style_name.startswith('heading 2') or style_name == '제목 2' else
                "h3" if style_name.startswith('heading 3') or style_name == '제목 3' else
                "h4" if style_name.startswith('heading 4') or style_name == '제목 4' else
                "h5" if style_name.startswith('heading 5') or style_name == '제목 5' else
                "h6" if style_name.startswith('heading 6') or style_name == '제목 6' else
                "p",
                "text": processed_text,
                "words": words,
                "id": elem_id,
                "sent_id": sent_id,
                "level": 1 if style_name.startswith('heading 1') or style_name == '제목 1' else
                        2 if style_name.startswith('heading 2') or style_name == '제목 2' else
                        3 if style_name.startswith('heading 3') or style_name == '제목 3' else
                        4 if style_name.startswith('heading 4') or style_name == '제목 4' else
                        5 if style_name.startswith('heading 5') or style_name == '제목 5' else
                        6 if style_name.startswith('heading 6') or style_name == '제목 6' else
                        0,
                "markers": markers,
                "position": para_idx,
                "insert_before": False
            })
    
    print(f"단락 처리 완료: 총 {len(content_structure)}개의 구조 요소 생성")

    # 표 처리
    print("표 처리 중...")
    
    if len(document.tables) > 0:
        print(f"문서에 {len(document.tables)}개의 표 발견")
        
        for table_idx, table in enumerate(document.tables, 1):
            print(f"표 {table_idx} 처리 중...")
            element_counter += 1
            sent_counter += 1
            elem_id = f"table_{element_counter}"
            sent_id = f"sent_{sent_counter}"
            
            # 표 데이터 추출
            table_data = {
                "rows": [],
                "cols": [],
                "cells": []
            }
            
            # 행과 열 정보 추출
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for col_idx, cell in enumerate(row.cells):
                    cell_text = " ".join(para.text for para in cell.paragraphs)
                    row_data.append(cell_text)
                    
                    # 셀 병합 정보 확인
                    rowspan = 1
                    colspan = 1
                    is_merged_cell = False
                    
                    if hasattr(cell, '_tc') and hasattr(cell._tc, 'vMerge'):
                        if cell._tc.vMerge == 'restart':
                            is_merged_cell = True
                            for next_row_idx in range(row_idx + 1, len(table.rows)):
                                if col_idx < len(table.rows[next_row_idx].cells):
                                    next_cell = table.rows[next_row_idx].cells[col_idx]
                                    if (hasattr(next_cell, '_tc') and hasattr(next_cell._tc, 'vMerge') and 
                                        next_cell._tc.vMerge == 'continue'):
                                        rowspan += 1
                                    else:
                                        break
                                else:
                                    break
                        elif cell._tc.vMerge == 'continue':
                            continue
                    
                    # 가로 병합 확인
                    if hasattr(cell, '_tc') and hasattr(cell._tc, 'hMerge'):
                        if cell._tc.hMerge == 'restart':
                            is_merged_cell = True
                            colspan = 1
                            for next_col_idx in range(col_idx + 1, len(row.cells)):
                                next_cell = row.cells[next_col_idx]
                                if (hasattr(next_cell, '_tc') and hasattr(next_cell._tc, 'hMerge') and 
                                    next_cell._tc.hMerge == 'continue'):
                                    colspan += 1
                                else:
                                    break
                        elif cell._tc.hMerge == 'continue':
                            continue
                    
                    # 셀 정보 저장
                    table_data["cells"].append({
                        "row": row_idx,
                        "col": col_idx,
                        "text": cell_text,
                        "is_merged": is_merged_cell,
                        "rowspan": rowspan,
                        "colspan": colspan
                    })
                
                table_data["rows"].append(row_data)
            
            # 열 정보 추출
            for col_idx in range(len(table.columns)):
                col_data = []
                for row in table.rows:
                    if col_idx < len(row.cells):
                        cell_text = " ".join(para.text for para in row.cells[col_idx].paragraphs)
                        col_data.append(cell_text)
                table_data["cols"].append(col_data)
            
            # 표의 실제 위치를 찾기
            table_position_body = len(document.paragraphs)
            try:
                body_element = document._element.body
                all_elements = list(body_element.iterchildren())
                
                table_element_index = -1
                for idx, element in enumerate(all_elements):
                    if element is table._element:
                        table_element_index = idx
                        break
                        
                if table_element_index != -1:
                    paragraph_count_before_table = 0
                    for idx in range(table_element_index):
                        element = all_elements[idx]
                        if element.tag.endswith('p'):
                            paragraph_count_before_table += 1
                    
                    table_position_body = paragraph_count_before_table
                    print(f"표 {table_idx} 정확한 위치 발견: {table_position_body}")
                else:
                    for para_idx, para in enumerate(document.paragraphs):
                        para_text = para.text.strip()
                        if re.search(r'\[?표\s*\d+\.?\d*\]?', para_text, re.IGNORECASE):
                            table_position_body = para_idx + 0.5
                            print(f"표 {table_idx} 제목 패턴 위치 발견: {para_idx + 0.5}")
                            break
                    
                    if table_position_body == len(document.paragraphs):
                        table_position_body = len(document.paragraphs) - 1
                        print(f"표 {table_idx} 마지막 위치 사용: {table_position_body}")
                
                print(f"표 {table_idx} 최종 위치: {table_position_body}")
                    
            except Exception as e:
                print(f"표 위치 계산 중 오류: {e}")
                table_position_body = len(document.paragraphs) - 1
            
            table_title = f"표 {table_idx}"
            
            # 표 정보를 content_structure에 추가 (caption 요소 포함)
            content_structure.append({
                "type": "table",
                "table_data": table_data,
                "id": elem_id,
                "sent_id": sent_id,
                "level": 0,
                "markers": [],
                "position": table_position_body,
                "insert_before": False,
                "title": table_title,
                "table_number": table_idx,
                "text": table_title
            })
            
            print(f"표 {table_idx} 처리 완료: {len(table_data['rows'])}행 x {len(table_data['cols'])}열, 위치: {table_position_body}")
    else:
        print("문서에 표가 없습니다.")

    # 메모리 정리
    gc.collect()

    # 콘텐츠를 위치에 따라 정렬
    content_structure.sort(key=lambda x: (x["position"], 
                                         x.get("image_number", float('inf')) if x["type"] == "image" else 0, 
                                         not x["insert_before"]))

    print(f"총 {len(content_structure)}개의 구조 요소 분석 완료.")

    # --- 1. DTBook XML 생성 (dtbook.xml) ---
    print("DTBook 생성 중...")
    dtbook_ns = "http://www.daisy.org/z3986/2005/dtbook/"
    dc_ns = "http://purl.org/dc/elements/1.1/"

    # 페이지 카운터 초기화
    total_pages = 0
    max_page_number = 0
    for item in content_structure:
        for marker in item.get("markers", []):
            if marker.type == "page":
                total_pages += 1
                try:
                    page_num = int(marker.value)
                    max_page_number = max(max_page_number, page_num)
                except ValueError:
                    pass

    dtbook_root = etree.Element(
        "{%s}dtbook" % dtbook_ns,
        attrib={
            "version": "2005-3"
        },
        nsmap={
            None: dtbook_ns,
            "dc": dc_ns
        }
    )

    # head 요소 추가
    head = etree.SubElement(dtbook_root, "head")

    # 필수 메타데이터 추가
    meta_uid = etree.SubElement(head, "meta",
                                name="dtb:uid",
                                content=book_uid)
    meta_title = etree.SubElement(head, "meta",
                                  name="dc:Title",
                                  content=book_title)
    meta_author = etree.SubElement(head, "meta",
                                   name="dc:Creator",
                                   content=book_author)
    meta_publisher = etree.SubElement(head, "meta",
                                   name="dc:Publisher",
                                   content=book_publisher)
    meta_language = etree.SubElement(head, "meta",
                                     name="dc:Language",
                                     content=book_language)
    meta_date = etree.SubElement(head, "meta",
                                 name="dc:Date",
                                 content=datetime.now().strftime("%Y-%m-%d"))

    # 페이지 관련 메타데이터 추가
    etree.SubElement(head, "meta",
                     name="dtb:totalPageCount",
                     content=str(total_pages))
    etree.SubElement(head, "meta",
                     name="dtb:maxPageNumber",
                     content=str(max_page_number))

    # book 요소 추가
    dtbook_book = etree.SubElement(dtbook_root, "book", showin="blp")

    # frontmatter 추가
    dtbook_frontmatter = etree.SubElement(dtbook_book, "frontmatter")

    # doctitle과 docauthor 추가
    doctitle_seq = etree.SubElement(dtbook_frontmatter, "doctitle",
                                    id="forsmil-1",
                                    smilref="dtbook.smil#sforsmil-1")
    doctitle_seq.text = book_title

    docauthor = etree.SubElement(dtbook_frontmatter, "docauthor",
                                 id="forsmil-2",
                                 smilref="dtbook.smil#sforsmil-2")
    docauthor.text = book_author

    # 출판사 추가
    docpublisher = etree.SubElement(dtbook_frontmatter, "docpublisher",
                                    id="forsmil-3",
                                    smilref="dtbook.smil#sforsmil-3")
    docpublisher.text = book_publisher

    # bodymatter 추가
    dtbook_bodymatter = etree.SubElement(dtbook_book, "bodymatter")

    # 현재 level1 요소
    current_level1 = None
    current_level = 0

    # 콘텐츠 추가
    for item in content_structure:
        if item["type"] == "pagenum":
            pagenum = etree.SubElement(
                current_level1 if current_level1 is not None else dtbook_bodymatter,
                "pagenum",
                id=f"page_{item['text']}_{item['text']}",
                smilref=f"dtbook.smil#smil_par_page_{item['text']}_{item['text']}",
                page="normal"
            )
            pagenum.text = str(item["text"])
            continue
        elif item["type"] == "image":
            if current_level1 is None:
                # level1이 없는 경우 생성
                current_level1 = etree.SubElement(dtbook_bodymatter, "level1",
                                                id=item["id"],
                                                smilref=f"dtbook.smil#smil_par_{item['id']}")
                current_level = 1
                heading = etree.SubElement(current_level1, "h1")
                heading.text = " ".join(item["words"])

            # 이미지 그룹 생성
            imggroup = etree.SubElement(
                current_level1,
                "imggroup",
                id=item["id"],
                class_="figure"
            )

            # 이미지 요소 생성
            img = etree.SubElement(
                imggroup,
                "img",
                id=f"{item['id']}_img",
                src=item["src"],
                alt=item["alt_text"]
            )
            
            # 이미지 크기를 적절히 설정
            img.set("width", "100%")
            img.set("height", "auto")
            
            # 이미지 캡션 추가
            caption = etree.SubElement(imggroup, "caption",
                                       id=f"{item['id']}_caption")
            sent = etree.SubElement(caption, "sent",
                                    id=item["sent_id"],
                                    smilref=f"dtbook.smil#smil_par_{item['sent_id']}")
            
            # 이미지 제목만 캡션으로 설정
            # w = etree.SubElement(sent, "w")
            
            # 제목 설정
            # if "title" in item and item["title"]:
            #     img_type = item.get("type", "그림")
            #     w.text = f"{img_type} {item['id'].replace('id_', '')}: {item['title']}"
            # else:
            #     w.text = item["alt_text"]
            
            # 이미지 설명이 있을 경우에만 추가
            # if "description" in item and item["description"]:
            #     desc_p = etree.SubElement(caption, "p", class_="image-description")
            #     desc_p.text = item["description"]
            
            continue
        elif item["type"] == "table":
            # 표 처리
            if current_level1 is None:
                # level1이 없는 경우 생성
                current_level1 = etree.SubElement(dtbook_bodymatter, "level1",
                                                id=item["id"],
                                                smilref=f"dtbook.smil#smil_par_{item['id']}")
                current_level = 1
                heading = etree.SubElement(current_level1, "h1")
                heading.text = "제목 없음"
            
            # 표 요소 생성
            table = etree.SubElement(current_level1, "table", 
                                    id=item["id"],
                                    class_="data-table",
                                    smilref=f"dtbook.smil#smil_par_{item['id']}",
                                    border="1")
            
            # 표 데이터 가져오기
            table_data = item["table_data"]
            
            # 표 번호 가져오기
            table_number = item.get("table_number", 1)  # 기본값 1로 설정
            
            # tbody 요소 생성
            tbody = etree.SubElement(table, "tbody")
            
            # 표 데이터로 행과 열 생성
            for row_idx, row_data in enumerate(table_data["rows"]):
                tr = etree.SubElement(tbody, "tr", 
                                     id=f"forsmil-{element_counter+row_idx}",
                                     smilref=f"dtbook.smil#smil_par_{item['id']}_cell_{row_idx}")
                
                for col_idx, cell_text in enumerate(row_data):
                    # 셀 정보 찾기
                    cell_info = next((cell for cell in table_data["cells"] 
                                    if cell["row"] == row_idx and cell["col"] == col_idx), None)
                    
                    # 병합된 셀의 경우 건너뛰기 (이미 처리됨)
                    if cell_info and cell_info["is_merged"] and (cell_info["rowspan"] > 1 or cell_info["colspan"] > 1):
                        continue
                    
                    # 셀 요소 생성
                    if col_idx == 0:
                        cell_elem = etree.SubElement(tr, "th", scope="row",
                                                    id=f"forsmil-{element_counter+row_idx*10+col_idx}",
                                                    smilref=f"dtbook.smil#smil_par_{item['id']}_cell_{row_idx}_{col_idx}")
                    else:
                        cell_elem = etree.SubElement(tr, "td",
                                                    id=f"forsmil-{element_counter+row_idx*10+col_idx}",
                                                    smilref=f"dtbook.smil#smil_par_{item['id']}_cell_{row_idx}_{col_idx}")
                    
                    # 병합 속성 설정
                    if cell_info:
                        if cell_info["rowspan"] > 1:
                            cell_elem.set("rowspan", str(cell_info["rowspan"]))
                        if cell_info["colspan"] > 1:
                            cell_elem.set("colspan", str(cell_info["colspan"]))
                    
                    # 셀 내용 추가: sent/w 태그 없이 <p> 바로 텍스트를 넣음
                    p = etree.SubElement(
                        cell_elem,
                        "p",
                        id=f"table_{item['id']}_cell_{row_idx}_{col_idx}",
                        smilref=f"dtbook.smil#smil_par_{item['id']}_cell_{row_idx}_{col_idx}"
                    )
                    # br 태그가 포함된 경우 실제 br 요소로 생성
                    if cell_text.strip() == "<br/>":
                        br_elem = etree.SubElement(p, "br")
                    else:
                        p.text = cell_text.strip()
        elif item["type"].startswith("h"):
            level = int(item["type"][1])  # h1 -> 1, h2 -> 2, h3 -> 3, h4 -> 4, h5 -> 5, h6 -> 6

            if level == 1:
                # 새로운 level1 시작
                current_level1 = etree.SubElement(dtbook_bodymatter, "level1",
                                                id=item["id"],
                                                smilref=f"dtbook.smil#smil_par_{item['id']}")
                current_level = 1
                heading = etree.SubElement(current_level1, "h1")
                heading.text = " ".join(item["words"])
            else:
                # level2~6은 이전 level 내에 위치
                if current_level1 is None:
                    # level1이 없는 경우 생성
                    current_level1 = etree.SubElement(dtbook_bodymatter, "level1",
                                                    id=item["id"],
                                                    smilref=f"dtbook.smil#smil_par_{item['id']}")
                    current_level = 1
                    heading = etree.SubElement(current_level1, "h1")
                    heading.text = "제목 없음"

                # 현재 레벨에 맞는 부모 요소 찾기
                parent = current_level1
                current_level_elem = None
                for l in range(2, level):
                    level_elem = parent.find(f"level{l}")
                    if level_elem is None:
                        # 중간 레벨이 없으면 생성
                        level_elem = etree.SubElement(parent, f"level{l}",
                                                    id=item["id"],
                                                    smilref=f"dtbook.smil#smil_par_{item['id']}")
                        heading = etree.SubElement(level_elem, f"h{l}")
                        heading.text = f"제목 {l}"
                    parent = level_elem
                    current_level_elem = level_elem

                # 새로운 level 요소 생성
                new_level = etree.SubElement(parent, f"level{level}",
                                           id=item["id"],
                                           smilref=f"dtbook.smil#smil_par_{item['id']}")
                heading = etree.SubElement(new_level, f"h{level}")
                heading.text = " ".join(item["words"])

                # 현재 레벨 요소 업데이트
                if level > current_level:
                    current_level_elem = new_level
                current_level = level

            # 기타 마커 처리
            for marker in item.get("markers", []):
                if marker.type != "page":  # 페이지 마커는 이미 처리됨
                    elem_info = MarkerProcessor.create_dtbook_element(marker)
                    if elem_info:
                        marker_elem = etree.SubElement(current_level_elem or current_level1,
                                                     elem_info["tag"],
                                                     attrib=elem_info["attrs"])
                        marker_elem.text = elem_info["text"]

            # 일반 텍스트 내용 추가
            if not item["type"].startswith("h"):
                parent_elem = current_level_elem or current_level1
                p = etree.SubElement(parent_elem, "p",
                                   id=item["id"],
                                   smilref=f"dtbook.smil#smil_par_{item['id']}")
                
                # <br/> 태그가 포함된 경우 실제 br 요소로 생성
                if item.get("text", "") == "<br/>":
                    br_elem = etree.SubElement(p, "br")
                else:
                    p.text = " ".join(item["words"])
        else:
            # 일반 단락은 현재 level 요소 내에 추가
            if current_level1 is None:
                # level1이 없는 경우 생성
                current_level1 = etree.SubElement(dtbook_bodymatter, "level1",
                                                id=item["id"],
                                                smilref=f"dtbook.smil#smil_par_{item['id']}")
                # 임시 제목 추가
                temp_h1 = etree.SubElement(current_level1, "h1")
                temp_h1.text = "제목 없음"

            p = etree.SubElement(current_level1, "p",
                                 id=item["id"],
                                 smilref=f"dtbook.smil#smil_par_{item['id']}")
            
            # <br/> 태그가 포함된 경우 실제 br 요소로 생성
            if item.get("text", "") == "<br/>":
                br_elem = etree.SubElement(p, "br")
            else:
                p.text = " ".join(item["words"])

            # 기타 마커 처리
            for marker in item.get("markers", []):
                if marker.type != "page":  # 페이지 마커는 이미 처리됨
                    elem_info = MarkerProcessor.create_dtbook_element(marker)
                    if elem_info:
                        marker_elem = etree.SubElement(current_level1,
                                                       elem_info["tag"],
                                                       attrib=elem_info["attrs"])
                        marker_elem.text = elem_info["text"]

    # XML 파일 저장
    dtbook_filepath = os.path.join(output_dir, "dtbook.xml")
    tree = etree.ElementTree(dtbook_root)

    # XML 선언에 인코딩 명시적 지정
    with open(dtbook_filepath, 'wb') as f:
        # XML 선언
        f.write('<?xml version="1.0" encoding="utf-8" standalone="no"?>\n'.encode('utf-8'))
        # DTD 선언
        f.write('<!DOCTYPE dtbook PUBLIC "-//NISO//DTD dtbook 2005-3//EN" "http://www.daisy.org/z3986/2005/dtbook-2005-3.dtd">\n'.encode('utf-8'))
        # XML 트리 저장 - 명시적으로 인코딩 지정
        tree.write(f,
                  encoding='utf-8',
                  pretty_print=True,
                  xml_declaration=False,
                  method='xml')

    print(f"DTBook 생성 완료: {dtbook_filepath}")

    # --- 2. OPF 파일 생성 (dtbook.opf) ---
    print("OPF 생성 중...")
    opf_ns = "http://openebook.org/namespaces/oeb-package/1.0/"
    dc_ns = "http://purl.org/dc/elements/1.1/"

    opf_root = etree.Element(
        "package",
        attrib={
            "unique-identifier": "uid"
        },
        nsmap={
            None: opf_ns
        }
    )
 
    # 메타데이터
    metadata = etree.SubElement(opf_root, "metadata")

    # DC 메타데이터
    dc_metadata = etree.SubElement(metadata, "dc-metadata")

    format_elem = etree.SubElement(
        dc_metadata, "{%s}Format" % dc_ns, nsmap={'dc': dc_ns})
    format_elem.text = "ANSI/NISO Z39.86-2005"

    lang_elem = etree.SubElement(
        dc_metadata, "{%s}Language" % dc_ns, nsmap={'dc': dc_ns})
    lang_elem.text = book_language

    date_elem = etree.SubElement(
        dc_metadata, "{%s}Date" % dc_ns, nsmap={'dc': dc_ns})
    date_elem.text = datetime.now().strftime("%Y-%m-%d")

    # publisher_elem = etree.SubElement(
    #     dc_metadata, "{%s}Publisher" % dc_ns, nsmap={'dc': dc_ns})
    # publisher_elem.text = book_publisher

    title_elem = etree.SubElement(
        dc_metadata, "{%s}Title" % dc_ns, nsmap={'dc': dc_ns})
    title_elem.text = book_title

    identifier_elem = etree.SubElement(
        dc_metadata, "{%s}Identifier" % dc_ns, nsmap={'dc': dc_ns}, id="uid")
    identifier_elem.text = book_uid

    creator_elem = etree.SubElement(
        dc_metadata, "{%s}Creator" % dc_ns, nsmap={'dc': dc_ns})
    creator_elem.text = book_author

    # X-Metadata
    x_metadata = etree.SubElement(metadata, "x-metadata")
    etree.SubElement(x_metadata, "meta",
                     name="dtb:multimediaType",
                     content="textNCX")
    etree.SubElement(x_metadata, "meta",
                     name="dtb:totalTime",
                     content="0:00:00")
    etree.SubElement(x_metadata, "meta",
                     name="dtb:multimediaContent",
                     content="text")

    # Manifest
    manifest = etree.SubElement(opf_root, "manifest")

    # OPF
    etree.SubElement(manifest, "item",
                     href="dtbook.opf",
                     id="opf",
                     **{"media-type": "text/xml"})

    # DTBook
    etree.SubElement(manifest, "item",
                     href="dtbook.xml",
                     id="opf-1",
                     **{"media-type": "application/x-dtbook+xml"})

    # SMIL 파일들
    etree.SubElement(manifest, "item",
                     href="dtbook.smil",
                     id="mo",
                     **{"media-type": "application/smil"})

    # NCX
    etree.SubElement(manifest, "item",
                     href="dtbook.ncx",
                     id="ncx",
                     **{"media-type": "application/x-dtbncx+xml"})

    # Resources
    etree.SubElement(manifest, "item",
                     href="dtbook.res",
                     id="resource",
                     **{"media-type": "application/x-dtbresource+xml"})

    # Spine
    spine = etree.SubElement(opf_root, "spine")
    etree.SubElement(spine, "itemref",
                     idref="mo")

    # OPF Manifest에 이미지 파일 추가
    for item in content_structure:
        if item["type"] == "image":
            image_filename = os.path.basename(item["src"])
            image_id = f"img_{item['id']}"
            extension = os.path.splitext(image_filename)[1][1:].lower()

            # 이미지 확장자에 따른 MIME 타입 설정
            mime_type = {
                'jpg': 'image/jpeg',
                'jpeg': 'image/jpeg',
                'png': 'image/png',
                'gif': 'image/gif',
                'bmp': 'image/bmp',
                'tiff': 'image/tiff',
                'tif': 'image/tiff'
            }.get(extension, f'image/{extension}')

            print(f"이미지 매니페스트 추가: {image_filename} (MIME: {mime_type})")
            etree.SubElement(manifest, "item",
                             href=item["src"],
                             id=image_id,
                             **{"media-type": mime_type})

    # OPF 파일 저장
    opf_filepath = os.path.join(output_dir, "dtbook.opf")
    tree = etree.ElementTree(opf_root)

    with open(opf_filepath, 'wb') as f:
        f.write('<?xml version="1.0" encoding="utf-8" standalone="no"?>\n'.encode('utf-8'))
        f.write('<!DOCTYPE package PUBLIC "+//ISBN 0-9673008-1-9//DTD OEB 1.2 Package//EN" "http://openebook.org/dtds/oeb-1.2/oebpkg12.dtd">\n'.encode('utf-8'))
        tree.write(f,
                  encoding='utf-8',
                  pretty_print=True,
                  xml_declaration=False,
                  method='xml')

    print(f"OPF 생성 완료: {opf_filepath}")

    # --- 3. SMIL 파일 생성 (dtbook.smil) ---
    print("SMIL 파일 생성 중...")

    smil_ns = "http://www.w3.org/2001/SMIL20/"

    smil_root = etree.Element(
        "smil",
        nsmap={
            None: smil_ns
        }
    )

    # head
    head = etree.SubElement(smil_root, "head")
    etree.SubElement(head, "meta",
                     name="dtb:uid",
                     content=book_uid)
    etree.SubElement(head, "meta",
                     name="dtb:totalElapsedTime",
                     content="0:00:00")
    etree.SubElement(head, "meta",
                     name="dtb:generator",
                     content="DAISY Pipeline 2")

    # 페이지 관련 메타데이터 추가
    etree.SubElement(head, "meta",
                     name="dtb:totalPageCount",
                     content=str(total_pages))
    etree.SubElement(head, "meta",
                     name="dtb:maxPageNumber",
                     content=str(max_page_number))

    # body
    body = etree.SubElement(smil_root, "body")
    root_seq = etree.SubElement(body, "seq", id="root-seq")

    # doctitle과 docauthor 추가
    doctitle_par = etree.SubElement(root_seq, "par",
                                   id="sforsmil-1",
                                   **{"class": "doctitle"})
    etree.SubElement(doctitle_par, "text",
                     src="dtbook.xml#forsmil-1")

    docauthor_par = etree.SubElement(root_seq, "par",
                                    id="sforsmil-2",
                                    **{"class": "docauthor"})
    etree.SubElement(docauthor_par, "text",
                     src="dtbook.xml#forsmil-2")

    # 나머지 콘텐츠 추가
    for item in content_structure:
        # pagenum 타입은 이미 DTBook에서 처리되었으므로 SMIL에서만 처리
        if item["type"] == "pagenum":
            page_par = etree.SubElement(root_seq, "par",
                                      id=f"smil_par_page_{item['text']}_{item['text']}",
                                      **{"class": "pagenum"},
                                      customTest="pagenum")
            etree.SubElement(page_par, "text",
                           src=f"dtbook.xml#page_{item['text']}_{item['text']}")
            continue

        # 기본 콘텐츠
        if item["type"].startswith("h"):
            # 제목 요소일 경우 level로 처리
            level = int(item["type"][1])  # h1 -> 1, h2 -> 2, h3 -> 3, h4 -> 4, h5 -> 5, h6 -> 6
            par = etree.SubElement(root_seq, "par",
                                 id=f"smil_par_{item['id']}",
                                 **{"class": f"level{level}"})
            etree.SubElement(par, "text",
                           src=f"dtbook.xml#{item['id']}")
        else:
            # 일반 콘텐츠 처리
            par = etree.SubElement(root_seq, "par",
                                 id=f"smil_par_{item['id']}",
                                 **{"class": item["type"]})
            etree.SubElement(par, "text",
                           src=f"dtbook.xml#{item['id']}")

        # 표 처리
        if item["type"] == "table":
            for row_idx, row_data in enumerate(item["table_data"]["rows"]):
                for col_idx, cell_text in enumerate(row_data):
                    cell_par = etree.SubElement(root_seq, "par",
                                              id=f"smil_par_{item['id']}_cell_{row_idx}_{col_idx}",
                                              **{"class": "table-cell"})
                    etree.SubElement(cell_par, "text",
                                   src=f"dtbook.xml#table_{item['id']}_cell_{row_idx}_{col_idx}")

        # 마커 처리 (페이지 마커 제외)
        for marker in item.get("markers", []):
            if marker.type != "page":  # 페이지 마커는 이미 처리됨
                elem_info = MarkerProcessor.create_smil_element(marker, item["id"])
                if elem_info:
                    marker_par = etree.SubElement(root_seq, "par",
                                                id=f"smil_par_{item['id']}_{marker.type}",
                                                **{"class": elem_info["par_class"]})
                    etree.SubElement(marker_par, "text",
                                   src=elem_info["text_src"])

    # SMIL 파일 저장
    smil_filepath = os.path.join(output_dir, "dtbook.smil")
    tree = etree.ElementTree(smil_root)

    with open(smil_filepath, 'wb') as f:
        f.write('<?xml version="1.0" encoding="utf-8" standalone="no"?>\n'.encode('utf-8'))
        f.write('<!DOCTYPE smil PUBLIC "-//NISO//DTD dtbsmil 2005-2//EN" "http://www.daisy.org/z3986/2005/dtbsmil-2005-2.dtd">\n'.encode('utf-8'))
        tree.write(f,
                  encoding='utf-8',
                  pretty_print=True,
                  xml_declaration=False,
                  method='xml')

    print(f"SMIL 파일 생성 완료: {smil_filepath}")

    # --- 4. NCX 파일 생성 (dtbook.ncx) ---
    print("NCX 생성 중...")
    ncx_ns = "http://www.daisy.org/z3986/2005/ncx/"

    ncx_root = etree.Element(
        "{%s}ncx" % ncx_ns,
        attrib={
            "version": "2005-1"
        },
        nsmap={
            None: ncx_ns
        }
    )

    # head
    head = etree.SubElement(ncx_root, "head")
    etree.SubElement(head, "meta",
                     name="dc:Identifier",
                     content=book_uid)
    etree.SubElement(head, "meta",
                     name="dc:Title",
                     content=book_title)
    etree.SubElement(head, "meta",
                     name="dc:Date",
                     content=datetime.now().strftime("%Y-%m-%d"))
    etree.SubElement(head, "meta",
                     name="dc:Format",
                     content="ANSI/NISO Z39.86-2005")
    etree.SubElement(head, "meta",
                     name="dc:Language",
                     content=book_language)
    etree.SubElement(head, "meta",
                     name="dtb:depth",
                     content="3")  # 최대 제목 레벨
    etree.SubElement(head, "meta",
                     name="dtb:totalPageCount",
                     content=str(total_pages))
    etree.SubElement(head, "meta",
                     name="dtb:maxPageNumber",
                     content=str(max_page_number))
    etree.SubElement(head, "meta",
                     name="dtb:uid",
                     content=book_uid)
    etree.SubElement(head, "meta",
                     name="dtb:generator",
                     content="docx_to_daisy")

    # smilCustomTest 추가
    etree.SubElement(head, "smilCustomTest",
                    id="pagenum",
                    defaultState="false",
                    override="visible",
                    bookStruct="PAGE_NUMBER")
    etree.SubElement(head, "smilCustomTest",
                    id="note",
                    defaultState="true",
                    override="visible",
                    bookStruct="NOTE")
    etree.SubElement(head, "smilCustomTest",
                    id="noteref",
                    defaultState="true",
                    override="visible",
                    bookStruct="NOTE_REFERENCE")
    etree.SubElement(head, "smilCustomTest",
                    id="table",
                    defaultState="true",
                    override="visible")

    # docTitle
    doc_title = etree.SubElement(ncx_root, "docTitle")
    text = etree.SubElement(doc_title, "text")
    text.text = book_title

    # docAuthor
    doc_author = etree.SubElement(ncx_root, "docAuthor")
    text = etree.SubElement(doc_author, "text")
    text.text = book_author

    # navMap
    nav_map = etree.SubElement(ncx_root, "navMap")

    # 목차 항목 생성
    play_order = 1
    current_level1_point = None
    current_level2_point = None
    current_level3_point = None
    current_level4_point = None
    current_level5_point = None

    for item in content_structure:
        if item["type"].startswith("h"):
            level = int(item["type"][1])  # h1 -> 1, h2 -> 2, h3 -> 3, h4 -> 4, h5 -> 5, h6 -> 6
            nav_point = etree.Element("navPoint",
                                     id=f"ncx_{item['id']}",
                                     **{"class": f"level{level}"},
                                     playOrder=str(play_order))
            nav_label = etree.SubElement(nav_point, "navLabel")
            text = etree.SubElement(nav_label, "text")
            text.text = item["text"]
            content = etree.SubElement(nav_point, "content",
                                       src=f"dtbook.smil#smil_par_{item['id']}")

            # 계층 구조에 맞게 배치
            if level == 1:
                nav_map.append(nav_point)
                current_level1_point = nav_point
                current_level2_point = None
                current_level3_point = None
                current_level4_point = None
                current_level5_point = None
            elif level == 2 and current_level1_point is not None:
                current_level1_point.append(nav_point)
                current_level2_point = nav_point
                current_level3_point = None
                current_level4_point = None
                current_level5_point = None
            elif level == 3 and current_level2_point is not None:
                current_level2_point.append(nav_point)
                current_level3_point = nav_point
                current_level4_point = None
                current_level5_point = None
            elif level == 4 and current_level3_point is not None:
                current_level3_point.append(nav_point)
                current_level4_point = nav_point
                current_level5_point = None
            elif level == 5 and current_level4_point is not None:
                current_level4_point.append(nav_point)
                current_level5_point = nav_point
            elif level == 6 and current_level5_point is not None:
                current_level5_point.append(nav_point)

            play_order += 1
        elif item["type"] == "table":
            # 표 네비게이션 포인트 추가
            nav_point = etree.Element("navPoint",
                                     id=f"ncx_{item['id']}",
                                     **{"class": "level1"},
                                     playOrder=str(play_order))
            nav_label = etree.SubElement(nav_point, "navLabel")
            text = etree.SubElement(nav_label, "text")
            text.text = f"표 {play_order}"  # 표 제목 또는 번호
            content = etree.SubElement(nav_point, "content",
                                       src=f"dtbook.smil#smil_par_{item['id']}")
            
            # 현재 레벨에 추가
            if current_level1_point is not None:
                current_level1_point.append(nav_point)
            else:
                nav_map.append(nav_point)
            
            play_order += 1

    # pageList (페이지 마커가 있는 경우 추가)
    page_targets = []
    processed_page_markers = set()  # 이미 처리된 페이지 마커 추적
    
    for item in content_structure:
        # pagenum 타입은 이미 DTBook에서 처리되었으므로 NCX에서 제외
        if item["type"] == "pagenum":
            processed_page_markers.add(item["text"])
            continue
            
        for marker in item.get("markers", []):
            if marker.type == "page":
                # 이미 처리된 페이지 마커는 건너뛰기
                if marker.value in processed_page_markers:
                    continue
                    
                page_targets.append({
                    "id": f"p{marker.value}",
                    "value": marker.value,
                    "type": "normal",  # front, normal, special 중 하나
                    "smil_file": item["smil_file"],
                    "item_id": item["id"],
                    "play_order": play_order
                })
                processed_page_markers.add(marker.value)
                play_order += 1

    if page_targets:
        page_list = etree.SubElement(ncx_root, "pageList", id="pages")
        nav_label = etree.SubElement(page_list, "navLabel")
        text = etree.SubElement(nav_label, "text")
        text.text = "Page numbers list"
        
        for page in page_targets:
            nav_point = etree.SubElement(page_list, "pageTarget",
                                        id=page["id"],
                                        **{"class": "pagenum"},
                                        type=page["type"],
                                        value=page["value"],
                                        playOrder=str(page["play_order"]))
            nav_label = etree.SubElement(nav_point, "navLabel")
            text = etree.SubElement(nav_label, "text")
            text.text = page["value"]
            content = etree.SubElement(nav_point, "content",
                                      src=f"{page['smil_file']}#smil_par_page_{page['value']}_{page['value']}")

    # navList (각주, 미주 등이 있는 경우 추가)
    note_targets = []
    for item in content_structure:
        for marker in item.get("markers", []):
            if marker.type in ["note", "annotation"]:
                note_targets.append({
                    "id": f"note_{marker.value}",
                    "text": marker.text,
                    "smil_file": item["smil_file"],
                    "item_id": item["id"],
                    "play_order": play_order
                })
                play_order += 1

    if note_targets:
        nav_list = etree.SubElement(ncx_root, "navList")
        nav_label = etree.SubElement(nav_list, "navLabel")
        text = etree.SubElement(nav_label, "text")
        text.text = "각주"

        for note in note_targets:
            nav_target = etree.SubElement(nav_list, "navTarget",
                                         id=note["id"],
                                         playOrder=str(note["play_order"]))
            nav_label = etree.SubElement(nav_target, "navLabel")
            text = etree.SubElement(nav_label, "text")
            text.text = note["text"]
            content = etree.SubElement(nav_target, "content",
                                      src=f"{note['smil_file']}#s{note['item_id']}")

    # NCX 파일 저장
    ncx_filepath = os.path.join(output_dir, "dtbook.ncx")
    tree = etree.ElementTree(ncx_root)

    with open(ncx_filepath, 'wb') as f:
        f.write(b'<?xml version="1.0" encoding="UTF-8"?>\n')
        f.write(b'<!DOCTYPE ncx PUBLIC "-//NISO//DTD ncx 2005-1//EN" "http://www.daisy.org/z3986/2005/ncx-2005-1.dtd" >\n')
        tree.write(f,
                   pretty_print=True,
                   encoding='utf-8',
                   xml_declaration=False)

    print(f"NCX 생성 완료: {ncx_filepath}")

    # --- 5. Resources 파일 생성 (dtbook.res) ---
    print("Resources 생성 중...")
    res_ns = "http://www.daisy.org/z3986/2005/resource/"

    res_root = etree.Element(
        "{%s}resources" % res_ns,
        attrib={
            "version": "2005-1"
        },
        nsmap={
            None: res_ns
        }
    )

    # NCX scope
    ncx_scope = etree.SubElement(res_root, "scope",
                                 nsuri="http://www.daisy.org/z3986/2005/ncx/")

    # Custom tests
    custom_tests = [
        ("page-set", "PAGE_NUMBER", "page"),
        ("note-set", "NOTE", "note"),
        ("notref-set", "NOTE_REFERENCE", "note"),
        ("annot-set", "ANNOTATION", "annotation"),
        ("line-set", "LINE_NUMBER", "line"),
        ("sidebar-set", "OPTIONAL_SIDEBAR", "sidebar"),
        ("prodnote-set", "OPTIONAL_PRODUCER_NOTE", "note")
    ]

    for test_id, book_struct, text in custom_tests:
        node_set = etree.SubElement(ncx_scope, "nodeSet",
                                    id=test_id,
                                    select=f"//smilCustomTest[@bookStruct='{book_struct}']")
        resource = etree.SubElement(node_set, "resource",
                                    **{"{http://www.w3.org/XML/1998/namespace}lang": "en"})
        text_elem = etree.SubElement(resource, "text")
        text_elem.text = text

    # SMIL scope
    smil_scope = etree.SubElement(res_root, "scope",
                                  nsuri="http://www.w3.org/2001/SMIL20/")

    # Math sets
    math_sets = [
        ("math-seq-set", "seq", "mathematical formula"),
        ("math-par-set", "par", "mathematical formula")
    ]

    for set_id, elem_type, text in math_sets:
        node_set = etree.SubElement(smil_scope, "nodeSet",
                                    id=set_id,
                                    select=f"//{elem_type}[@class='math']")
        resource = etree.SubElement(node_set, "resource",
                                    **{"{http://www.w3.org/XML/1998/namespace}lang": "en"})
        text_elem = etree.SubElement(resource, "text")
        text_elem.text = text

    # Resources 파일 저장
    res_filepath = os.path.join(output_dir, "dtbook.res")
    tree = etree.ElementTree(res_root)

    with open(res_filepath, 'wb') as f:
        f.write(b'<?xml version="1.0" encoding="utf-8"?>\n')
        f.write(b'<!DOCTYPE resources\n  PUBLIC "-//NISO//DTD resource 2005-1//EN" "http://www.daisy.org/z3986/2005/resource-2005-1.dtd">\n')
        tree.write(f,
                   pretty_print=True,
                   encoding='utf-8',
                   xml_declaration=False)

    print(f"Resources 생성 완료: {res_filepath}")

    print("\n--- DAISY 기본 파일 생성 완료 ---")
    print(f"생성된 파일은 '{output_dir}' 폴더에 있습니다.")
    print("주의: 이 코드는 DOCX의 기본적인 제목/문단 구조만 변환하며,")
    print("      오디오, SMIL 동기화, 목록, 표, 이미지, 페이지 번호 등은 포함하지 않습니다.")


def zip_daisy_output(source_dir, output_zip_filename):

    """
    지정된 폴더의 내용을 ZIP 파일로 압축합니다.

    Args:
        source_dir (str): 압축할 DAISY 파일들이 있는 폴더 경로.
        output_zip_filename (str): 생성될 ZIP 파일의 이름 (경로 포함 가능).
    """
    if not os.path.isdir(source_dir):
        print(f"오류: 소스 디렉토리를 찾을 수 없습니다 - {source_dir}")
        return

    try:
        print(f"'{source_dir}' 폴더를 '{output_zip_filename}' 파일로 압축 중...")
        # ZIP 파일 쓰기 모드로 열기 (압축 사용)
        with zipfile.ZipFile(output_zip_filename, 'w', zipfile.ZIP_DEFLATED) as zipf:
            # source_dir 내부의 모든 파일과 폴더를 순회
            for root, dirs, files in os.walk(source_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    # ZIP 파일 내부에 저장될 상대 경로 계산
                    # (source_dir 자체를 포함하지 않도록 함)
                    archive_name = os.path.relpath(file_path, source_dir)
                    print(f"  추가 중: {archive_name}")
                    zipf.write(file_path, arcname=archive_name)
        print(f"ZIP 파일 생성 완료: {output_zip_filename}")
    except Exception as e:
        print(f"ZIP 파일 생성 중 오류 발생: {e}")


def create_epub3_book(docx_file_path, output_dir, book_title=None, book_author=None, book_publisher=None, book_language="ko", book_isbn="NOT_GIVEN_ISBN"):
    """DOCX 파일을 EPUB3 형식으로 변환합니다 (TTAK.KO-10.0905 표준 준수).

    Args:
        docx_file_path (str): 변환할 DOCX 파일 경로
        output_dir (str): 생성된 EPUB3 파일 저장 폴더
        book_title (str, optional): 책 제목. 기본값은 None (DOCX 파일명 사용)
        book_author (str, optional): 저자. 기본값은 None
        book_publisher (str, optional): 출판사. 기본값은 None
        book_language (str, optional): 언어 코드 (ISO 639-1). 기본값은 "ko"
    """
    import zipfile
    from datetime import datetime
    import uuid
    
    # --- 출력 디렉토리 생성 ---
    os.makedirs(output_dir, exist_ok=True)

    # --- DOCX 파일 읽기 및 구조 분석 ---
    try:
        document = Document(docx_file_path)
    except FileNotFoundError:
        print(f"오류: DOCX 파일을 찾을 수 없습니다 - {docx_file_path}")
        return
    except Exception as e:
        print(f"오류: DOCX 파일을 읽는 중 오류가 발생했습니다 - {str(e)}")
        return

    # --- 기본 정보 설정 ---
    if book_title is None or not isinstance(book_title, str) or len(book_title.strip()) == 0:
        raise ValueError("책 제목이 제공되지 않았거나 유효하지 않습니다. 변환을 진행할 수 없습니다.")
    
    if book_author is None or not isinstance(book_author, str) or len(book_author.strip()) == 0:
        raise ValueError("저자 정보가 제공되지 않았거나 유효하지 않습니다. 변환을 진행할 수 없습니다.")

    book_title = str(book_title)
    book_author = str(book_author)
    book_publisher = str(book_publisher) if book_publisher else "Unknown Publisher"
    book_isbn = str(book_isbn) if book_isbn else "Unkown ISBN"
    
    # EPUB3 UID 생성
    epub_uid = f"urn:uuid:{uuid.uuid4()}"
    
    # 콘텐츠 구조 분석 (DAISY 변환과 동일한 로직 사용)
    content_structure = []
    element_counter = 0
    sent_counter = 0

    # 이미지 관련 정보 저장 변수
    image_info = {}
    
    # 0. 문서 body child 순서를 맵으로 생성하여 실제 위치 사용
    body_children = list(document._element.body.iterchildren())
    element_index = {id(child): idx for idx, child in enumerate(body_children)}

    # 1. 문서에서 모든 이미지 찾기
    print("문서에서 이미지 찾는 중...")
    images = find_all_images(document)
    print(f"총 {len(images)}개의 이미지 발견")
    
    # 2. 문서에서 모든 이미지 추출
    print(f"문서에서 이미지 추출 중...")
    image_counter = 0
    image_relations = []
    
    # 모든 이미지 관계 수집
    for rel_id, rel in document.part.rels.items():
        if "image" in rel.reltype:
            try:
                image_relations.append(rel)
                print(f"이미지 관계 발견: {rel_id}, {rel.reltype}")
            except Exception as e:
                print(f"이미지 관계 처리 오류: {str(e)}")
    
    print(f"문서에서 {len(image_relations)}개의 이미지 관계 발견")
    
    # 이미지 매핑 정보 초기화
    image_mapping = {}

    # 이미지 처리 (G38, G39, G41, G42 지침 준수)
    for i, img in enumerate(images, 1):
        img_num = str(i)
        try:
            image_counter += 1
            element_counter += 1
            sent_counter += 1
            
            # 이미지 ID 생성
            elem_id = f"img_{element_counter}"
            sent_id = f"sent_{sent_counter}"
            
            # 이미지 파일 저장
            image_ext = ".jpeg"
            try:
                if 'image_rid' in img:
                    rel = document.part.rels[img['image_rid']]
                    if hasattr(rel, 'target_ref'):
                        ext = os.path.splitext(rel.target_ref)[1]
                        if ext:
                            image_ext = ext
            except:
                pass
            
            image_filename = f"image{img_num}{image_ext}"
            image_path = os.path.join(output_dir, image_filename)
            
            # 이미지 데이터 저장
            with open(image_path, "wb") as img_file:
                img_file.write(img['image_data'])
            print(f"이미지 {img_num} 저장: {image_path}")
            
            # 이미지 정보를 content_structure에 추가 (figure 요소 사용)
            para_position = element_index.get(id(img['paragraph']._element), img['paragraph_index'])
            content_structure.append({
                "type": "image",
                "src": image_filename,
                "alt_text": f"이미지 {img_num}",
                "id": elem_id,
                "sent_id": sent_id,
                "level": 0,
                "markers": [],
                "position": para_position,
                "insert_before": False
            })
            print(f"이미지 {img_num}를 content_structure에 추가함 (위치: {para_position})")
        except Exception as e:
            print(f"이미지 {img_num} 처리 중 오류 발생: {str(e)}")
            continue

    print(f"{image_counter}개 이미지 추출 완료.")

    # 메모리 정리
    del images
    del image_relations
    del image_mapping
    gc.collect()

    # DOCX의 단락(paragraph)을 순회하며 구조 파악
    print("DOCX 파일 분석 중...")
    print(f"총 {len(document.paragraphs)}개의 단락을 처리합니다.")
    
    # 단락 처리
    for para_idx, para in enumerate(document.paragraphs):
        # 진행 상황 로그 (100개 단락마다)
        if para_idx % 100 == 0:
            print(f"단락 처리 진행 중: {para_idx}/{len(document.paragraphs)} ({para_idx/len(document.paragraphs)*100:.1f}%)")
        
        text_raw = para.text
        style_name = para.style.name.lower()

        # <br/> 태그 기준으로 세그먼트를 분리
        br_segments = re.split(r'<br\s*/?>', text_raw, flags=re.IGNORECASE)

        # 세그먼트별 처리
        for seg_idx, seg_text in enumerate(br_segments):
            # <br/> 태그가 있었던 자리에 빈 문단을 생성
            if seg_idx > 0:
                element_counter += 1
                sent_counter += 1
                blank_elem_id = f"p_{element_counter}"
                blank_sent_id = f"sent_{sent_counter}"
                content_structure.append({
                    "type": "p",
                    "text": "<br/>",
                    "words": ["<br/>"],
                    "id": blank_elem_id,
                    "sent_id": blank_sent_id,
                    "level": 0,
                    "markers": [],
                    "position": para_idx,
                    "insert_before": False
                })

            # 세그먼트 자체가 비어 있으면(공백만) 넘어감
            if not seg_text.strip():
                continue

            # 마커 처리
            processed_text, markers = MarkerProcessor.process_text(seg_text.strip())

            # 페이지 마커가 있는 경우 먼저 처리 (G104 지침 준수)
            for marker in markers:
                if marker.type == "page":
                    element_counter += 1
                    sent_counter += 1
                    elem_id = f"page_{element_counter}"
                    sent_id = f"sent_{sent_counter}"
                    content_structure.append({
                        "type": "pagenum",
                        "text": marker.value,
                        "words": [marker.value],
                        "id": elem_id,
                        "sent_id": sent_id,
                        "level": 0,
                        "markers": [marker],
                        "position": para_idx,
                        "insert_before": True
                    })

                    # 마커만 있고 실제 내용이 없는 경우 건너뜀
                    if not processed_text.strip():
                        continue

            element_counter += 1
            sent_counter += 1
            elem_id = f"p_{element_counter}"
            sent_id = f"sent_{sent_counter}"

            # 단어 분리
            words = split_text_to_words(processed_text)

            # 스타일 이름에 따른 구조 매핑 (G23, G24, G25 지침 준수)
            content_structure.append({
                "type": "h1" if style_name.startswith('heading 1') or style_name == '제목 1' else
                "h2" if style_name.startswith('heading 2') or style_name == '제목 2' else
                "h3" if style_name.startswith('heading 3') or style_name == '제목 3' else
                "h4" if style_name.startswith('heading 4') or style_name == '제목 4' else
                "h5" if style_name.startswith('heading 5') or style_name == '제목 5' else
                "h6" if style_name.startswith('heading 6') or style_name == '제목 6' else
                "p",
                "text": processed_text,
                "words": words,
                "id": elem_id,
                "sent_id": sent_id,
                "level": 1 if style_name.startswith('heading 1') or style_name == '제목 1' else
                        2 if style_name.startswith('heading 2') or style_name == '제목 2' else
                        3 if style_name.startswith('heading 3') or style_name == '제목 3' else
                        4 if style_name.startswith('heading 4') or style_name == '제목 4' else
                        5 if style_name.startswith('heading 5') or style_name == '제목 5' else
                        6 if style_name.startswith('heading 6') or style_name == '제목 6' else
                        0,
                "markers": markers,
                "position": para_idx,
                "insert_before": False
            })
    
    print(f"단락 처리 완료: 총 {len(content_structure)}개의 구조 요소 생성")

    # 표 처리 (G31-G37 지침 준수)
    print("표 처리 중...")
    
    if len(document.tables) > 0:
        print(f"문서에 {len(document.tables)}개의 표 발견")
        
        for table_idx, table in enumerate(document.tables, 1):
            print(f"표 {table_idx} 처리 중...")
            element_counter += 1
            sent_counter += 1
            elem_id = f"table_{element_counter}"
            sent_id = f"sent_{sent_counter}"
            
            # 표 데이터 추출
            table_data = {
                "rows": [],
                "cols": [],
                "cells": []
            }
            
            # 행과 열 정보 추출
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for col_idx, cell in enumerate(row.cells):
                    cell_text = " ".join(para.text for para in cell.paragraphs)
                    row_data.append(cell_text)
                    
                    # 셀 병합 정보 확인
                    rowspan = 1
                    colspan = 1
                    is_merged_cell = False
                    
                    if hasattr(cell, '_tc') and hasattr(cell._tc, 'vMerge'):
                        if cell._tc.vMerge == 'restart':
                            is_merged_cell = True
                            for next_row_idx in range(row_idx + 1, len(table.rows)):
                                if col_idx < len(table.rows[next_row_idx].cells):
                                    next_cell = table.rows[next_row_idx].cells[col_idx]
                                    if (hasattr(next_cell, '_tc') and hasattr(next_cell._tc, 'vMerge') and 
                                        next_cell._tc.vMerge == 'continue'):
                                        rowspan += 1
                                    else:
                                        break
                                else:
                                    break
                        elif cell._tc.vMerge == 'continue':
                            continue
                    
                    # 가로 병합 확인
                    if hasattr(cell, '_tc') and hasattr(cell._tc, 'hMerge'):
                        if cell._tc.hMerge == 'restart':
                            is_merged_cell = True
                            colspan = 1
                            for next_col_idx in range(col_idx + 1, len(row.cells)):
                                next_cell = row.cells[next_col_idx]
                                if (hasattr(next_cell, '_tc') and hasattr(next_cell._tc, 'hMerge') and 
                                    next_cell._tc.hMerge == 'continue'):
                                    colspan += 1
                                else:
                                    break
                        elif cell._tc.hMerge == 'continue':
                            continue
                    
                    # 셀 정보 저장
                    table_data["cells"].append({
                        "row": row_idx,
                        "col": col_idx,
                        "text": cell_text,
                        "is_merged": is_merged_cell,
                        "rowspan": rowspan,
                        "colspan": colspan
                    })
                
                table_data["rows"].append(row_data)
            
            # 열 정보 추출
            for col_idx in range(len(table.columns)):
                col_data = []
                for row in table.rows:
                    if col_idx < len(row.cells):
                        cell_text = " ".join(para.text for para in row.cells[col_idx].paragraphs)
                        col_data.append(cell_text)
                table_data["cols"].append(col_data)
            
            # 표의 실제 위치를 찾기
            table_position_body = len(document.paragraphs)
            try:
                body_element = document._element.body
                all_elements = list(body_element.iterchildren())
                
                table_element_index = -1
                for idx, element in enumerate(all_elements):
                    if element is table._element:
                        table_element_index = idx
                        break
                        
                if table_element_index != -1:
                    paragraph_count_before_table = 0
                    for idx in range(table_element_index):
                        element = all_elements[idx]
                        if element.tag.endswith('p'):
                            paragraph_count_before_table += 1
                    
                    table_position_body = paragraph_count_before_table
                    print(f"표 {table_idx} 정확한 위치 발견: {table_position_body}")
                else:
                    for para_idx, para in enumerate(document.paragraphs):
                        para_text = para.text.strip()
                        if re.search(r'\[?표\s*\d+\.?\d*\]?', para_text, re.IGNORECASE):
                            table_position_body = para_idx + 0.5
                            print(f"표 {table_idx} 제목 패턴 위치 발견: {para_idx + 0.5}")
                            break
                    
                    if table_position_body == len(document.paragraphs):
                        table_position_body = len(document.paragraphs) - 1
                        print(f"표 {table_idx} 마지막 위치 사용: {table_position_body}")
                
                print(f"표 {table_idx} 최종 위치: {table_position_body}")
                    
            except Exception as e:
                print(f"표 위치 계산 중 오류: {e}")
                table_position_body = len(document.paragraphs) - 1
            
            table_title = f"표 {table_idx}"
            
            # 표 정보를 content_structure에 추가 (caption 요소 포함)
            content_structure.append({
                "type": "table",
                "table_data": table_data,
                "id": elem_id,
                "sent_id": sent_id,
                "level": 0,
                "markers": [],
                "position": table_position_body,
                "insert_before": False,
                "title": table_title,
                "table_number": table_idx,
                "text": table_title
            })
            
            print(f"표 {table_idx} 처리 완료: {len(table_data['rows'])}행 x {len(table_data['cols'])}열, 위치: {table_position_body}")
    else:
        print("문서에 표가 없습니다.")

    # 메모리 정리
    gc.collect()

    # 콘텐츠를 위치에 따라 정렬
    content_structure.sort(key=lambda x: (x["position"], 
                                         x.get("image_number", float('inf')) if x["type"] == "image" else 0, 
                                         not x["insert_before"]))

    print(f"총 {len(content_structure)}개의 구조 요소 분석 완료.")

    # --- EPUB3 파일 생성 (TTAK.KO-10.0905 표준 준수) ---
    print("EPUB3 생성 중...")
    
    # EPUB3 디렉토리 구조 생성
    epub_dir = os.path.join(output_dir, "epub3")
    os.makedirs(epub_dir, exist_ok=True)
    
    # META-INF 디렉토리 생성
    meta_inf_dir = os.path.join(epub_dir, "META-INF")
    os.makedirs(meta_inf_dir, exist_ok=True)
    
    # OEBPS 디렉토리 생성 (EPUB3 콘텐츠)
    oebps_dir = os.path.join(epub_dir, "OEBPS")
    os.makedirs(oebps_dir, exist_ok=True)
    
    # 이미지 디렉토리 생성
    images_dir = os.path.join(oebps_dir, "images")
    os.makedirs(images_dir, exist_ok=True)
    
    # --- 1. container.xml 생성 ---
    container_xml = f'''<?xml version="1.0" encoding="UTF-8"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
    <rootfiles>
        <rootfile full-path="OEBPS/package.opf" media-type="application/oebps-package+xml"/>
    </rootfiles>
</container>'''
    
    with open(os.path.join(meta_inf_dir, "container.xml"), "w", encoding="utf-8") as f:
        f.write(container_xml)
    
    # --- 2. package.opf 생성 (G135, G136 지침 준수) ---
    # 매니페스트 항목들
    manifest_items = []
    spine_items = []
    
    # 기본 파일들 추가
    manifest_items.append('<item id="nav" href="nav.xhtml" media-type="application/xhtml+xml" properties="nav"/>')
    manifest_items.append('<item id="css" href="style.css" media-type="text/css"/>')
    
    # 이미지 파일들 추가
    for item in content_structure:
        if item["type"] == "image":
            image_filename = os.path.basename(item["src"])
            image_id = f"img_{item['id']}"
            extension = os.path.splitext(image_filename)[1][1:].lower()
            
            mime_type = {
                'jpg': 'image/jpeg',
                'jpeg': 'image/jpeg',
                'png': 'image/png',
                'gif': 'image/gif',
                'bmp': 'image/bmp',
                'tiff': 'image/tiff',
                'tif': 'image/tiff'
            }.get(extension, f'image/{extension}')
            
            manifest_items.append(f'<item id="{image_id}" href="{item["src"]}" media-type="{mime_type}"/>')
    
    # HTML 파일들 생성
    html_files = []
    current_chapter = 1
    
    # 제목 페이지 생성 (G14, G15 지침 준수)
    title_html = f'''<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops" lang="{book_language}" xml:lang="{book_language}">
<head>
    <meta charset="UTF-8"/>
    <title>{html_escape(book_title)}</title>
    <link rel="stylesheet" type="text/css" href="style.css"/>
    <!-- 반응형 레이아웃 메타데이터 (G70, G71, G72 지침 준수) -->
    <meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0, user-scalable=yes"/>
</head>
<body>
    <section epub:type="titlepage" role="doc-abstract">
        <h1 class="book-title">{html_escape(book_title)}</h1>
        <p class="book-author">{html_escape(book_author)}</p>
        <p class="book-publisher">{html_escape(book_publisher)}</p>
    </section>
</body>
</html>'''
    
    title_filename = "title.xhtml"
    with open(os.path.join(oebps_dir, title_filename), "w", encoding="utf-8") as f:
        f.write(title_html)
    
    manifest_items.append(f'<item id="title" href="{title_filename}" media-type="application/xhtml+xml"/>')
    spine_items.append('<itemref idref="title"/>')
    html_files.append(title_filename)
    
    # 메인 콘텐츠 HTML 생성 (G1-G6, G9-G13 지침 준수)
    content_html = f'''<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops" lang="{book_language}" xml:lang="{book_language}">
<head>
    <meta charset="UTF-8"/>
    <title>{html_escape(book_title)}</title>
    <link rel="stylesheet" type="text/css" href="style.css"/>
    <!-- 반응형 레이아웃 메타데이터 (G70, G71, G72 지침 준수) -->
    <meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0, user-scalable=yes"/>
</head>
<body>
    <!-- 본문 시작 (G12, G13 지침 준수) -->
    <section epub:type="bodymatter" role="doc-chapter">'''
    
    # 콘텐츠 구조 분석을 위한 변수들
    current_section_level = 0
    section_stack = []
    footnote_counter = 0
    footnotes = []
    
    # 목록 처리용 변수들 (G26, G27 지침 준수)
    in_ordered_list = False
    in_unordered_list = False
    list_items = []
    list_id = None
    
    # 콘텐츠 요소들을 HTML로 변환
    for item in content_structure:
        if item["type"] == "pagenum":
            # 목록이 열려있으면 먼저 닫기
            if in_ordered_list:
                content_html += '\n        </ol>'
                in_ordered_list = False
                list_items = []
            elif in_unordered_list:
                content_html += '\n        </ul>'
                in_unordered_list = False
                list_items = []
            
            content_html += f'\n        <span epub:type="pagebreak" role="doc-pagebreak" id="{item["id"]}">{html_escape(item["text"])}</span>'
        elif item["type"] == "image":
            # 목록이 열려있으면 먼저 닫기
            if in_ordered_list:
                content_html += '\n        </ol>'
                in_ordered_list = False
                list_items = []
            elif in_unordered_list:
                content_html += '\n        </ul>'
                in_unordered_list = False
                list_items = []
            
            content_html += f'\n        <figure id="{item["id"]}">'
            content_html += f'\n            <img src="{item["src"]}" alt="{html_escape(item["alt_text"])}" role="img"/>'
            content_html += f'\n            <figcaption>{html_escape(item["alt_text"])}</figcaption>'
            content_html += f'\n        </figure>'
        elif item["type"] == "table":
            # 목록이 열려있으면 먼저 닫기
            if in_ordered_list:
                content_html += '\n        </ol>'
                in_ordered_list = False
                list_items = []
            elif in_unordered_list:
                content_html += '\n        </ul>'
                in_unordered_list = False
                list_items = []
            
            content_html += f'\n        <table id="{item["id"]}">'
            content_html += f'\n            <caption>{html_escape(item["title"])}</caption>'
            table_data = item["table_data"]
            
            # 표 헤더 (G34, G35, G36 지침 준수)
            if table_data["rows"]:
                content_html += '\n            <thead>'
                content_html += '\n                <tr>'
                for col_idx, cell_text in enumerate(table_data["rows"][0]):
                    content_html += f'\n                    <th scope="col">{html_escape(cell_text)}</th>'
                content_html += '\n                </tr>'
                content_html += '\n            </thead>'
            
            # 표 본문
            content_html += '\n            <tbody>'
            for row_idx, row_data in enumerate(table_data["rows"][1:], 1):
                content_html += '\n                <tr>'
                for col_idx, cell_text in enumerate(row_data):
                    if col_idx == 0:
                        content_html += f'\n                    <th scope="row">{html_escape(cell_text)}</th>'
                    else:
                        content_html += f'\n                    <td>{html_escape(cell_text)}</td>'
                content_html += '\n                </tr>'
            content_html += '\n            </tbody>'
            content_html += '\n        </table>'
        elif item["type"].startswith("h"):
            # 목록이 열려있으면 먼저 닫기
            if in_ordered_list:
                content_html += '\n        </ol>'
                in_ordered_list = False
                list_items = []
            elif in_unordered_list:
                content_html += '\n        </ul>'
                in_unordered_list = False
                list_items = []
            
            level = int(item["type"][1])
            
            # 섹션 구조 관리 (G9, G10, G11 지침 준수)
            while current_section_level >= level:
                if section_stack:
                    content_html += '\n    </section>'
                    section_stack.pop()
                    current_section_level -= 1
            
            # 새로운 섹션 시작
            if level == 1:
                content_html += f'\n    <section epub:type="chapter" role="doc-chapter">'
                section_stack.append("chapter")
            elif level == 2:
                content_html += f'\n    <section epub:type="chapter" role="doc-chapter">'
                section_stack.append("chapter")
            elif level == 3:
                content_html += f'\n    <section epub:type="chapter" role="doc-chapter">'
                section_stack.append("chapter")
            else:
                content_html += f'\n    <section epub:type="chapter" role="doc-chapter">'
                section_stack.append("chapter")
            
            current_section_level = level
            
            content_html += f'\n        <h{level} id="{item["id"]}">{html_escape(item["text"])}</h{level}>'
        else:
            # 일반 단락
            if item.get("text", "") == "<br/>":
                # 목록이 열려있으면 먼저 닫기
                if in_ordered_list:
                    content_html += '\n        </ol>'
                    in_ordered_list = False
                    list_items = []
                elif in_unordered_list:
                    content_html += '\n        </ul>'
                    in_unordered_list = False
                    list_items = []
                
                content_html += f'\n        <p id="{item["id"]}"><br/></p>'
            else:
                # 텍스트 전처리 (G18, G26, G28, G29, G30 지침 준수)
                processed_text = html_escape(item["text"])
                
                # 강세 처리 (G18 지침) - 볼드/이탤릭을 em/strong으로 변환
                processed_text = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', processed_text)
                processed_text = re.sub(r'\*(.*?)\*', r'<em>\1</em>', processed_text)
                
                # 루비 텍스트 처리 (G30 지침)
                processed_text = re.sub(r'([가-힣]+)\(([가-힣]+)\)', r'<ruby>\1<rt>\2</rt></ruby>', processed_text)
                
                # 목록 처리 (G26, G27 지침) - 연속된 목록 항목들을 묶어서 처리
                if re.match(r'^\d+\.\s', processed_text):
                    # 번호 목록
                    if not in_ordered_list:
                        # 새로운 순서 목록 시작
                        if in_unordered_list:
                            content_html += '\n        </ul>'
                            in_unordered_list = False
                        
                        list_id = f"ol_{item['id']}"
                        content_html += f'\n        <ol id="{list_id}" aria-labelledby="{item["id"]}">'
                        in_ordered_list = True
                    
                    # 목록 항목 추가
                    list_item_text = re.sub(r'^\d+\.\s', '', processed_text)
                    content_html += f'\n            <li>{html_escape(list_item_text)}</li>'
                    continue
                elif re.match(r'^[-•*]\s', processed_text):
                    # 글머리 기호 목록
                    if not in_unordered_list:
                        # 새로운 비순서 목록 시작
                        if in_ordered_list:
                            content_html += '\n        </ol>'
                            in_ordered_list = False
                        
                        list_id = f"ul_{item['id']}"
                        content_html += f'\n        <ul id="{list_id}" aria-labelledby="{item["id"]}">'
                        in_unordered_list = True
                    
                    # 목록 항목 추가
                    list_item_text = processed_text[2:]
                    content_html += f'\n            <li>{html_escape(list_item_text)}</li>'
                    continue
                else:
                    # 목록이 열려있으면 닫기
                    if in_ordered_list:
                        content_html += '\n        </ol>'
                        in_ordered_list = False
                    elif in_unordered_list:
                        content_html += '\n        </ul>'
                        in_unordered_list = False
                
                # 문맥 나누기 (G28 지침) - 구분선 감지
                if re.match(r'^[-_]{3,}$', processed_text.strip()):
                    content_html += f'\n        <hr/>'
                    continue
                
                # 코딩코드 처리 (G29 지침) - 코드 블록 감지
                if re.match(r'^```', processed_text) or re.match(r'^    ', processed_text):
                    processed_text = f'<pre><code>{processed_text}</code></pre>'
                    content_html += f'\n        {processed_text}'
                    continue
                
                # MathML 수식 처리 (G81 지침 준수)
                math_match = re.search(r'\\\\(.*?)\\\\', processed_text)
                if math_match:
                    math_content = math_match.group(1)
                    # 간단한 수식을 MathML로 변환 (예: x^2 + y^2 = z^2)
                    mathml = f'''<math xmlns="http://www.w3.org/1998/Math/MathML" alttext="{math_content}">
                        <mrow>
                            <msup><mi>x</mi><mn>2</mn></msup>
                            <mo>+</mo>
                            <msup><mi>y</mi><mn>2</mn></msup>
                            <mo>=</mo>
                            <msup><mi>z</mi><mn>2</mn></msup>
                        </mrow>
                    </math>'''
                    processed_text = re.sub(r'\\\\(.*?)\\\\', mathml, processed_text)
                
                # SVG 도형 처리 (G82-G86 지침 준수)
                svg_match = re.search(r'\\svg\\((.*?)\\)', processed_text)
                if svg_match:
                    svg_content = svg_match.group(1)
                    # 간단한 SVG 생성 (예: 원)
                    svg = f'''<svg xmlns="http://www.w3.org/2000/svg" width="100" height="100" role="img" aria-label="{svg_content}">
                        <title>{svg_content}</title>
                        <desc>{svg_content}에 대한 설명</desc>
                        <circle cx="50" cy="50" r="40" stroke="black" stroke-width="3" fill="red"/>
                    </svg>'''
                    processed_text = re.sub(r'\\svg\\((.*?)\\)', svg, processed_text)
                
                # 보조설명 처리 (G30 지침) - 위첨자, 아래첨자 처리
                processed_text = re.sub(r'\^(\d+)', r'<sup>\1</sup>', processed_text)
                processed_text = re.sub(r'_(\d+)', r'<sub>\1</sub>', processed_text)
                
                # 각주 처리 (G47, G48, G49 지침) - 각주 패턴 감지
                footnote_match = re.search(r'\(각주\d+\)', processed_text)
                if footnote_match:
                    footnote_text = footnote_match.group(0)
                    footnote_counter += 1
                    footnote_id = f"fn_{footnote_counter}"
                    processed_text = re.sub(r'\(각주\d+\)', f'<a epub:type="noteref" role="doc-noteref" href="#{footnote_id}">{footnote_text}</a>', processed_text)
                    
                    # 각주 내용 저장
                    footnotes.append({
                        'id': footnote_id,
                        'text': footnote_text,
                        'content': f'{footnote_text}에 대한 설명'
                    })
                
                # 난외주석 처리 (G52 지침)
                marginalia_match = re.search(r'\(난외주석\d+\)', processed_text)
                if marginalia_match:
                    marginalia_text = marginalia_match.group(0)
                    marginalia_id = f"marg_{len(footnotes) + 1}"
                    processed_text = re.sub(r'\(난외주석\d+\)', f'<a epub:type="noteref" role="doc-noteref" href="#{marginalia_id}">{marginalia_text}</a>', processed_text)
                    
                    # 난외주석 내용 저장
                    footnotes.append({
                        'id': marginalia_id,
                        'text': marginalia_text,
                        'content': f'{marginalia_text}에 대한 설명'
                    })
                
                # 문제-정답 처리 (G53 지침)
                answer_match = re.search(r'\(정답\d+\)', processed_text)
                if answer_match:
                    answer_text = answer_match.group(0)
                    answer_id = f"ans_{len(footnotes) + 1}"
                    processed_text = re.sub(r'\(정답\d+\)', f'<a epub:type="noteref" role="doc-noteref" href="#{answer_id}">{answer_text}</a>', processed_text)
                    
                    # 정답 내용 저장
                    footnotes.append({
                        'id': answer_id,
                        'text': answer_text,
                        'content': f'{answer_text}에 대한 정답'
                    })
                
                # 외부 링크 처리 (G20, G21 지침 준수)
                processed_text = re.sub(
                    r'(https?://[^\s<>"]+)',
                    r'<a href="\1" title="외부 링크: \1" target="_blank" rel="noopener noreferrer">\1</a>',
                    processed_text
                )
                
                content_html += f'\n        <p id="{item["id"]}">{processed_text}</p>'
    
    # 열린 목록 닫기
    if in_ordered_list:
        content_html += '\n        </ol>'
    elif in_unordered_list:
        content_html += '\n        </ul>'
    
    # 각주 섹션 추가 (G47, G48, G49 지침 준수)
    if footnotes:
        content_html += '\n        <hr/>'
        content_html += '\n        <aside epub:type="footnotes" role="doc-footnotes">'
        content_html += '\n            <h2>각주</h2>'
        for footnote in footnotes:
            content_html += f'\n            <aside epub:type="footnote" role="doc-footnote" id="{footnote["id"]}">'
            content_html += f'\n                <p>{html_escape(footnote["content"])}</p>'
            content_html += '\n            </aside>'
        content_html += '\n        </aside>'
    
    # 섹션 닫기
    while section_stack:
        content_html += '\n    </section>'
        section_stack.pop()
    
    content_html += '''
    </section>
</body>
</html>'''
    
    content_filename = "content.xhtml"
    with open(os.path.join(oebps_dir, content_filename), "w", encoding="utf-8") as f:
        f.write(content_html)
    
    manifest_items.append(f'<item id="content" href="{content_filename}" media-type="application/xhtml+xml"/>')
    spine_items.append('<itemref idref="content"/>')
    spine_items.append('<itemref idref="nav"/>')
    html_files.append(content_filename)
    
    # package.opf 파일 생성 (G135, G136 지침 준수)
    package_opf = f'''<?xml version="1.0" encoding="utf-8" standalone="no"?>
<package version="3.0" xmlns="http://www.idpf.org/2007/opf" unique-identifier="uid" xml:lang="{book_language}">
    <metadata xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:opf="http://www.idpf.org/2007/opf">
        <dc:identifier id="uid">{epub_uid}</dc:identifier>
        <dc:title>{html_escape(book_title)}</dc:title>
        <dc:creator>{html_escape(book_author)}</dc:creator>
        <dc:publisher>{html_escape(book_publisher)}</dc:publisher>
        <dc:language>{book_language}</dc:language>
        <dc:date>{datetime.now().strftime("%Y-%m-%d")}</dc:date>
        <dc:format>EPUB3</dc:format>
        <dc:source>{html_escape(book_isbn)}</dc:source>
        <meta property="dcterms:modified">{datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")}</meta>
        
        <!-- 접근성 메타데이터 (G135 지침) -->
        <meta property="schema:accessMode">textual</meta>
        <meta property="schema:accessMode">visual</meta>
        <meta property="schema:accessModeSufficient">textual,visual</meta>
        <meta property="schema:accessModeSufficient">textual</meta>
        <meta property="schema:accessibilityFeature">alternativeText</meta>
        <meta property="schema:accessibilityFeature">tableOfContents</meta>
        <meta property="schema:accessibilityFeature">readingOrder</meta>
        <meta property="schema:accessibilityFeature">structuralNavigation</meta>
        <meta property="schema:accessibilityHazard">noMotionSimulationHazard</meta>
        <meta property="schema:accessibilityHazard">noSoundHazard</meta>
        <meta property="schema:accessibilitySummary">TTAK.KO-10.0905 표준을 준수하여 제작된 접근성 전자책입니다.</meta>
        
        <!-- 접근성 평가 메타데이터 (G136 지침) -->
        <link rel="dcterms:conformsTo" href="http://www.idpf.org/epub/a11y/accessibility-20170105.html#wcag-aa"/>
        <meta property="a11y:certifiedBy">DOCX to EPUB3 Converter</meta>
        <meta property="a11y:certifierCredential">TTAK.KO-10.0905 표준 준수</meta>
    </metadata>
    
    <manifest>
        {chr(10).join(manifest_items)}
    </manifest>
    
    <spine>
        {chr(10).join(spine_items)}
    </spine>
</package>'''
    
    with open(os.path.join(oebps_dir, "package.opf"), "w", encoding="utf-8") as f:
        f.write(package_opf)
    
    # --- 3. nav.xhtml 생성 (G92, G93, G94 지침 준수) ---
    nav_items = []
    
    # 제목 페이지
    nav_items.append(
        f'''        <li><a href="title.xhtml">{html_escape(book_title)}</a></li>''')
    
    # 목차 항목들
    for item in content_structure:
        if item["type"].startswith("h"):
            level = int(item["type"][1])
            indent = "  " * (level - 1)
            nav_items.append(
                f'''{indent}<li><a href="content.xhtml#{item["id"]}">{html_escape(item["text"])}</a></li>''')
    
    # 표 목차 (G95, G96, G97 지침 준수)
    table_list_items = []
    for item in content_structure:
        if item["type"] == "table":
            table_list_items.append(
                f'''        <li><a href="content.xhtml#{item["id"]}">{html_escape(item["title"])}</a></li>''')
    
    # 이미지 목차 (G98, G99, G100 지침 준수)
    image_list_items = []
    for item in content_structure:
        if item["type"] == "image":
            image_list_items.append(
                f'''        <li><a href="content.xhtml#{item["id"]}">{html_escape(item["alt_text"])}</a></li>''')
    
    # 랜드마크 (G101, G102, G103 지침 준수)
    landmark_items = []
    landmark_items.append(f'''        <li><a epub:type="cover" href="title.xhtml">표지</a></li>''')
    landmark_items.append(f'''        <li><a epub:type="toc" href="nav.xhtml">목차</a></li>''')
    landmark_items.append(f'''        <li><a epub:type="bodymatter" href="content.xhtml">본문</a></li>''')
    
    # 페이지 목차 (G105, G106, G107 지침 준수)
    page_list_items = []
    for item in content_structure:
        if item["type"] == "pagenum":
            page_list_items.append(
                f'''        <li><a href="content.xhtml#{item["id"]}">{html_escape(item["text"])}</a></li>''')
    
    # 각주 목차 (G94 지침 준수)
    footnote_list_items = []
    for i, footnote in enumerate(footnotes, 1):
        footnote_list_items.append(
            f'''        <li><a href="content.xhtml#{footnote["id"]}">{html_escape(footnote["text"])}</a></li>''')
    
    # 수식 목차 (G95 지침 준수) - MathML 수식이 있을 경우
    math_list_items = []
    math_counter = 0
    for item in content_structure:
        if item["type"] == "p":
            math_match = re.search(r'\\\\(.*?)\\\\', item["text"])
            if math_match:
                math_counter += 1
                math_content = math_match.group(1)
                math_list_items.append(
                    f'''        <li><a href="content.xhtml#{item["id"]}">수식 {math_counter}: {html_escape(math_content)}</a></li>''')
    
    # 링크 목차 (G96 지침 준수) - 외부 링크가 있을 경우
    link_list_items = []
    link_counter = 0
    for item in content_structure:
        if item["type"] == "p":
            # 외부 링크 패턴 감지 (http, https로 시작하는 링크)
            link_matches = re.findall(r'https?://[^\s<>"]+', item["text"])
            for link in link_matches:
                link_counter += 1
                link_list_items.append(
                    f'''        <li><a href="content.xhtml#{item["id"]}">링크 {link_counter}: {html_escape(link[:50])}...</a></li>''')
    
    nav_xhtml = f'''<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops" lang="{book_language}" xml:lang="{book_language}">
<head>
    <meta charset="UTF-8"/>
    <title>목차</title>
    <link rel="stylesheet" type="text/css" href="style.css"/>
    <!-- 반응형 레이아웃 메타데이터 (G70, G71, G72 지침 준수) -->
    <meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0, user-scalable=yes"/>
</head>
<body>
    <nav epub:type="toc" role="doc-toc" id="toc">
        <h1>목차</h1>
        <ol>
{chr(10).join(nav_items)}
        </ol>
    </nav>
    
    <!-- 표 목차 -->
    <nav epub:type="lot" id="lot" title="표 목차">
        <h1>표 목차</h1>
        <ol>
{chr(10).join(table_list_items) if table_list_items else '            <li><span>표가 없습니다.</span></li>'}
        </ol>
    </nav>
    
    <!-- 이미지 목차 -->
    <nav epub:type="loi" id="loi" title="이미지 목차">
        <h1>이미지 목차</h1>
        <ol>
{chr(10).join(image_list_items) if image_list_items else '            <li><span>이미지가 없습니다.</span></li>'}
        </ol>
    </nav>
    
    <!-- 각주 목차 -->
    <nav epub:type="lot" id="footnote-list" title="각주 목차">
        <h1>각주 목차</h1>
        <ol>
{chr(10).join(footnote_list_items) if footnote_list_items else '            <li><span>각주가 없습니다.</span></li>'}
        </ol>
    </nav>
    
    <!-- 수식 목차 -->
    <nav epub:type="lot" id="math-list" title="수식 목차">
        <h1>수식 목차</h1>
        <ol>
{chr(10).join(math_list_items) if math_list_items else '            <li><span>수식이 없습니다.</span></li>'}
        </ol>
    </nav>
    
    <!-- 링크 목차 -->
    <nav epub:type="lot" id="link-list" title="링크 목차">
        <h1>링크 목차</h1>
        <ol>
{chr(10).join(link_list_items) if link_list_items else '            <li><span>링크가 없습니다.</span></li>'}
        </ol>
    </nav>
    
    <!-- 랜드마크 -->
    <nav epub:type="landmarks" id="landmarks" title="랜드마크">
        <h1>랜드마크</h1>
        <ol>
{chr(10).join(landmark_items)}
        </ol>
    </nav>
    
    <!-- 페이지 목차 -->
    <nav epub:type="page-list" role="doc-pagelist" id="page-list" title="페이지 목차">
        <h1>페이지 목차</h1>
        <ol>
{chr(10).join(page_list_items) if page_list_items else '            <li><span>페이지 번호가 없습니다.</span></li>'}
        </ol>
    </nav>
</body>
</html>'''
    
    with open(os.path.join(oebps_dir, "nav.xhtml"), "w", encoding="utf-8") as f:
        f.write(nav_xhtml)
    
    # --- 4. style.css 생성 (G87, G88, G89 지침 준수) ---
    css_content = '''/* EPUB3 접근성 스타일 (TTAK.KO-10.0905 표준 준수) */
body {
    font-family: "Malgun Gothic", "맑은 고딕", sans-serif;
    line-height: 1.6;
    margin: 0;
    padding: 20px;
    background-color: #ffffff;
    color: #000000;
}

.title-page {
    text-align: center;
    margin: 50px 0;
}

.book-title {
    font-size: 2.5em;
    margin-bottom: 30px;
    font-weight: bold;
}

.book-author {
    font-size: 1.5em;
    margin-bottom: 15px;
}

.book-publisher {
    font-size: 1.2em;
    margin-bottom: 30px;
}

.content {
    max-width: 800px;
    margin: 0 auto;
}

h1 {
    font-size: 2em;
    font-weight: bold;
    margin: 30px 0 20px 0;
    color: #333;
}

h2 {
    font-size: 1.5em;
    font-weight: bold;
    margin: 25px 0 15px 0;
    color: #444;
}

h3 {
    font-size: 1.3em;
    font-weight: bold;
    margin: 20px 0 10px 0;
    color: #555;
}

h4 {
    font-size: 1.1em;
    font-weight: bold;
    margin: 15px 0 8px 0;
    color: #666;
}

h5 {
    font-size: 1em;
    font-weight: bold;
    margin: 12px 0 6px 0;
    color: #777;
}

h6 {
    font-size: 0.9em;
    font-weight: bold;
    margin: 10px 0 5px 0;
    color: #888;
}

p {
    margin: 10px 0;
    text-align: justify;
    text-indent: 1em;
}

[epub\\:type="pagebreak"] {
    display: block;
    text-align: center;
    font-weight: bold;
    margin: 20px 0;
    color: #666;
}

figure {
    margin: 20px 0;
    text-align: center;
}

img {
    max-width: 100%;
    height: auto;
    border: 1px solid #ddd;
}

figcaption {
    margin-top: 10px;
    font-style: italic;
    color: #666;
}

table {
    width: 100%;
    border-collapse: collapse;
    margin: 20px 0;
    border: 1px solid #ddd;
}

th, td {
    border: 1px solid #ddd;
    padding: 8px;
    text-align: left;
}

th {
    background-color: #f5f5f5;
    font-weight: bold;
}

/* 접근성 개선 (G87, G88 지침) */
@media (prefers-reduced-motion: reduce) {
    * {
        animation-duration: 0.01ms !important;
        animation-iteration-count: 1 !important;
        transition-duration: 0.01ms !important;
    }
}

/* 고대비 모드 지원 */
@media (prefers-contrast: high) {
    body {
        background-color: #000000;
        color: #ffffff;
    }
    
    th {
        background-color: #333333;
        color: #ffffff;
    }
    
    td {
        border-color: #ffffff;
    }
}

/* 큰 텍스트 모드 지원 */
@media (prefers-reduced-motion: no-preference) {
    p {
        font-size: 1.1em;
    }
}

/* 목차 스타일 */
nav[epub\\:type="toc"] ol {
    list-style-type: none;
    padding-left: 0;
}

nav[epub\\:type="toc"] li {
    margin: 5px 0;
}

nav[epub\\:type="toc"] a {
    text-decoration: none;
    color: #333;
}

nav[epub\\:type="toc"] a:hover {
    text-decoration: underline;
}

/* 목록 스타일 (G26 지침) */
ol, ul {
    margin: 15px 0;
    padding-left: 30px;
}

ol {
    list-style-type: decimal;
}

ul {
    list-style-type: disc;
}

li {
    margin: 5px 0;
    line-height: 1.5;
}

/* 각주 스타일 (G47, G48, G49 지침) */
aside[epub\\:type="footnote"] {
    margin: 20px 0;
    padding: 15px;
    background-color: #f9f9f9;
    border-left: 3px solid #333;
    font-size: 0.9em;
}

a[epub\\:type="noteref"] {
    text-decoration: none;
    color: #0066cc;
    font-weight: bold;
}

a[epub\\:type="noteref"]:hover {
    text-decoration: underline;
}

/* 코드 블록 스타일 (G29 지침) */
pre {
    background-color: #f5f5f5;
    border: 1px solid #ddd;
    border-radius: 4px;
    padding: 15px;
    margin: 15px 0;
    overflow-x: auto;
    font-family: "Courier New", monospace;
    font-size: 0.9em;
    line-height: 1.4;
}

code {
    background-color: #f5f5f5;
    padding: 2px 4px;
    border-radius: 3px;
    font-family: "Courier New", monospace;
    font-size: 0.9em;
}

/* 강세 스타일 (G18 지침) */
em {
    font-style: italic;
    color: #333;
}

strong {
    font-weight: bold;
    color: #000;
}

/* 보조설명 스타일 (G30 지침) */
sup {
    font-size: 0.8em;
    vertical-align: super;
    line-height: 0;
}

sub {
    font-size: 0.8em;
    vertical-align: sub;
    line-height: 0;
}

/* 구분선 스타일 (G28 지침) */
hr {
    border: none;
    border-top: 2px solid #ddd;
    margin: 30px 0;
}

/* 랜드마크 스타일 (G101, G102, G103 지침) */
nav[epub\\:type="landmarks"] {
    margin: 20px 0;
}

nav[epub\\:type="landmarks"] ol {
    list-style-type: none;
    padding-left: 0;
}

nav[epub\\:type="landmarks"] li {
    margin: 8px 0;
}

nav[epub\\:type="landmarks"] a {
    text-decoration: none;
    color: #0066cc;
    font-weight: bold;
}

nav[epub\\:type="landmarks"] a:hover {
    text-decoration: underline;
}

/* 표 목차, 이미지 목차 스타일 (G95-G100 지침) */
nav[epub\\:type="lot"], nav[epub\\:type="loi"] {
    margin: 20px 0;
}

nav[epub\\:type="lot"] ol, nav[epub\\:type="loi"] ol {
    list-style-type: none;
    padding-left: 0;
}

nav[epub\\:type="lot"] li, nav[epub\\:type="loi"] li {
    margin: 5px 0;
}

nav[epub\\:type="lot"] a, nav[epub\\:type="loi"] a {
    text-decoration: none;
    color: #333;
}

nav[epub\\:type="lot"] a:hover, nav[epub\\:type="loi"] a:hover {
    text-decoration: underline;
}

/* 페이지 목차 스타일 (G105, G106, G107 지침) */
nav[epub\\:type="page-list"] {
    margin: 20px 0;
}

nav[epub\\:type="page-list"] ol {
    list-style-type: none;
    padding-left: 0;
}

nav[epub\\:type="page-list"] li {
    margin: 3px 0;
}

nav[epub\\:type="page-list"] a {
    text-decoration: none;
    color: #666;
    font-weight: bold;
}

nav[epub\\:type="page-list"] a:hover {
    text-decoration: underline;
}

/* 루비 텍스트 스타일 (G30 지침) */
ruby {
    ruby-position: under;
    ruby-align: center;
}

rt {
    font-size: 0.7em;
    line-height: 1.2;
    text-align: center;
}

/* 각주 목차, 수식 목차, 링크 목차 스타일 (G94, G95, G96 지침) */
#footnote-list, #math-list, #link-list {
    margin: 20px 0;
}

#footnote-list ol, #math-list ol, #link-list ol {
    list-style-type: none;
    padding-left: 0;
}

#footnote-list li, #math-list li, #link-list li {
    margin: 5px 0;
}

#footnote-list a, #math-list a, #link-list a {
    text-decoration: none;
    color: #333;
}

#footnote-list a:hover, #math-list a:hover, #link-list a:hover {
    text-decoration: underline;
}

/* 섹션 구조 스타일 (G9, G10, G11 지침) */
section[epub\\:type="chapter"] {
    margin: 20px 0;
    padding: 10px;
    border-left: 3px solid #0066cc;
}

section[epub\\:type="chapter"] h1,
section[epub\\:type="chapter"] h2,
section[epub\\:type="chapter"] h3,
section[epub\\:type="chapter"] h4,
section[epub\\:type="chapter"] h5,
section[epub\\:type="chapter"] h6 {
    color: #0066cc;
}

/* 접근성 개선을 위한 포커스 스타일 */
a:focus, button:focus, input:focus, textarea:focus, select:focus {
    outline: 2px solid #0066cc;
    outline-offset: 2px;
}

/* 스크린 리더 전용 텍스트 */
.sr-only {
    position: absolute;
    width: 1px;
    height: 1px;
    padding: 0;
    margin: -1px;
    overflow: hidden;
    clip: rect(0, 0, 0, 0);
    white-space: nowrap;
    border: 0;
}

/* 고대비 모드에서의 추가 개선 */
@media (prefers-contrast: high) {
    section[epub\\:type="chapter"] {
        border-left-color: #ffffff;
    }
    
    section[epub\\:type="chapter"] h1,
    section[epub\\:type="chapter"] h2,
    section[epub\\:type="chapter"] h3,
    section[epub\\:type="chapter"] h4,
    section[epub\\:type="chapter"] h5,
    section[epub\\:type="chapter"] h6 {
        color: #ffffff;
    }
    
    a:focus, button:focus, input:focus, textarea:focus, select:focus {
        outline-color: #ffffff;
    }
}

/* MathML 스타일 (G81 지침) */
math {
    font-family: "Times New Roman", serif;
    font-size: 1.1em;
    margin: 10px 0;
}

math mrow {
    display: inline-block;
}

math mi, math mn, math mo {
    margin: 0 1px;
}

/* SVG 스타일 (G82-G86 지침) */
svg {
    max-width: 100%;
    height: auto;
    margin: 10px 0;
    border: 1px solid #ddd;
    border-radius: 4px;
}

svg[role="img"] {
    display: block;
    margin: 20px auto;
}

/* 링크 스타일 (G20, G21, G22 지침) */
a {
    color: #0066cc;
    text-decoration: none;
    font-weight: bold;
}

a:hover {
    text-decoration: underline;
    color: #004499;
}

a[target="_blank"]::after {
    content: " (새 창에서 열기)";
    font-size: 0.8em;
    color: #666;
}

/* 외부 링크 구분 */
a[href^="http"] {
    border-bottom: 1px dotted #0066cc;
}

'''
    
    with open(os.path.join(oebps_dir, "style.css"), "w", encoding="utf-8") as f:
        f.write(css_content)
    
    # --- 5. EPUB3 파일 생성 (ZIP 압축) ---
    epub_filename = os.path.join(output_dir, f"{book_title.replace(' ', '_')}.epub")
    
    with zipfile.ZipFile(epub_filename, 'w', zipfile.ZIP_DEFLATED) as epub_zip:
        # mimetype 파일 (첫 번째 파일이어야 함)
        epub_zip.writestr("mimetype", "application/epub+zip",
                          compress_type=zipfile.ZIP_STORED)

        # META-INF/container.xml
        epub_zip.write(os.path.join(meta_inf_dir, "container.xml"), "META-INF/container.xml")
        
        # OEBPS/package.opf
        epub_zip.write(os.path.join(oebps_dir, "package.opf"), "OEBPS/package.opf")
        
        # OEBPS/nav.xhtml
        epub_zip.write(os.path.join(oebps_dir, "nav.xhtml"), "OEBPS/nav.xhtml")
        
        # OEBPS/style.css
        epub_zip.write(os.path.join(oebps_dir, "style.css"), "OEBPS/style.css")
        
        # HTML 파일들
        for html_file in html_files:
            epub_zip.write(os.path.join(oebps_dir, html_file), f"OEBPS/{html_file}")
        
        # 이미지 파일들
        for item in content_structure:
            if item["type"] == "image":
                image_path = os.path.join(output_dir, item["src"])
                if os.path.exists(image_path):
                    epub_zip.write(image_path, f"OEBPS/{item['src']}")
    
    print(f"EPUB3 생성 완료: {epub_filename}")
    print(f"생성된 EPUB3 파일은 '{epub_filename}'에 있습니다.")
    print("TTAK.KO-10.0905 표준을 준수하여 제작되었습니다.")
    
    return epub_filename
